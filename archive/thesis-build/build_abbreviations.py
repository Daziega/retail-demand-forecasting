"""Insert a List of Abbreviations into the front matter, between the List of
Figures and Acknowledgements, matching the existing front-matter page-break
pattern (a run-level <w:br type="page"/> inside an empty paragraph)."""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = "/Users/desmond/Capstone Project/retail-demand-forecasting/TFM_Final_formatted.docx"

ABBREVIATIONS = [
    ("AI", "Artificial Intelligence"),
    ("ANN", "Artificial Neural Network"),
    ("ARIMA", "Autoregressive Integrated Moving Average"),
    ("BI", "Business Intelligence (as in Power BI)"),
    ("CNN", "Convolutional Neural Network"),
    ("CPI", "Consumer Price Index"),
    ("CRISP-DM", "Cross-Industry Standard Process for Data Mining"),
    ("CSV", "Comma-Separated Values"),
    ("DRL", "Deep Reinforcement Learning"),
    ("EOQ", "Economic Order Quantity"),
    ("ERP", "Enterprise Resource Planning"),
    ("GNN", "Graph Neural Network"),
    ("GPU", "Graphics Processing Unit"),
    ("HC", "Holding Cost"),
    ("KNN", "K-Nearest Neighbours"),
    ("KPI", "Key Performance Indicator"),
    ("L1 / L2", "L1/L2 regularisation (Lasso/Ridge penalty terms)"),
    ("LSTM", "Long Short-Term Memory (neural network architecture)"),
    ("M5", "M5-Forecasting (Walmart) dataset / competition"),
    ("MA", "Moving Average"),
    ("MAE", "Mean Absolute Error"),
    ("MAPE", "Mean Absolute Percentage Error"),
    ("ML", "Machine Learning"),
    ("MLP", "Multi-Layer Perceptron"),
    ("MSE", "Mean Squared Error"),
    ("OC", "Ordering Cost"),
    ("OUTL", "Order-Up-To-Level (inventory review policy)"),
    ("P10 / P50 / P90 / P95 / P99", "10th / 50th / 90th / 95th / 99th percentile of the forecast demand distribution"),
    ("POS", "Point-of-Sale"),
    ("RAM", "Random Access Memory"),
    ("RF", "Random Forest"),
    ("RMDN", "Recurrent Mixture Density Network"),
    ("RMSE", "Root Mean Squared Error"),
    ("SARIMA / SARIMAX", "Seasonal ARIMA / Seasonal ARIMA with eXogenous variables"),
    ("SBA", "Syntetos–Boylan Approximation (as in Croston-SBA)"),
    ("SC", "Stockout Cost"),
    ("SES", "Simple Exponential Smoothing"),
    ("SHAP", "SHapley Additive exPlanations"),
    ("SKU", "Stock-Keeping Unit"),
    ("SL", "Service Level"),
    ("SME", "Small and Medium-sized Enterprise"),
    ("SNAP", "Supplemental Nutrition Assistance Program"),
    ("SS", "Safety Stock"),
    ("SVM / SVR", "Support Vector Machine / Support Vector Regression"),
    ("TC", "Total Cost"),
    ("TFT", "Temporal Fusion Transformer"),
    ("WMAPE", "Weighted Mean Absolute Percentage Error"),
]


def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def make_page_break_paragraph(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r._r.append(br)
    return p


def main():
    doc = Document(SRC)

    # Locate the run-level page-break paragraph that currently separates
    # "List of Figures" content from the "Acknowledgements" heading.
    ack_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "Acknowledgements" and p.style.name == "Heading 1":
            ack_idx = i
            break
    if ack_idx is None:
        raise RuntimeError("Could not locate 'Acknowledgements' heading")

    # The immediately preceding paragraph should be the page-break separator.
    separator_p = doc.paragraphs[ack_idx - 1]._p

    # ---- Build the new content as a scratch area at the end of the doc,
    # then relocate each element to just after the separator (in order). ----
    heading = doc.add_heading("List of Abbreviations", level=1)
    for r in heading.runs:
        r.font.name = "Arial"

    table = doc.add_table(rows=1 + len(ABBREVIATIONS), cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(["Abbreviation", "Definition"]):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(10)
        set_cell_shading(hdr[i], "D5E8F0")
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for r_idx, (abbr, definition) in enumerate(ABBREVIATIONS, start=1):
        cells = table.rows[r_idx].cells
        cells[0].text = ""
        run0 = cells[0].paragraphs[0].add_run(abbr)
        run0.bold = True
        run0.font.name = "Arial"
        run0.font.size = Pt(10)
        cells[1].text = ""
        run1 = cells[1].paragraphs[0].add_run(definition)
        run1.font.name = "Arial"
        run1.font.size = Pt(10)

    trailing_break = make_page_break_paragraph(doc)

    # ---- Relocate: heading, table, trailing_break -> right after separator_p ----
    anchor = separator_p
    for element in [heading._p, table._tbl, trailing_break._p]:
        anchor.addnext(element)
        anchor = element

    doc.save(SRC)
    print(f"Saved: {SRC}")
    print(f"Inserted List of Abbreviations with {len(ABBREVIATIONS)} entries "
          f"between List of Figures and Acknowledgements.")


if __name__ == "__main__":
    main()
