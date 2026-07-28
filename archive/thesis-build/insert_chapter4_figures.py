"""Insert SHAP, model-comparison, and fairness figures into Chapter 4.

Strategy: open Chapter_4_Model_Results.docx, find each target section heading,
insert figure(s) + caption just before the next section heading.
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy

SRC = "/Users/desmond/Capstone Project/retail-demand-forecasting/Chapter_4_Model_Results.docx"
OUT = "/Users/desmond/Capstone Project/retail-demand-forecasting/Chapter_4_Model_Results_with_figures.docx"
FIG_DIR = "/Users/desmond/Capstone Project/retail-demand-forecasting/figures"

# Map of "insert before this heading" → list of (figure path, caption, width inches)
INSERTIONS = {
    "4.10 ": [
        (
            f"{FIG_DIR}/model_comparison.png",
            "Figure 4.1. Consolidated model comparison across all twelve forecasting "
            "approaches evaluated. Bars show MAE on the test set (lower is better), "
            "ordered from best to worst. The stacking ensemble (STACK_Ridge_trees) "
            "achieves the lowest MAE at 0.952, narrowly beating the three individual "
            "tree-based ML models. The LSTM and Seasonal Naive baselines are the "
            "weakest performers; Croston-SBA leads the classical baselines.",
            6.0,
        ),
    ],
    "4.12 ": [
        (
            f"{FIG_DIR}/shap_importance.png",
            "Figure 4.2. SHAP feature importance ranking for the LightGBM model "
            "(mean absolute SHAP values, computed on a 5,000-row stratified sample of "
            "the test set). The top five features — rolling_mean_7, rolling_mean_28, "
            "item_month_mean, rolling_std_7, and item_dow_mean — together account for "
            "78.6% of model behaviour, confirming that engineered temporal statistics "
            "and target-encoded item-level seasonality are the dominant predictive "
            "drivers.",
            6.0,
        ),
        (
            f"{FIG_DIR}/shap_beeswarm.png",
            "Figure 4.3. SHAP beeswarm plot showing the directional effect of each "
            "feature on individual predictions. Each point represents one observation; "
            "horizontal position shows the SHAP value (positive = higher predicted "
            "demand) and colour encodes the feature value (red = high, blue = low). "
            "The plot reveals that high recent demand (rolling_mean_7) and high item-"
            "specific historical levels push predictions upward, while is_snap_day and "
            "weekend indicators contribute meaningful but smaller signals.",
            6.0,
        ),
    ],
    "4.13 ": [
        (
            f"{FIG_DIR}/fairness_check.png",
            "Figure 4.4. Performance fairness analysis. WMAPE is reported by product "
            "category (left) and by state (right). Variation across categories (26.8% "
            "coefficient of variation) and across states (19.8%) reflects underlying "
            "demand structure (intermittence levels, demand volumes) rather than "
            "systematic algorithmic bias: no category or state performs more than 2× "
            "worse than another on any metric. Foods, the highest-volume category, has "
            "the highest absolute MAE but the best WMAPE and Pred(10%).",
            6.0,
        ),
    ],
}


def main():
    doc = Document(SRC)

    # Index headings by their prefix
    heading_index = {}
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        style = p.style.name if p.style else ""
        if "Heading" in style and text:
            for prefix in INSERTIONS:
                if text.startswith(prefix):
                    heading_index[prefix] = i
                    break

    print("Located headings:")
    for prefix, idx in heading_index.items():
        print(f"  '{prefix}...' at paragraph {idx}")

    # For each target heading, build figure paragraphs and insert before that heading
    total_inserted = 0
    for heading_prefix, figs in INSERTIONS.items():
        if heading_prefix not in heading_index:
            print(f"  WARNING: heading prefix '{heading_prefix}' not found — skipping")
            continue
        target_p = doc.paragraphs[heading_index[heading_prefix]]._p

        for img_path, caption, width in figs:
            # Create image paragraph (centered)
            img_para = doc.add_paragraph()
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = img_para.add_run()
            run.add_picture(img_path, width=Inches(width))

            # Create caption paragraph (centered, italic, small)
            cap_para = doc.add_paragraph()
            cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_para.paragraph_format.space_after = Pt(18)
            cap_run = cap_para.add_run(caption)
            cap_run.italic = True
            cap_run.font.name = "Arial"
            cap_run.font.size = Pt(10)

            # Move to just before the target heading
            target_p.addprevious(deepcopy(img_para._p))
            target_p.addprevious(deepcopy(cap_para._p))
            # Remove the duplicates we created at end
            img_para._p.getparent().remove(img_para._p)
            cap_para._p.getparent().remove(cap_para._p)
            total_inserted += 1

    doc.save(OUT)
    print(f"\nSaved: {OUT}")
    print(f"Inserted {total_inserted} figures.")


if __name__ == "__main__":
    main()
