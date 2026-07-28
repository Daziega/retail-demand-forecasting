"""Build the thesis front matter: title page, abstract, resumen, TOC, list of
tables, list of figures.

The TOC / list-of-tables / list-of-figures use Word field codes — they
auto-populate when the user opens the final combined document and right-clicks
'Update field' (or presses F9).
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/Users/desmond/Capstone Project/retail-demand-forecasting/TFM_Front_Matter.docx"


def add_field(paragraph, instr):
    """Insert a Word field (TOC, TOF, etc.) into a paragraph."""
    run = paragraph.add_run()
    r = run._r
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.set(qn("xml:space"), "preserve")
    instrText.text = instr
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "separate")
    fldChar3 = OxmlElement("w:t")
    fldChar3.text = (
        "Right-click here in Word and choose 'Update Field' "
        "after assembling the full thesis document."
    )
    fldChar4 = OxmlElement("w:fldChar")
    fldChar4.set(qn("w:fldCharType"), "end")
    r.append(fldChar1)
    r.append(instrText)
    r.append(fldChar2)
    r.append(fldChar3)
    r.append(fldChar4)


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def add_centered(doc, text, size, bold=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_para(doc, text, *, bold=False, italic=False, size=11, space_after=8,
             alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


# ─────────────────────────────────────────────────────────────────────────────
doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(11)

# ── TITLE PAGE ────────────────────────────────────────────────────────────
for _ in range(6):
    doc.add_paragraph()  # vertical spacing

add_centered(doc, "Master's Programme in Data Science", 14, italic=True)
add_centered(doc, "Spain  |  2026", 12, italic=True, space_after=24)

add_centered(doc, "DEMAND FORECASTING AND INVENTORY", 22, bold=True)
add_centered(doc, "OPTIMISATION USING MACHINE LEARNING", 22, bold=True, space_after=18)
add_centered(doc,
    "A simulation-based evaluation on the M5-Forecasting (Walmart) dataset, "
    "with an SME-accessible Power BI dashboard.",
    13, italic=True, space_after=36)

add_centered(doc, "Author", 11, bold=True, space_after=4)
add_centered(doc, "Desmond Korbla Aziega", 13, space_after=18)

add_centered(doc, "Supervisor", 11, bold=True, space_after=4)
add_centered(doc, "[Supervisor name]", 13, space_after=36)

add_centered(doc, "Final Master's Thesis (TFM)", 11, italic=True)
add_centered(doc, "Master's Programme in Data Science", 11, italic=True)
add_centered(doc, "[Institution name]", 11, italic=True, space_after=4)

add_page_break(doc)

# ── ABSTRACT (English) ───────────────────────────────────────────────────
add_heading(doc, "Abstract", 1)
add_para(doc,
    "Retail demand forecasting and inventory optimisation are central to operations "
    "management, yet the operational value of machine-learning forecasting is more "
    "often asserted than measured. This thesis develops and evaluates an end-to-end "
    "forecasting-to-inventory framework on the M5-Forecasting (Walmart) dataset — a "
    "stratified subsample of 502 product–store series of daily, intermittent SKU-"
    "level demand — comparing classical statistical baselines, gradient-boosting and "
    "deep-learning models, a stacking ensemble, and LightGBM quantile regression, and "
    "integrating their outputs with a periodic (R, s, S) Order-Up-To-Level inventory "
    "policy. On forecast accuracy, the stacking ensemble is the most accurate model "
    "but improves on the strongest classical baseline by only about 4.5% (MAE), well "
    "below the 10–25% reported in prior work. Critically, when inventory outcomes are "
    "measured by forward simulation over the real test window rather than assumed, all "
    "policies achieve full service with no stockouts, so neither the more accurate "
    "point forecast nor a quantile-based safety-stock reformulation lowers inventory "
    "cost: the apparent saving reported under assumed service levels does not survive "
    "measurement. The thesis interprets this as a finding rather than a failure. The "
    "10–25% and 5–15% benchmarks of the literature were obtained on aggregated data "
    "and do not transfer to SKU-daily intermittent demand, and the operational value "
    "of ML forecasting is conditional on the demand regime — material only where the "
    "baseline policy faces genuine stockout risk. The contribution is a rigorous, "
    "reproducible, simulation-based characterisation of when ML forecasting and "
    "distributional safety stock affect inventory outcomes, delivered through an "
    "accessible Power BI dashboard designed for SME retailers.")
add_para(doc,
    "Keywords: demand forecasting; machine learning; inventory optimisation; "
    "intermittent demand; M5; safety stock; quantile regression.",
    italic=True, size=10, space_after=6)

add_page_break(doc)

# ── RESUMEN (Spanish) ────────────────────────────────────────────────────
add_heading(doc, "Resumen", 1)
add_para(doc,
    "La previsión de la demanda y la optimización de inventario son centrales en la "
    "gestión de operaciones del comercio minorista, pero el valor operativo de la "
    "previsión basada en aprendizaje automático se suele afirmar más que medir. Esta "
    "tesis desarrolla y evalúa un marco integral de previsión e inventario sobre el "
    "conjunto de datos M5-Forecasting (Walmart) —una submuestra estratificada de 502 "
    "series producto–tienda de demanda diaria, intermitente y a nivel de SKU—, "
    "comparando modelos estadísticos clásicos, modelos de gradient boosting y "
    "aprendizaje profundo, un ensemble por apilamiento (stacking) y regresión "
    "cuantílica con LightGBM, e integrando sus salidas con una política periódica "
    "(R, s, S) de nivel de reposición (Order-Up-To-Level). En precisión de previsión, "
    "el ensemble es el modelo más preciso, pero mejora al mejor modelo clásico en "
    "apenas un 4,5% (MAE), muy por debajo del 10–25% reportado en trabajos previos. "
    "De forma decisiva, cuando los resultados de inventario se miden mediante "
    "simulación hacia delante sobre la ventana de test real, en lugar de asumirse, "
    "todas las políticas alcanzan servicio completo sin roturas de stock, de modo que "
    "ni la mayor precisión de la previsión ni la reformulación cuantílica del stock "
    "de seguridad reducen el coste de inventario: el ahorro aparente obtenido bajo "
    "niveles de servicio asumidos no sobrevive a la medición. La tesis interpreta esto "
    "como un hallazgo, no como un fracaso. Los umbrales del 10–25% y del 5–15% de la "
    "literatura se obtuvieron sobre datos agregados y no se trasladan a la demanda "
    "intermitente a nivel SKU-día, y el valor operativo de la previsión con ML es "
    "condicional al régimen de demanda, relevante solo cuando la política base se "
    "expone a un riesgo real de rotura. La contribución es una caracterización "
    "rigurosa, reproducible y basada en simulación de cuándo la previsión con ML y el "
    "stock de seguridad distribucional afectan a los resultados de inventario, "
    "presentada mediante un cuadro de mando accesible en Power BI diseñado para pymes "
    "minoristas.")
add_para(doc,
    "Palabras clave: previsión de la demanda; aprendizaje automático; optimización de "
    "inventario; demanda intermitente; M5; stock de seguridad; regresión cuantílica.",
    italic=True, size=10, space_after=6)

add_page_break(doc)

# ── TABLE OF CONTENTS ────────────────────────────────────────────────────
add_heading(doc, "Table of Contents", 1)
add_para(doc,
    "(Place cursor below this line, right-click, and choose 'Update Field' once the "
    "full thesis is assembled. The TOC will auto-populate from the Heading 1, 2, and "
    "3 styles of the combined document.)",
    italic=True, size=9, space_after=12, alignment=WD_ALIGN_PARAGRAPH.LEFT)

toc_para = doc.add_paragraph()
add_field(toc_para, r'TOC \o "1-3" \h \z \u')

add_page_break(doc)

# ── LIST OF TABLES ───────────────────────────────────────────────────────
add_heading(doc, "List of Tables", 1)
add_para(doc,
    "(Right-click and 'Update Field' after assembly. Populated from Word caption "
    "items labelled 'Table'.)",
    italic=True, size=9, space_after=12, alignment=WD_ALIGN_PARAGRAPH.LEFT)

tot_para = doc.add_paragraph()
add_field(tot_para, r'TOC \h \z \c "Table"')

add_page_break(doc)

# ── LIST OF FIGURES ──────────────────────────────────────────────────────
add_heading(doc, "List of Figures", 1)
add_para(doc,
    "(Right-click and 'Update Field' after assembly. Populated from Word caption "
    "items labelled 'Figure'.)",
    italic=True, size=9, space_after=12, alignment=WD_ALIGN_PARAGRAPH.LEFT)

tof_para = doc.add_paragraph()
add_field(tof_para, r'TOC \h \z \c "Figure"')

add_page_break(doc)

# ── ACKNOWLEDGEMENTS placeholder ─────────────────────────────────────────
add_heading(doc, "Acknowledgements", 1)
add_para(doc,
    "[Placeholder — add personal acknowledgements to supervisor, family, peers, etc. "
    "Keep to half a page.]",
    italic=True)

add_page_break(doc)

doc.save(OUT)
print(f"Saved: {OUT}")
