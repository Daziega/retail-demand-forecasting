"""Generate Chapter 6 - Discussion."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_para(doc, text, *, italic=False, bold=False, indent_left=None, font_size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if indent_left:
        p.paragraph_format.left_indent = Inches(indent_left)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(font_size)
    run.italic = italic
    run.bold = bold
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    return p


def add_table_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(10)


def add_simple_table(doc, headers, rows, header_fill="D5E8F0"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.autofit = True

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(10)
        set_cell_shading(hdr_cells[i], header_fill)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for r_idx, row in enumerate(rows, start=1):
        cells = table.rows[r_idx].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ""
            p = cells[c_idx].paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = "Arial"
            run.font.size = Pt(10)
            cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return table


# ─────────────────────────────────────────────────────────────────────────────
doc = Document()

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(11)

# ── TITLE ─────────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("CHAPTER 6")
run.bold = True
run.font.name = "Arial"
run.font.size = Pt(20)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Discussion")
run.bold = True
run.font.name = "Arial"
run.font.size = Pt(16)
subtitle.paragraph_format.space_after = Pt(18)

# ── 6.1 INTRODUCTION ─────────────────────────────────────────────────────
add_heading(doc, "6.1 Introduction", 1)
add_para(doc,
    "This chapter discusses the empirical findings of the thesis in the context of the research "
    "gaps identified in Chapter 2 and the broader academic literature. The discussion proceeds "
    "in seven parts. Section 6.2 synthesises the principal results from Chapters 4 and 5. "
    "Section 6.3 revisits each of the four research gaps from the literature review and "
    "evaluates the extent to which the thesis has closed them. Section 6.4 positions the "
    "findings against the comparable studies referenced in Chapter 2. Sections 6.5 and 6.6 "
    "discuss the theoretical and practical contributions of the work. Section 6.7 addresses "
    "validity threats. Section 6.8 expands on the limitations introduced in Chapters 4 and 5. "
    "Section 6.9 identifies promising directions for future research. Section 6.10 summarises "
    "the chapter.")

# ── 6.2 SYNTHESIS ────────────────────────────────────────────────────────
add_heading(doc, "6.2 Synthesis of Empirical Findings", 1)
add_para(doc,
    "Three findings emerge from the empirical work and form the basis for the discussion that "
    "follows.")
add_para(doc,
    "Finding 1: Forecast accuracy improvements at SKU-daily granularity are modest. The "
    "stacking ensemble of LightGBM, XGBoost, and Random Forest base learners combined through "
    "a non-negative Ridge meta-learner achieves the best point-forecast accuracy across all "
    "twelve models evaluated (MAE 0.952, WMAPE 73.58%). Improvements over the best classical "
    "baselines are 4.5% on MAE, 9.5% on RMSE, and approximately flat on MAPE. These figures "
    "sit below the 10–25% target band specified in Chapter 1 — an outcome that, while "
    "initially appearing as a shortfall, is consistent with the granularity-dependence "
    "documented in the M5 retrospective literature. The headline accuracy figures reported by "
    "Barghi (2025) and Seyedan et al. (2023) are obtained at weekly or category-level "
    "aggregation, where averaging across observations softens the relative weight of zero-"
    "actual days that distort per-row percentage metrics.")
add_para(doc,
    "Finding 2: The Gaussian safety-stock formula systematically under-allocates buffer "
    "inventory for intermittent demand. The empirical 95th-percentile demand spread, computed "
    "from the LightGBM quantile-regression model, is 1.36 times the Gaussian-implied buffer at "
    "the same nominal service level. The Gaussian formula, by assuming normality on a right-"
    "skewed and zero-inflated demand distribution, prescribes a buffer that covers "
    "approximately the 83rd percentile rather than the targeted 95th. This phenomenon is "
    "implicit in the textbook treatment of safety stock (Silver, Pyke & Thomas, 2017) but has "
    "not, to the author's knowledge, been quantitatively measured on a standard retail "
    "benchmark such as M5 in prior literature.")
add_para(doc,
    "Finding 3: Operational cost reduction is substantial and robust. When the Gaussian under-"
    "allocation gap is closed through quantile-based safety-stock calibration, total annual "
    "inventory cost reduces by 21% at the central scenario (L = 14, SL = 95%, baseline "
    "stockout cost) and is robust across all four lead-time scenarios specified in Chapter 1 "
    "(20.8–21.2%). Sensitivity analysis across MAPE (3–15%), stockout cost multiplier (0.4 to "
    "2.0×), and service level (90% to 99%) confirms that the result exceeds the Chapter 1 "
    "target of 5–15% even under the most conservative stockout cost assumption, and "
    "demonstrates that the framework's value scales most strongly with service level — making "
    "it most valuable for high-margin or reputation-critical SME inventory.")

# ── 6.3 RESEARCH GAPS ────────────────────────────────────────────────────
add_heading(doc, "6.3 Revisiting the Four Research Gaps", 1)
add_para(doc,
    "The literature review in Chapter 2 identified four research gaps that motivated the "
    "design of the thesis. This section returns to each gap and evaluates the extent to which "
    "the empirical work has closed it.")

add_heading(doc, "6.3.1 Gap 1 — Disconnect between forecast accuracy and inventory decisions", 2)
add_para(doc,
    "Gap 1, articulated explicitly by Seyedan et al. (2023), identifies that most comparative "
    "forecasting studies report accuracy metrics without translating improved forecasts into "
    "operational inventory outcomes. This thesis addresses Gap 1 directly through the OUTL "
    "integration and the quantile-based safety stock formulation in Chapter 5.")
add_para(doc,
    "The findings provide quantitative evidence of why the gap matters operationally. When ML "
    "forecasts are fed into the same Gaussian safety-stock formula as classical forecasts, the "
    "operational benefit is negligible (the ML-Gaussian policy saves approximately $20 per "
    "year per 502-series test set, or 0.02% of total annual cost). The forecast-to-inventory "
    "disconnect is not closed by improving forecast accuracy within the conventional Gaussian "
    "formulation; it is closed by changing the formulation itself to incorporate the empirical "
    "demand distribution that quantile regression makes accessible. This is a specific, "
    "actionable refinement of the broader gap identified by Seyedan et al. (2023): the "
    "missing methodological step is not better forecasts but a different mapping from forecast "
    "to inventory.")
add_para(doc,
    "This thesis therefore contributes to closing Gap 1 in two complementary ways. First, it "
    "provides the first quantitative measurement on the M5 benchmark of the magnitude of the "
    "Gaussian under-allocation problem (a factor of 1.36, or equivalently a ~12 percentage-"
    "point service-level under-shoot). Second, it provides a concrete, operationally deployable "
    "alternative — LightGBM quantile regression feeding empirically calibrated safety stock — "
    "and quantifies its cost impact across realistic operational scenarios.")

add_heading(doc, "6.3.2 Gap 2 — Real-world deployment and operational validation", 2)
add_para(doc,
    "Gap 2 concerns the limited focus in the literature on real-world deployment, model "
    "maintenance, and the practical constraints of live operation. The literature is dominated "
    "by historical-simulation evaluations that do not address the deployment layer.")
add_para(doc,
    "The thesis addresses Gap 2 through three concrete deliverables. First, the end-to-end "
    "pipeline produces a set of CSV data feeds that are designed for daily refresh in a "
    "production deployment — five files exported from the Python pipeline corresponding to the "
    "four dashboard views and a KPI summary. Second, the Power BI dashboard provides an "
    "interactive, accessible interface for non-technical retail managers, organising the "
    "framework's outputs into four operational views (Forecast Overview, Inventory Status, Cost "
    "Analysis, Sensitivity Analysis) and exposing the underlying parameters as configurable "
    "filters. Third, the four-dimensional sensitivity analysis quantifies how the operational "
    "result varies with the principal parameters that practitioners would actually adjust in "
    "deployment — lead time, service level, stockout cost, and forecast accuracy.")
add_para(doc,
    "While the thesis does not implement live deployment or model maintenance routines, the "
    "framework is structurally deployable: data feeds are CSV-based and refresh on a daily "
    "cycle, the modelling pipeline is reproducible from raw inputs, and the SME transferability "
    "specification in §3.11 documents the minimum data, tooling, hardware, and skills required "
    "to operate the framework outside a research environment. The gap is not fully closed — "
    "live deployment and incremental model updating are out of scope — but the framework is "
    "delivered in a form that significantly closes the distance between the published academic "
    "literature and operational deployment.")

add_heading(doc, "6.3.3 Gap 3 — SME accessibility and scalability", 2)
add_para(doc,
    "Gap 3 concerns the dominance in the literature of frameworks designed for enterprise-scale "
    "analytics infrastructure, with limited attention to whether the same methodologies are "
    "viable for small and medium-sized retailers operating with constrained resources. The "
    "thesis addresses this gap through the explicit SME transferability commitment of Chapter 1 "
    "and the corresponding specification in §3.11.")
add_para(doc,
    "Three specific design choices implement this commitment. First, the entire pipeline is "
    "built on free and open-source tooling — Python with mainstream libraries (pandas, scikit-"
    "learn, lightgbm, xgboost, tensorflow, statsforecast, shap) and Microsoft Power BI Desktop "
    "free tier. No commercial software licence or cloud-computing subscription is required. "
    "Second, the framework runs in under three hours on a standard 16 GB-RAM laptop with no "
    "GPU. Third, the data inputs required — daily SKU-store sales, weekly prices, and a "
    "calendar — are standard outputs of off-the-shelf SME point-of-sale systems, ensuring that "
    "no additional data infrastructure investment is needed.")
add_para(doc,
    "An important nuance is that SME transferability is not simply a tooling question but also "
    "a methodological one: the framework must work on the scale of inventory that SMEs operate, "
    "not just on the scale of enterprise benchmarks. The stratified subsample of 502 product-"
    "store series used in the thesis is a deliberate choice in this respect — a representative "
    "SME inventory might consist of 200–500 SKUs across a small number of locations, which the "
    "framework handles within a runtime envelope that permits daily refresh. The framework is "
    "therefore not merely portable to SME environments; it is sized for them.")

add_heading(doc, "6.3.4 Gap 4 — Responsible AI and model transparency", 2)
add_para(doc,
    "Gap 4 concerns the limited treatment of responsible AI principles — transparency, "
    "fairness, and data minimisation — in the retail forecasting literature. Ungureanu (2025) "
    "highlights this gap as under-addressed in comparable forecasting studies.")
add_para(doc,
    "The thesis addresses Gap 4 through three elements documented in §3.10 and reported in "
    "§4.11 and §4.12. Model transparency is established through SHAP analysis of the best-"
    "performing model, identifying the principal demand drivers (rolling means, target-encoded "
    "item-level seasonality, and demand volatility) and quantifying their relative "
    "contribution. The top five SHAP-ranked features account for 78.6% of model predictions, "
    "making the model interpretable at the feature level. Performance fairness is established "
    "through disaggregated error metrics by product category (Foods, Hobbies, Household) and "
    "by state (California, Texas, Wisconsin), with the observation that performance variation "
    "across these groups reflects underlying demand structure (intermittence levels, demand "
    "volumes) rather than systematic algorithmic bias. Data minimisation is implemented by "
    "retaining only features with non-trivial SHAP contribution.")
add_para(doc,
    "A consideration worth flagging is that the fairness analysis reveals coefficient-of-"
    "variation figures (26.8% for category, 19.8% for state) that initially appear above "
    "conventional algorithmic-bias thresholds. The discussion in §4.12 establishes that the "
    "variation is explained by structural differences in the underlying demand distributions — "
    "Foods has higher absolute MAE because its absolute sales volume is higher, not because "
    "the model treats Foods worse — and the fairness check ultimately passes. Future work "
    "should consider extending the fairness analysis to additional dimensions such as series "
    "intermittence level, which would isolate algorithmic behaviour from data structure more "
    "cleanly.")

# ── 6.4 LITERATURE COMPARISON ────────────────────────────────────────────
add_heading(doc, "6.4 Comparison with the Literature", 1)
add_para(doc,
    "This section positions the thesis findings against the principal comparable studies "
    "reviewed in Chapter 2. The comparison focuses on three axes: forecast accuracy, modelling "
    "architecture, and inventory cost impact.")
add_para(doc,
    "Forecast accuracy. Barghi (2025) reports MAPE of 6.48% (hybrid ensemble) reduced to 3.53% "
    "with fuzzy generalisation preprocessing, on the Rossmann Store Sales dataset at weekly "
    "aggregation. Seyedan et al. (2023) report MAPE of 5.22% (sports) and 9.58% (electronics) "
    "with their heterogeneous MLP-LSTM-CNN stacking ensemble on weekly retail data. Nasseri et "
    "al. (2023) report MAPE in a similar range with Extra Tree Regressors on perishable retail "
    "categories. The MAPE figures in this thesis — 53–58% across the model tiers at SKU-daily "
    "granularity — are an order of magnitude higher in absolute terms but consistent with the "
    "M5 retrospective literature at the same granularity. Single-digit MAPE on M5-style data "
    "is achievable only at weekly or category-level aggregation; the SKU-daily granularity "
    "used in this thesis is methodologically more demanding and operationally more "
    "informative, because inventory decisions are made at the SKU-store-day level rather than "
    "at aggregated levels.")
add_para(doc,
    "Modelling architecture. Barghi (2025) develops a sequential residual-correction ensemble "
    "(LightGBM → CatBoost → XGBoost) with a learning rate of 0.05 across all three stages. "
    "This thesis explicitly considered and rejected the sequential residual-correction design "
    "in favour of conventional stacking with a Ridge meta-learner, on the methodological "
    "grounds that sequential residual chaining tends to overfit successive layers to the noise "
    "of earlier layers and that stacking with non-negative weights is the more conventional "
    "heterogeneous ensemble architecture in the modern forecasting literature (Mitra et al., "
    "2022; Seyedan et al., 2023). The empirical results support this choice: the stacking "
    "ensemble achieves the best MAE and WMAPE among all twelve models evaluated.")
add_para(doc,
    "Inventory cost impact. Seyedan et al. (2023) report inventory cost reductions in the "
    "range of 8–14% for their ensemble against single-model baselines, on a weekly aggregated "
    "dataset. The thesis's headline figure of 21% (baseline stockout multiplier) and "
    "robustness range of 15–37% (across the stockout cost sensitivity) exceeds the Seyedan et "
    "al. result in absolute terms. The explanation lies in the mechanism: Seyedan et al. "
    "attribute their savings to improved point-forecast accuracy feeding the Gaussian formula, "
    "whereas the present thesis attributes its savings to the empirical quantile reformulation "
    "of the safety-stock calculation itself. The two mechanisms are complementary, and a "
    "future extension would combine them by applying the quantile reformulation to the "
    "Seyedan stacking ensemble.")

# ── 6.5 THEORETICAL CONTRIBUTIONS ────────────────────────────────────────
add_heading(doc, "6.5 Theoretical Contributions", 1)
add_para(doc,
    "The thesis makes three theoretical contributions.")
add_para(doc,
    "First, it provides the first quantitative empirical measurement on the M5 benchmark of "
    "the Gaussian safety-stock under-allocation gap for intermittent demand. The 1.36 factor "
    "between empirical 95th-percentile demand spread and the Gaussian-implied buffer is, to the "
    "author's knowledge, not previously reported in the M5 literature. This finding is "
    "directly relevant to the broader inventory theory literature, which acknowledges the "
    "Gaussian assumption's validity conditions (Silver et al., 2017) but rarely measures the "
    "gap on a standard benchmark.")
add_para(doc,
    "Second, the thesis demonstrates that the operational value of ML forecasting in retail "
    "inventory management is mediated by the inventory formulation, not by the forecasting "
    "model alone. The negligible improvement of the ML-Gaussian policy over Classical-Gaussian "
    "(0.02% of annual cost) and the substantial improvement of ML-Empirical-Quantile over "
    "Classical-Gaussian (21%) jointly establish that the inventory-formulation choice dominates "
    "the forecasting-model choice for operational cost outcomes. This is a refinement of the "
    "literature's common framing of forecasting-inventory integration as primarily a "
    "forecasting problem.")
add_para(doc,
    "Third, the thesis shows that the operational value of forecast-inventory integration is "
    "robust to MAPE level. Across the 3–15% MAPE sensitivity range, the cost reduction varies "
    "by less than 0.1 percentage points. The practical implication is that the marginal "
    "operational value of further forecast accuracy improvement, beyond the level needed to "
    "support the quantile-regression model, is small. This finding suggests a reordering of "
    "research priorities in retail demand forecasting: the next decimal point of MAPE is less "
    "operationally valuable than improvements in the inventory-decision layer that translate "
    "any forecast (with or without ML) into stock-keeping decisions.")

# ── 6.6 PRACTICAL CONTRIBUTIONS ──────────────────────────────────────────
add_heading(doc, "6.6 Practical Contributions", 1)
add_para(doc,
    "Three practical contributions complement the theoretical findings.")
add_para(doc,
    "An end-to-end, reproducible, SME-deployable framework. The thesis delivers a complete "
    "workflow from raw M5 ingestion through to the Power BI dashboard, implemented in free and "
    "open-source tooling and runnable on a standard laptop. The framework is documented to the "
    "level of detail required for independent replication, including the stratified subsample "
    "construction (§3.3.3), the chronological train/validation/test split (§3.6.2), the "
    "feature engineering specification (§3.5), the model configurations (§3.6.1), and the OUTL "
    "calculations (§3.7).")
add_para(doc,
    "A decision-support dashboard for non-technical managers. The Power BI dashboard "
    "translates the modelling outputs into four operational views that surface the forecast, "
    "the inventory parameters, the cost decomposition, and the sensitivity-analysis curves. "
    "The dashboard supports practitioner what-if analysis — changing the service-level target, "
    "the assumed lead time, or the stockout cost multiplier and observing the impact on "
    "expected savings — without requiring the user to understand the underlying methodology.")
add_para(doc,
    "A defensible quantitative argument for the operational value of forecasting investment. "
    "By demonstrating that the headline 21% cost reduction is robust to MAPE level and "
    "achieves the Chapter 1 target band even under the most conservative stockout cost "
    "assumption (15% at m = 0.4), the thesis provides SME inventory managers with a defensible "
    "quantitative argument for adopting ML-based demand forecasting integrated with quantile-"
    "based safety-stock calibration. The sensitivity analysis curves (§5.7) provide the basis "
    "for adapting the headline figure to the specific operational parameters of any individual "
    "SME.")

# ── 6.7 VALIDITY THREATS ─────────────────────────────────────────────────
add_heading(doc, "6.7 Validity Threats", 1)
add_para(doc,
    "Three classes of validity threat are considered: internal validity, external validity, "
    "and construct validity.")
add_para(doc,
    "Internal validity concerns whether the observed cost reduction is attributable to the "
    "framework rather than to confounding factors in the implementation. The principal "
    "internal-validity threat is the chronological split between training, validation, and "
    "test sets. The methodology in §3.6.2 specifies strict chronological partitioning to "
    "eliminate data leakage from future to past, and the evaluation results in Chapters 4 and "
    "5 are computed on a hold-out test set that was not observed during training or "
    "hyperparameter tuning. A secondary threat is the target encoding of high-cardinality "
    "categorical features, which is also subject to leakage if not computed on training data "
    "only; §3.4.6 specifies that target encodings are computed exclusively on the training "
    "window and propagated unchanged to validation and test, eliminating this leakage channel.")
add_para(doc,
    "External validity concerns the generalisability of the findings to other retail contexts. "
    "The M5 dataset originates from a single large enterprise retailer (Walmart) operating "
    "primarily in three US states; the SME accessibility argument therefore rests on the "
    "transferability of the methodology, not on direct empirical validation in an SME "
    "environment. Specifically, the relative magnitudes of demand intermittence, the validity "
    "of the Gaussian under-allocation factor (1.36 in this dataset), and the cost components "
    "may differ materially in other retail contexts. The sensitivity analysis (§5.7) partially "
    "mitigates this threat by reporting the result across a wide range of parameter "
    "assumptions; a stronger mitigation would be replication on an actual SME inventory.")
add_para(doc,
    "Construct validity concerns whether the metrics used measure what they are intended to "
    "measure. The Chapter 1 MAPE improvement target is the principal construct-validity "
    "concern: at SKU-daily granularity on intermittent demand, MAPE is well-known to be an "
    "unstable metric, and the supplementary WMAPE used in this thesis is more robust for the "
    "data structure. The decision to retain MAPE while adding WMAPE as a supplementary metric "
    "preserves comparability with the literature (which uses MAPE) while providing a more "
    "honest signal for the operational interpretation. The cost reduction figure in Chapter 5 "
    "is computed analytically rather than through full inventory simulation, which is a "
    "deliberate methodological choice for tractability; the construct it measures is the "
    "expected annual cost under the analytical (R, s, S) policy, which is the construct used "
    "in Silver et al. (2017) and the related literature.")

# ── 6.8 EXTENDED LIMITATIONS ─────────────────────────────────────────────
add_heading(doc, "6.8 Extended Limitations", 1)
add_para(doc,
    "Five limitations specific to Chapter 5's inventory analysis were noted in §5.10. This "
    "section adds three further limitations spanning the thesis as a whole.")
add_para(doc,
    "First, the empirical safety stock derived from the per-day quantile spread is scaled to "
    "the (R+L)-day horizon using a √(R+L) factor under the same i.i.d. assumption as the "
    "Gaussian formula. This is a deliberate methodological choice for like-for-like comparison, "
    "but it inherits the i.i.d. limitation of the Gaussian formula. A more rigorous "
    "specification would train a quantile model directly on cumulative (R+L)-day demand "
    "rather than per-day demand, which would relax the i.i.d. assumption but at substantial "
    "computational cost. This refinement is left to future work.")
add_para(doc,
    "Second, the LSTM model in Tier 3 underperforms the tree-based models and the stacking "
    "ensemble. While this finding is consistent with the M5 retrospective literature, it does "
    "leave open the possibility that a more elaborate sequence architecture (e.g., Temporal "
    "Fusion Transformers, N-BEATS) or per-series LSTM specialisation would close the gap. The "
    "thesis evaluates only the methodology-specified LSTM architecture (128/64 stacked, "
    "dropout 0.2, 28-day lookback) and does not exhaust the deep learning design space.")
add_para(doc,
    "Third, the achieved service level under the Gaussian policy is computed by projection "
    "onto the empirical quantile distribution rather than by direct simulation of stockout "
    "events. A future extension would simulate demand drawn from the empirical distribution "
    "and explicitly count stockout events under each policy, providing a direct rather than "
    "inferred measurement of the under-allocation gap.")

# ── 6.9 FUTURE RESEARCH ──────────────────────────────────────────────────
add_heading(doc, "6.9 Implications for Future Research", 1)
add_para(doc,
    "The findings of this thesis suggest five directions for productive future research.")
add_para(doc,
    "First, the Gaussian under-allocation factor of 1.36 measured on M5 SKU-daily data should "
    "be re-measured on other retail benchmarks (Rossmann, Favorita, regional grocery datasets) "
    "to establish how the factor varies with category, granularity, and intermittence level. "
    "If the factor varies systematically with category structure, a category-specific "
    "calibration scheme could be developed.")
add_para(doc,
    "Second, the quantile-based safety stock formulation should be evaluated against more "
    "sophisticated probabilistic forecasting methods (e.g., DeepAR, Bayesian neural networks, "
    "conformal prediction intervals) to assess whether quantile regression is the right "
    "probabilistic model class for this purpose, or whether alternative probabilistic "
    "specifications yield further operational improvement.")
add_para(doc,
    "Third, the framework should be extended to incorporate live deployment routines: rolling "
    "retraining schedules, model drift detection, and the operational monitoring of forecast "
    "calibration over time. This extension would more fully close Gap 2 from Chapter 2.")
add_para(doc,
    "Fourth, the sensitivity analysis should be extended to incorporate non-linear stockout "
    "cost functions (e.g., increasing marginal cost as stockout frequency rises, reflecting "
    "customer-loyalty erosion) and supplier lead-time uncertainty. Both refinements would "
    "produce a more realistic operational cost model, although at the expense of analytical "
    "tractability.")
add_para(doc,
    "Fifth, the framework's transferability to SME environments should be validated through "
    "deployment in actual SME retail operations. The §3.11 transferability specification "
    "documents the requirements for such deployment, but an empirical case study in an SME "
    "context would substantially strengthen the practical contribution of the work.")

# ── 6.10 SUMMARY ─────────────────────────────────────────────────────────
add_heading(doc, "6.10 Chapter Summary", 1)
add_para(doc,
    "This chapter has discussed the empirical findings of Chapters 4 and 5 in the context of "
    "the four research gaps identified in Chapter 2 and the broader academic literature. Three "
    "principal findings were synthesised: forecast accuracy improvements at SKU-daily "
    "granularity are modest (~4.5% MAE, ~9.5% RMSE) and below the Chapter 1 target band; the "
    "Gaussian safety-stock formula systematically under-allocates by a factor of 1.36 on "
    "intermittent demand; and the quantile-based inventory reformulation reduces total annual "
    "cost by 21–27% across realistic operational scenarios, exceeding the Chapter 1 target "
    "band.")
add_para(doc,
    "The thesis closes Gap 1 (forecast-inventory disconnect) through the quantile-based "
    "reformulation, closes Gap 3 (SME accessibility) through tooling and infrastructure "
    "choices, and partially closes Gap 2 (deployment) through the dashboard and CSV refresh "
    "architecture and Gap 4 (responsible AI) through SHAP and fairness analysis. The literature "
    "comparison establishes that the thesis's headline cost reduction figure exceeds "
    "comparable studies in absolute magnitude, with the mechanism (quantile reformulation) "
    "complementary to the mechanism reported in those studies (improved point forecasting).")
add_para(doc,
    "Three theoretical contributions are identified: the first quantitative M5 measurement of "
    "the Gaussian under-allocation gap; the demonstration that the inventory formulation "
    "dominates the forecasting model for operational outcomes; and the empirical robustness of "
    "the result to forecast accuracy level. Three practical contributions complement these: an "
    "end-to-end deployable framework; a decision-support dashboard; and a defensible "
    "quantitative argument for SME forecasting investment. Validity threats and limitations "
    "are acknowledged, and five directions for future research are identified.")
add_para(doc,
    "Chapter 7 concludes the thesis. It returns to the research objectives stated in Chapter "
    "1, summarises the extent to which each has been achieved, and consolidates the "
    "contribution to the demand forecasting and inventory optimisation literature.")

OUT = "/Users/desmond/Capstone Project/retail-demand-forecasting/Chapter_6_Discussion.docx"
doc.save(OUT)
print(f"Saved: {OUT}")
