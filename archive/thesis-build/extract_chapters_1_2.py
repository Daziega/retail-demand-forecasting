"""Extract Chapter 1 (supervisor-revised) and Chapter 2 into clean standalone docx
files so the final assembly is a simple linear sequence."""
from docx import Document
from copy import deepcopy
from docx.shared import Pt, Inches


def extract(src_path, start_idx, end_idx_exclusive, out_path,
            skip_part_b_header=False):
    """Copy paragraphs [start_idx, end_idx_exclusive) from src_path → out_path."""
    src = Document(src_path)
    out = Document()
    # Page margins
    for section in out.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    style = out.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    body = out.element.body
    sectPr = body[-1]

    for i in range(start_idx, end_idx_exclusive):
        p = src.paragraphs[i]
        if skip_part_b_header and p.text.strip().startswith("PART B"):
            continue
        new_p = deepcopy(p._p)
        sectPr.addprevious(new_p)

    out.save(out_path)
    return out


# Chapter 1 — supervisor's revised version. The file has Part A (abstract) and
# Part B (Chapter 1). Skip the Part A portion (already in TFM_Front_Matter)
# and the "PART B" header line itself.
src = "/Users/desmond/Capstone Project/retail-demand-forecasting/TFM_Abstract_and_Chapter1_revised.docx"
doc = Document(src)
n = len(doc.paragraphs)
extract(src,
        start_idx=11,        # "PART B — CHAPTER 1 (revised; changes in red)"
        end_idx_exclusive=n,
        out_path="/Users/desmond/Capstone Project/retail-demand-forecasting/TFM_Chapter1_revised.docx",
        skip_part_b_header=True)
print(f"Saved Chapter 1 → TFM_Chapter1_revised.docx (paragraphs 11–{n})")

# Chapter 2 — from the combined thesis. Para 92 = "CHAPTER 2" header,
# para 234 = "CHAPTER 3" header (excluded).
src2 = "/Users/desmond/Capstone Project/retail-demand-forecasting/Demand_Forecast_Using_ML_full.docx"
extract(src2,
        start_idx=92,
        end_idx_exclusive=234,
        out_path="/Users/desmond/Capstone Project/retail-demand-forecasting/TFM_Chapter2.docx")
print("Saved Chapter 2 → TFM_Chapter2.docx (paragraphs 92–233)")
