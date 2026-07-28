"""Standardise Chapter 1 (and tidy Chapter 2) title banners to match Chapters
3-7, and add centred page numbers to every page."""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = "/Users/desmond/Capstone Project/retail-demand-forecasting/TFM_Final_formatted.docx"
OUT = SRC  # overwrite the single current formatted file


def style_banner_line(p, size, *, page_break=False):
    p.style = p.part.document.styles["Normal"]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    if page_break:
        p.paragraph_format.page_break_before = True
    for r in p.runs:
        r.font.name = "Arial"
        r.font.size = Pt(size)
        r.font.bold = True


def set_text(p, text):
    # collapse to a single run carrying `text`
    for r in list(p.runs):
        r._r.getparent().remove(r._r)
    p.add_run(text)


def add_footer_page_number(doc):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        # reuse first footer paragraph or create one
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        # clear it
        for r in list(p.runs):
            r._r.getparent().remove(r._r)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        def _r(kind=None, txt=None):
            r = OxmlElement("w:r")
            if kind in ("begin", "separate", "end"):
                fc = OxmlElement("w:fldChar")
                fc.set(qn("w:fldCharType"), kind)
                r.append(fc)
            elif kind == "instr":
                it = OxmlElement("w:instrText")
                it.set(qn("xml:space"), "preserve")
                it.text = txt
                r.append(it)
            else:
                t = OxmlElement("w:t")
                t.text = txt or ""
                r.append(t)
            return r

        for el in [_r("begin"), _r("instr", " PAGE "), _r("separate"),
                   _r(txt="1"), _r("end")]:
            p._p.append(el)


def main():
    doc = Document(SRC)
    paras = doc.paragraphs

    # ---- Chapter 1: convert "Chapter 1: Introduction" -> two-line banner ----
    for i, p in enumerate(paras):
        if p.text.strip() == "Chapter 1: Introduction":
            set_text(p, "CHAPTER 1")
            style_banner_line(p, 20, page_break=True)
            # add "Introduction" subtitle right after
            new_p = doc.add_paragraph()
            set_text(new_p, "Introduction")
            style_banner_line(new_p, 16)
            p._p.addnext(new_p._p)
            print("Chapter 1 banner standardised.")
            break

    # ---- Chapter 2: tidy the banner to 20pt / 16pt centred ----
    for i, p in enumerate(paras):
        if p.text.strip() == "CHAPTER 2":
            style_banner_line(p, 20, page_break=True)
            # next non-empty paragraph is the subtitle
            for j in range(i + 1, i + 4):
                if paras[j].text.strip():
                    style_banner_line(paras[j], 16)
                    break
            print("Chapter 2 banner tidied.")
            break

    # ---- Page numbers ----
    add_footer_page_number(doc)
    print("Centred page-number footer added.")

    doc.save(OUT)
    print(f"Saved: {OUT}")


if __name__ == "__main__":
    main()
