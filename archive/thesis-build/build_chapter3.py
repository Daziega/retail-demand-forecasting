"""Generate Chapter 3 - Methodology as a .docx file using python-docx."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_para(doc, text, *, italic=False, bold=False, indent_left=None, font_size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if indent_left:
        p.paragraph_format.left_indent = Inches(indent_left)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(font_size)
    run.italic = italic
    run.bold = bold
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    return p


def add_formula(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Cambria Math"
    run.font.size = Pt(11)
    run.italic = True
    return p


def add_table_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(10)


def add_simple_table(doc, headers, rows, header_fill="D5E8F0"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.autofit = True

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(10)
        set_cell_shading(hdr_cells[i], header_fill)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # Body rows
    for r_idx, row in enumerate(rows, start=1):
        cells = table.rows[r_idx].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ""
            p = cells[c_idx].paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = "Arial"
            run.font.size = Pt(10)
            cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    return table


# ─────────────────────────────────────────────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Default style
style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(11)

# ── TITLE ─────────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("CHAPTER 3")
run.bold = True
run.font.name = "Arial"
run.font.size = Pt(20)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Methodology")
run.bold = True
run.font.name = "Arial"
run.font.size = Pt(16)
subtitle.paragraph_format.space_after = Pt(18)

# ── 3.1 INTRODUCTION ─────────────────────────────────────────────────────
add_heading(doc, "3.1 Introduction", 1)
add_para(doc,
    "This chapter presents the research design, data sources, preprocessing pipeline, "
    "feature engineering strategy, modelling framework, evaluation methodology, and inventory "
    "policy integration procedure that constitute the empirical backbone of this study. Every "
    "decision documented here is anchored in the four research gaps identified in Chapter 2 and "
    "guided by the five objectives stated in Chapter 1: comparative evaluation of forecasting "
    "models, systematic feature engineering, integration of the best-performing model with the "
    "Order-Up-To-Level (OUTL) inventory policy, quantification of operational impact relative to "
    "classical baselines, and delivery of results through an accessible Power BI dashboard.")
add_para(doc,
    "The overarching methodological framework is the Cross-Industry Standard Process for Data "
    "Mining (CRISP-DM), which structures the research as an iterative, six-phase process — "
    "Business Understanding, Data Understanding, Data Preparation, Modelling, Evaluation, and "
    "Deployment — progressing from problem definition through to operational output (Barghi, "
    "2025; Ungureanu, 2025). CRISP-DM was selected over alternative frameworks because its "
    "iterative structure accommodates the cycles of feature refinement and model calibration that "
    "characterise retail demand forecasting research, because it has been adopted in closely "
    "related studies (Bastos, 2023; Ungureanu, 2025), and because its explicit deployment phase "
    "anchors every technical decision in the operational business problem.")
add_para(doc,
    "The remainder of this chapter is structured as follows. Section 3.2 describes the research "
    "design philosophy and the role of CRISP-DM. Section 3.3 provides a detailed description of "
    "the M5-Forecasting (Walmart) dataset and the stratified subsample used for modelling. "
    "Section 3.4 documents the data preprocessing pipeline. Section 3.5 presents the feature "
    "engineering strategy. Section 3.6 describes model development and the comparative evaluation "
    "framework across five model tiers. Section 3.7 details the OUTL inventory policy integration "
    "procedure, including a quantile-based robustness check and stockout cost modelling. Section "
    "3.8 outlines the Power BI dashboard development approach. Section 3.9 explains the "
    "sensitivity analysis design. Section 3.10 addresses ethical considerations governing the "
    "study. Section 3.11 specifies the SME transferability requirements. Section 3.12 summarises "
    "the chapter.")

# ── 3.2 RESEARCH DESIGN ──────────────────────────────────────────────────
add_heading(doc, "3.2 Research Design and Methodological Framework", 1)
add_heading(doc, "3.2.1 Research Philosophy", 2)
add_para(doc,
    "This study is positioned within a post-positivist research philosophy. It seeks to identify "
    "causal relationships between modelling architecture choices and operational inventory "
    "outcomes, using quantitative measurement of forecast accuracy and inventory cost as primary "
    "evidence. The study does not claim to generate universal laws but rather to produce "
    "replicable, empirically grounded findings that can inform decision-making in comparable "
    "retail environments (Seyedan, 2023). A deductive approach is adopted: hypotheses about "
    "model superiority and inventory cost reduction are formulated from the literature in "
    "Chapter 2 and tested against the M5 dataset.")
add_para(doc,
    "The research strategy is a quantitative computational experiment — a form of empirical "
    "investigation in which multiple modelling approaches are systematically compared on a shared "
    "dataset under controlled conditions (Nasseri et al., 2023; Mitra et al., 2022). This "
    "strategy was chosen over survey or case study designs because it permits direct, "
    "metric-grounded comparison of forecasting and inventory outcomes, which is precisely what "
    "the literature gaps identified in Chapter 2 require.")

add_heading(doc, "3.2.2 CRISP-DM as the Structuring Framework", 2)
add_para(doc,
    "The CRISP-DM framework organises the research into six iterative phases. Phase 1 (Business "
    "Understanding) is documented in Chapter 1. Phase 2 (Data Understanding) is addressed in "
    "Sections 3.3 and 3.4. Phase 3 (Data Preparation) covers Sections 3.4 and 3.5. Phase 4 "
    "(Modelling) is addressed in Section 3.6. Phase 5 (Evaluation) spans Sections 3.6 and 3.7 "
    "and Chapter 4. Phase 6 (Deployment) is addressed in Section 3.8 through the Power BI "
    "dashboard. The iterative character of CRISP-DM is reflected in feedback loops between "
    "feature engineering and model evaluation: features were refined based on pilot model "
    "performance, and the inventory policy formulation was extended with a quantile-based "
    "robustness check after preliminary results revealed the limitations of the Gaussian "
    "assumption on intermittent demand series.")

# ── 3.3 DATA SOURCE ──────────────────────────────────────────────────────
add_heading(doc, "3.3 Data Source: The M5-Forecasting (Walmart) Dataset", 1)
add_heading(doc, "3.3.1 Dataset Description and Provenance", 2)
add_para(doc,
    "The primary data source for this study is the M5-Forecasting (Walmart) dataset, sourced "
    "from the Kaggle platform (Obi, 2024; Haque et al., 2023). The dataset was originally "
    "prepared for the M5 Accuracy and Uncertainty competitions hosted by the Makridakis Open "
    "Forecasting Center and provides one of the most comprehensive publicly available benchmarks "
    "for hierarchical retail demand forecasting.")
add_para(doc,
    "The dataset contains daily unit sales for 3,049 individual product-store series, organised "
    "across three product categories — Hobbies, Foods, and Household — and seven product "
    "departments. Products are sold across ten Walmart stores in three US states: California "
    "(CA1–CA4), Texas (TX1–TX3), and Wisconsin (WI1–WI3). The publicly distributed "
    "sales_train_validation.csv file spans 1,913 consecutive days from 29 January 2011 to 24 "
    "April 2016. The dataset is distributed across three raw files: sales_train_validation.csv "
    "(daily unit sales), calendar.csv (date, day-of-week, SNAP eligibility flags, event labels), "
    "and sell_prices.csv (weekly selling prices per product-store combination).")

add_heading(doc, "3.3.2 Justification of Dataset Choice", 2)
add_para(doc,
    "The M5 dataset was preferred over alternatives such as the Rossmann Store Sales dataset "
    "(Barghi, 2025) or the Favorita Grocery Sales dataset (Ungureanu, 2025) for four specific "
    "reasons. First, its scale — 3,049 product-store series over more than five years — provides "
    "a richer basis for model generalisation than smaller benchmarks. Second, its hierarchical "
    "structure mirrors the organisational structure of real SME retail operations, making the "
    "methodology transferable. Third, the inclusion of enriched contextual variables — SNAP "
    "benefit eligibility flags, granular event categorisation, and weekly price data — reflects "
    "the types of operational data that medium-sized retailers already collect, fulfilling the "
    "SME accessibility principle stated in Chapter 1. Fourth, it constitutes an established "
    "academic benchmark, permitting direct comparison with published results from Obi (2024), "
    "Haque et al. (2023), and the broader M5 retrospective literature.")

add_heading(doc, "3.3.3 Stratified Subsample for Computational Tractability", 2)
add_para(doc,
    "The full M5 dataset expands to approximately 56.6 million product-store-day records once "
    "reshaped from wide to long format. While computationally tractable for tree-based gradient "
    "boosting, this scale prohibits rigorous walk-forward cross-validation and recurrent neural "
    "network training within the project's compute budget. A precedent for working with "
    "subsamples exists in the literature: Nasseri et al. (2023) report results on a subset of "
    "approximately 5.2 million records, and Seyedan et al. (2023) restrict their evaluation to "
    "selected product categories.")
add_para(doc,
    "This study uses a stratified random subsample of 502 product-store series, drawn "
    "proportionally across the nine strata defined by the three product categories crossed with "
    "the three states. The sampling is performed with a fixed random seed (42) to ensure "
    "reproducibility. The subsample preserves the categorical structure of the parent dataset "
    "while reducing total observations to approximately 933,000 — a 60-fold reduction that "
    "enables identical model architectures to be trained across all five tiers within consistent "
    "computational constraints, and permits the LSTM and Random Forest models to be evaluated "
    "rigorously rather than approximately. Table 3.1 reports the subsample composition.")

add_table_caption(doc, "Table 3.1: Subsample composition by category and state")
add_simple_table(doc,
    headers=["Category", "California", "Texas", "Wisconsin", "Row total"],
    rows=[
        ["Foods", "94", "71", "71", "236"],
        ["Hobbies", "37", "28", "28", "93"],
        ["Household", "69", "52", "52", "173"],
        ["Column total", "200", "151", "151", "502"],
    ])

# ── 3.4 PREPROCESSING ────────────────────────────────────────────────────
add_heading(doc, "3.4 Data Preprocessing Pipeline", 1)

add_heading(doc, "3.4.1 Data Integration", 2)
add_para(doc,
    "The three source files are merged into a single unified modelling dataset prior to any "
    "preprocessing. The sales file is first reshaped from wide format (one column per day) to "
    "long format (one row per product-store-day combination). The calendar file is joined on the "
    "date field, adding day-of-week, month, year, event type, event name, and SNAP flag "
    "variables. The sell_prices file is joined on the item_id, store_id, and wm_yr_wk composite "
    "key, adding the sell_price variable. The integrated long-format dataset is then filtered to "
    "the stratified subsample of 502 series.")

add_heading(doc, "3.4.2 Handling Missing Values", 2)
add_para(doc,
    "Missing values in the sell_price variable arise because prices are only recorded for weeks "
    "in which a product was actively stocked. Forward-fill imputation is applied first, using the "
    "most recent observed price for that product-store combination. Where no prior price "
    "observation exists, backward-fill imputation is applied from the first available price. "
    "This approach preserves the temporal structure of price variation and is consistent with "
    "retail practice, where products carry a stable price between changes (Nasseri et al., 2023; "
    "Barghi, 2025). Missing values in calendar event fields are treated as absence of an event "
    "and imputed with the label \"No Event\". No missing values are present in the core sales "
    "columns, as zero-sales records are valid observations.")

add_heading(doc, "3.4.3 Handling Zero-Sale Records", 2)
add_para(doc,
    "Zero-sale records may represent genuine zero demand, temporary stockouts, or store "
    "closures. Following Barghi (2025), zero-sales records corresponding to confirmed store "
    "closure days — identifiable from the calendar event_type field — are removed from the "
    "training set, as they do not represent purchasable demand. Zero-sale records on non-closure "
    "days are retained, as they represent genuine demand signals that are particularly important "
    "for safety stock calibration on intermittent series.")

add_heading(doc, "3.4.4 Outlier Detection and Treatment", 2)
add_para(doc,
    "Extreme sales outliers — defined as daily unit sales exceeding three standard deviations "
    "above the rolling 28-day mean for a given product-store combination — are identified and "
    "flagged. Outliers corresponding to confirmed promotional events or special calendar events "
    "are retained, as they represent genuine demand spikes that the model should learn to "
    "predict. Outliers that cannot be explained by any calendar or promotional event are "
    "winsorised to the 99th percentile for that product-store series, preventing extreme values "
    "from distorting gradient-based model training (Nasseri et al., 2023).")

add_heading(doc, "3.4.5 Categorical Variable Encoding", 2)
add_para(doc,
    "Identifier and categorical variables — item_id, dept_id, cat_id, store_id, state_id, "
    "event_type, and event_name — are encoded using label encoding. Label encoding is preferred "
    "over one-hot encoding for tree-based models because it preserves ordinality within compact "
    "integer representations without generating high-dimensional sparse matrices, which would "
    "substantially increase memory consumption. LightGBM natively handles label-encoded "
    "categoricals through its histogram-based splitting mechanism (Barghi, 2025; Obi, 2024).")

add_heading(doc, "3.4.6 Target-Encoded Mean Features", 2)
add_para(doc,
    "To complement the label-encoded identifiers, eight target-encoded mean features are "
    "constructed: item_mean, dept_mean, store_mean, state_mean, item_dow_mean, store_dow_mean, "
    "cat_month_mean, and item_month_mean. Each is the mean of the target variable computed "
    "exclusively over the training data, grouped by the relevant identifier(s). Target encoding "
    "is applied because it provides a strong prior for the expected demand level of each "
    "product-store combination, separating that signal cleanly from the temporal lag and rolling "
    "features. To prevent data leakage, all target encodings are fitted on the training window "
    "only and propagated unchanged to the validation and test windows. Missing combinations "
    "(categories not seen in training) are imputed with the global training mean. This technique "
    "is reported by M5 competition retrospectives as a significant contributor to top-performing "
    "solutions (Obi, 2024).")

# ── 3.5 FEATURE ENGINEERING ──────────────────────────────────────────────
add_heading(doc, "3.5 Feature Engineering", 1)
add_para(doc,
    "Feature engineering is the process of constructing derived variables from raw data that "
    "capture the demand patterns the models need to learn. The literature review identified "
    "feature engineering as contributing as much to predictive performance as model architecture "
    "choice: Barghi (2025) demonstrated that preprocessing alone reduced MAPE by 45.5% in a "
    "comparable retail forecasting study. The feature engineering strategy in this study is "
    "organised across six categories.")

add_heading(doc, "3.5.1 Lag Features", 2)
add_para(doc,
    "Lag features capture temporal autocorrelation by including historical sales values as "
    "predictors. Three lag windows are constructed for each product-store series: lag_7 (sales "
    "7 days prior, capturing weekly cyclical patterns), lag_14 (sales 14 days prior, capturing "
    "bi-weekly patterns), and lag_28 (sales 28 days prior, capturing monthly demand cycles and "
    "promotional recurrence). These windows are consistent with those used by Nasseri et al. "
    "(2023) and Obi (2024) on the M5 dataset.")

add_heading(doc, "3.5.2 Rolling Average Features", 2)
add_para(doc,
    "Rolling averages smooth demand signals across multiple time windows, providing estimates "
    "of recent trend and mean demand levels. Three rolling statistics are constructed: "
    "rolling_mean_7 (7-day rolling mean, capturing short-term trend), rolling_mean_28 (28-day "
    "rolling mean, capturing medium-term trend and reducing noise from single-day spikes), and "
    "rolling_std_7 (7-day rolling standard deviation, providing an estimate of recent demand "
    "volatility for safety stock calibration). All rolling statistics are computed over strictly "
    "historical windows to prevent data leakage.")

add_heading(doc, "3.5.3 Calendar and Temporal Features", 2)
add_para(doc,
    "Calendar features decompose the date variable into components that capture known demand "
    "seasonality patterns: day_of_week (0–6), week_of_year (1–52), month (1–12), year, "
    "is_weekend (binary indicator), days_since_last_event (count of days since the most recent "
    "calendar event, capturing pre-event demand run-up), and days_to_next_event (count of days "
    "until the next calendar event, capturing anticipation effects). Day-of-week is consistently "
    "the strongest temporal predictor across the literature; Seyedan et al. (2022) found "
    "Saturday sales exceeded midweek averages by 25–30% in comparable retail data.")

add_heading(doc, "3.5.4 Promotional and Event Features", 2)
add_para(doc,
    "Promotional and event indicators are among the strongest predictors of demand spikes in "
    "retail contexts (Ungureanu, 2025; Nasseri et al., 2023). The following binary indicators "
    "are constructed from the calendar file: is_snap_day (SNAP benefit eligibility in the "
    "relevant state, which drives demand for food products), is_sporting_event, "
    "is_cultural_event, is_national_event, and is_religious_event (derived from the event_type "
    "field).")

add_heading(doc, "3.5.5 Price Features", 2)
add_para(doc,
    "Two price-related features are constructed from the sell_prices file: sell_price (the "
    "weekly selling price of each product-store combination, a direct demand driver through "
    "price elasticity effects) and price_change (the week-on-week price difference, capturing "
    "the direction and magnitude of price movements and their potential demand impact).")

add_heading(doc, "3.5.6 Target-Encoded Mean Features", 2)
add_para(doc,
    "As specified in §3.4.6, eight target-encoded mean features are added to the feature set: "
    "item_mean, dept_mean, store_mean, state_mean, item_dow_mean, store_dow_mean, "
    "cat_month_mean, and item_month_mean. Table 3.2 summarises the complete engineered feature "
    "set.")

add_table_caption(doc, "Table 3.2: Complete feature engineering summary")
add_simple_table(doc,
    headers=["Category", "Features", "Demand signal captured"],
    rows=[
        ["Lag", "lag_7, lag_14, lag_28", "Historical sales at 7-, 14-, 28-day offsets"],
        ["Rolling statistics", "rolling_mean_7, rolling_mean_28, rolling_std_7",
         "Short- and medium-term trend; volatility"],
        ["Temporal", "day_of_week, week_of_year, month, year, is_weekend",
         "Weekly, seasonal and annual cycles"],
        ["Event proximity", "days_since_last_event, days_to_next_event",
         "Pre- and post-event demand run-up and decay"],
        ["Promotional / event", "is_snap_day, is_sporting_event, is_cultural_event, is_national_event, is_religious_event",
         "SNAP eligibility; event-type spikes"],
        ["Price", "sell_price, price_change",
         "Price level and week-on-week movement"],
        ["Identifier (label-encoded)", "item_id, dept_id, cat_id, store_id, state_id",
         "Cross-series generalisation"],
        ["Target-encoded mean", "item_mean, dept_mean, store_mean, state_mean, item_dow_mean, store_dow_mean, cat_month_mean, item_month_mean",
         "Demand-level prior per item, store, day-of-week and month"],
    ])

# ── 3.6 MODEL DEVELOPMENT ────────────────────────────────────────────────
add_heading(doc, "3.6 Model Development and Evaluation Framework", 1)

add_heading(doc, "3.6.1 Model Selection Rationale", 2)
add_para(doc,
    "Five model tiers are evaluated in this study, structured to provide a systematic "
    "performance progression from classical statistical baselines through to a quantile-"
    "regression model that supports the inventory policy integration. This tiered design "
    "reflects the comparative methodology adopted in the most rigorous studies reviewed in "
    "Chapter 2 — notably Barghi (2025), Nasseri et al. (2023), Mitra et al. (2022), and Seyedan "
    "et al. (2023) — and ensures that the performance improvement claimed for the proposed "
    "ensemble is substantiated against the full range of alternatives a practitioner might "
    "deploy.")

add_para(doc,
    "Tier 1 — Classical statistical baselines. Six classical models are implemented as "
    "performance baselines. The Moving Average (MA, 28-day window) provides a naïve demand "
    "estimate with no external features. Simple Exponential Smoothing (SES) applies "
    "exponentially decreasing weights to historical observations. AutoARIMA performs automatic "
    "parameter selection by AIC minimisation and is configured with a seasonal component of "
    "length seven to capture weekly cycles, encompassing both ARIMA and SARIMA as required by "
    "the data. Croston's Classic method and the Syntetos-Boylan Approximation (Croston-SBA) are "
    "added in recognition that retail SKU-daily demand is intermittent and contains many "
    "zero-sale days, conditions under which conventional methods like ARIMA and SES are known "
    "to underperform (Croston, 1972; Syntetos & Boylan, 2005). Seasonal Naive is included as a "
    "transparent reference baseline that simply forecasts the value observed seven days prior. "
    "These baselines represent methods that many SME retailers currently deploy in some form "
    "(Wahedi et al., 2023; Keith, 2023). All Tier 1 models are implemented via the "
    "statsforecast Python library, which provides parallel per-series fitting.")

add_para(doc,
    "Tier 2 — Standard machine learning models. Random Forest (RF) is a bagging ensemble of "
    "decision trees providing variance reduction and tolerance for non-linear interactions "
    "(Nasseri et al., 2023), implemented using scikit-learn. XGBoost is a gradient boosting "
    "algorithm with L1/L2 regularisation; LightGBM is a gradient boosting algorithm with leaf-"
    "wise tree growth and histogram-based binning that is well-suited to large tabular datasets "
    "(Obi, 2024; Barghi, 2025). For this study, both XGBoost and LightGBM are configured with "
    "the Tweedie loss function (variance power = 1.1), which is designed for non-negative count "
    "data with many zero observations and is consistently reported by M5 retrospective analyses "
    "as a significant contributor to top-performing solutions on intermittent demand.")

add_para(doc,
    "Tier 3 — Deep learning time-series model. A Long Short-Term Memory (LSTM) network captures "
    "long-range temporal dependencies through gated memory cells, learning multi-day "
    "promotional dynamics and seasonal rhythms that tree-based models treat as independent "
    "feature interactions (Punia et al., 2020; Nasseri et al., 2023). The LSTM architecture "
    "comprises two stacked recurrent layers (128 and 64 units respectively), each followed by a "
    "dropout layer with rate 0.2 for regularisation, and a single dense output unit. Input "
    "sequences use a 28-day lookback window. The target variable is transformed using log1p "
    "prior to training to handle the zero-inflated demand distribution and stabilise gradient "
    "updates; predictions are inverse-transformed using expm1 and clipped at zero. Training "
    "uses the Adam optimiser with mean-squared-error loss in log space, with early stopping on "
    "validation loss (patience = 5 epochs). Implementation uses TensorFlow/Keras.")

add_para(doc,
    "Tier 4 — Stacking ensemble (primary methodological contribution). The primary modelling "
    "contribution of this study is a stacking ensemble in which the three Tier 2 base learners "
    "— LightGBM, XGBoost, and Random Forest — produce predictions over the validation window, "
    "and these predictions serve as features for a Ridge regression meta-learner that learns "
    "the optimal combination weights:")
add_formula(doc, "ŷ_ensemble = w_LGBM · ŷ_LGBM + w_XGB · ŷ_XGB + w_RF · ŷ_RF + b")
add_para(doc,
    "where the weights w_i are constrained to be non-negative (forecast combination weights "
    "should not invert the base model's direction) and b is an intercept absorbing any "
    "systematic bias. The Ridge meta-learner is trained on the validation set predictions and "
    "applied to test set predictions to produce the final ensemble output. This architecture "
    "follows the stacking approach used by Seyedan et al. (2023) and Mitra et al. (2022), and "
    "is preferred over the sequential residual-correction design described in some prior work "
    "(Barghi, 2025) because (a) sequential residual chaining tends to overfit successive layers "
    "to the noise of earlier layers, and (b) the stacking approach with non-negative Ridge "
    "weights is the standard heterogeneous ensemble architecture in the modern forecasting "
    "literature.")

add_para(doc,
    "Tier 5 — LightGBM quantile regression (for inventory robustness). A fifth tier provides "
    "quantile forecasts at five probability levels (P10, P50, P90, P95, P99) by training five "
    "independent LightGBM models with objective=\"quantile\". These forecasts are used not as a "
    "competing point predictor against Tiers 1–4 (they predict quantiles, not means) but as the "
    "empirical input to the safety-stock robustness check defined in Section 3.7.4. This "
    "addresses the methodological gap identified by Seyedan et al. (2023): point forecast "
    "accuracy does not, on its own, govern inventory cost, because the standard Gaussian safety "
    "stock formula SS = z · σ · √(R + L) assumes normality of forecast residuals — an "
    "assumption that intermittent retail demand systematically violates.")

add_heading(doc, "3.6.2 Training, Validation, and Test Split", 2)
add_para(doc,
    "Temporal integrity is the critical constraint governing data splitting in time-series "
    "forecasting. Random splits risk data leakage — using future information to predict the "
    "past — which artificially inflates reported performance and renders results undeployable. "
    "All splits in this study strictly respect chronological order (Nasseri et al., 2023; "
    "Ungureanu, 2025).")
add_para(doc, "The dataset is partitioned into three chronologically ordered windows:")
add_bullet(doc,
    "Training set: January 2011 to 31 December 2015 (approximately 80% of the time series, "
    "covering five full calendar years and at least one complete cycle of seasonal and "
    "promotional patterns).")
add_bullet(doc,
    "Validation set: 1 January 2016 to 29 February 2016 (used for hyperparameter tuning, early "
    "stopping, and meta-learner training).")
add_bullet(doc,
    "Test set: 1 March 2016 to 24 April 2016 (the final eight weeks of the publicly distributed "
    "M5 data, reserved for out-of-sample evaluation and not observed during training or tuning).")
add_para(doc,
    "Walk-forward validation with three folds is additionally specified for hyperparameter "
    "tuning to ensure that performance estimates are stable across time windows: Fold 1 trains "
    "on 2011–2013 and validates on 2014; Fold 2 trains on 2011–2014 and validates on 2015; Fold "
    "3 trains on 2011–2015 and validates on Q1 2016.")

add_heading(doc, "3.6.3 Hyperparameter Optimisation", 2)
add_para(doc,
    "Hyperparameter optimisation uses a two-stage approach. In the first stage, a coarse random "
    "search over a broad parameter space identifies promising parameter regions efficiently. In "
    "the second stage, a fine-grained grid search is conducted over the promising regions with "
    "walk-forward cross-validation. Early stopping on the validation loss is applied for all "
    "gradient boosting and deep learning models. Table 3.3 documents the key hyperparameters.")

add_table_caption(doc, "Table 3.3: Hyperparameter search spaces")
add_simple_table(doc,
    headers=["Model", "Key hyperparameters", "Search strategy"],
    rows=[
        ["LightGBM (Tweedie)",
         "n_estimators, learning_rate, num_leaves, min_child_samples, tweedie_variance_power",
         "Random + Grid, walk-forward CV; early stopping on validation RMSE"],
        ["XGBoost (Tweedie)",
         "n_estimators, learning_rate, max_depth, subsample, colsample_bytree, reg_lambda, tweedie_variance_power",
         "Random + Grid, walk-forward CV; early stopping"],
        ["Random Forest",
         "n_estimators, max_depth, min_samples_leaf, max_features",
         "Grid Search"],
        ["LSTM",
         "units (L1, L2), dropout_rate, batch_size",
         "Fixed architecture (128/64, 0.2); learning rate and batch size grid; early stopping (patience 5)"],
        ["LightGBM Quantile",
         "n_estimators, learning_rate, num_leaves (per quantile level)",
         "Same as LightGBM, one fit per quantile α ∈ {0.10, 0.50, 0.90, 0.95, 0.99}"],
        ["Ridge meta-learner",
         "alpha, fit_intercept, positive",
         "Grid (alpha ∈ {0.1, 1.0, 10.0}); positive=True enforced"],
    ])

add_heading(doc, "3.6.4 Evaluation Metrics", 2)
add_para(doc,
    "Five complementary accuracy metrics evaluate forecasting performance across all models. "
    "Mean Absolute Error (MAE) measures the average absolute difference between predicted and "
    "actual demand in original demand units; it is the primary operational metric for store "
    "managers because it is expressed in the same units as daily sales. Root Mean Squared Error "
    "(RMSE) penalises large errors more heavily than MAE and is sensitive to promotional demand "
    "spikes. Mean Absolute Percentage Error (MAPE) normalises errors as a percentage of actual "
    "demand, enabling cross-product and cross-study comparisons; following standard practice, "
    "MAPE in this study masks zero-actual rows to avoid division blow-ups.")
add_para(doc,
    "To complement MAPE on intermittent demand series, Weighted Mean Absolute Percentage Error "
    "(WMAPE) is reported as a supplementary metric. WMAPE divides the sum of absolute errors by "
    "the sum of absolute actual values across all observations rather than averaging per-row "
    "percentage errors, and is therefore robust to the zero-actual rows that distort per-row "
    "MAPE on retail data. WMAPE is the metric family used by the M5 competition itself (in the "
    "form of WRMSSE) and by the broader modern forecasting literature for intermittent demand "
    "(Hyndman & Athanasopoulos, 2021). Finally, Pred(10%) — the percentage of forecasts falling "
    "within 10% of the actual value — is reported as a threshold-based precision measure "
    "directly relevant to inventory decisions (Barghi, 2025).")
add_para(doc,
    "The performance targets stated in Chapter 1 — 10–25% improvement in MAPE and 5–15% "
    "reduction in inventory cost relative to classical baselines — are informed by benchmarks "
    "reported in Barghi (2025) and Seyedan et al. (2023). However, the M5 retrospective "
    "literature indicates that double-digit MAPE improvements at the SKU-daily level are at the "
    "upper end of what is typically achievable, because per-series residual variance is bounded "
    "below by the inherent intermittence of the data. Chapter 4 will report results against "
    "both ends of the target range and discuss any deviation in light of these constraints, in "
    "line with Seyedan et al. (2023)'s observation that operational value at the inventory "
    "level can be substantial even when forecast accuracy improvements at the SKU-daily level "
    "are modest.")

# ── 3.7 INVENTORY ────────────────────────────────────────────────────────
add_heading(doc, "3.7 Inventory Policy Integration: The Order-Up-To-Level Framework", 1)

add_heading(doc, "3.7.1 Policy Selection and Justification", 2)
add_para(doc,
    "The Order-Up-To-Level (OUTL) policy, implemented as a periodic (R, s, S) review system, "
    "is selected as the inventory framework for this study. Under this policy, inventory is "
    "reviewed at fixed intervals R; if the inventory position falls at or below the reorder "
    "point s, a replenishment order is placed to raise the inventory level to the order-up-to "
    "level S. The (R, s, S) system is selected over alternative policies — such as continuous "
    "(s, Q) or Newsvendor formulations — because of its compatibility with batch data updates "
    "characteristic of retail sales systems (Barghi, 2025; Seyedan et al., 2023), its widespread "
    "use in the empirical literature, and its explicit parameterisation in terms of forecast "
    "accuracy and lead time, which enables direct quantification of the operational value of "
    "improved forecasting.")

add_heading(doc, "3.7.2 Gaussian Policy Parameter Formulations (Primary)", 2)
add_para(doc,
    "The primary OUTL parameters are computed following Silver, Pyke and Thomas (2017) and "
    "consistent with the formulations in Barghi (2025) and Seyedan et al. (2023). Safety Stock "
    "(SS) represents the buffer inventory maintained to absorb demand variability:")
add_formula(doc, "SS = z · σ · √(R + L)")
add_para(doc,
    "where z is the service-level z-score, σ is the standard deviation of forecast residuals "
    "derived from the forecasting model's prediction errors on the validation set, R is the "
    "review period in days, and L is the replenishment lead time in days. Crucially, σ is set "
    "to the standard deviation of forecast residuals rather than historical demand variability, "
    "representing the tighter coupling of forecasting and inventory stages identified as an "
    "unaddressed gap by Seyedan et al. (2023). A 95% service level is the primary target, "
    "corresponding to z = 1.6449.")
add_para(doc, "The Reorder Point (r) triggers replenishment:")
add_formula(doc, "r = d̄ · (R + L) + z · σ · √(R + L)")
add_para(doc,
    "where d̄ is the mean daily demand over the review and lead-time horizon, estimated from "
    "the model's point forecasts.")
add_para(doc, "The Order-Up-To Level (S) is the target inventory level to which stock is replenished:")
add_formula(doc, "S = r + Q")
add_para(doc,
    "where Q is the order quantity, computed from the Economic Order Quantity (EOQ) formula "
    "Q* = √(2 · D · K / h) with D the annual demand, K the fixed ordering cost, and h = "
    "unit_cost · carrying_rate the per-unit annual holding cost. EOQ is used in preference to a "
    "fixed Q = d̄ · R quantity because it decouples ordering frequency from forecast magnitude, "
    "ensuring that comparative cost analysis isolates the effect of forecast quality on "
    "inventory carrying costs rather than confounding it with ordering operations (Silver et "
    "al., 2017).")
add_para(doc, "Total annual cost integrates holding, ordering and stockout components:")
add_formula(doc, "TC = HC + OC + SC")
add_para(doc,
    "where holding cost HC = (SS + Q/2) · h, ordering cost OC = (D / Q) · K, and stockout cost "
    "SC is defined in §3.7.5. The carrying-rate is estimated at 25% of unit cost per year, "
    "reflecting the standard industry range of 10–30% documented by Silver et al. (2017). Unit "
    "cost is estimated as 60% of selling price (40% gross margin) following standard retail "
    "practice. Fixed ordering cost K is set at $50 per order, a typical value for small and "
    "medium retailers.")

add_heading(doc, "3.7.3 Lead Time Scenarios", 2)
add_para(doc,
    "OUTL parameters are evaluated across four lead-time scenarios: L = 7 days (short, regional "
    "distribution), L = 10 days (moderate), L = 14 days (medium, national distribution), and "
    "L = 21 days (extended, international or seasonal supply). This multi-scenario evaluation "
    "mirrors the approach of Seyedan et al. (2023), who demonstrate that the operational value "
    "of improved forecasting is most pronounced at longer lead times, because safety stock "
    "requirements grow with √(R + L), amplifying the cost impact of forecast accuracy "
    "differences. Inventory cost reductions of 5–15% relative to classical baseline-driven OUTL "
    "parameters are the primary operational performance target.")

add_heading(doc, "3.7.4 Robustness Check: Empirical Quantile Safety Stock", 2)
add_para(doc,
    "The standard Gaussian formula SS = z · σ · √(R + L) assumes that forecast residuals are "
    "normally distributed. For retail SKU-daily demand, this assumption is systematically "
    "violated: residuals are right-skewed (because demand has a long upper tail driven by "
    "promotions and events) and zero-inflated (because many days have no sales). The Gaussian "
    "formula consequently under-allocates safety stock relative to what is actually required to "
    "achieve the targeted service level on this class of data.")
add_para(doc,
    "To quantify and correct for this, the study computes a parallel, empirical safety stock "
    "derived from the LightGBM quantile-regression model of Tier 5. For each series, the per-"
    "day quantile spread q_α(x) − q_0.5(x) is computed at the service-level quantile α (e.g., "
    "0.95 for a 95% service level). This per-day spread is then scaled to the (R + L)-day "
    "horizon under the same i.i.d. assumption used in the Gaussian formula:")
add_formula(doc, "SS_empirical(α, L) = (q_α − q_0.5) · √(R + L)")
add_para(doc,
    "The empirical and Gaussian safety stocks are computed for the same series, lead times, "
    "and service levels, and OUTL parameters are recomputed for both. The comparison provides "
    "the methodological evidence required to assess whether the Gaussian formula systematically "
    "over- or under-allocates safety stock for intermittent retail demand — a question "
    "explicitly identified as unaddressed in the literature by Seyedan et al. (2023) and "
    "central to Gap 1 of this thesis.")

add_heading(doc, "3.7.5 Stockout Cost Modelling", 2)
add_para(doc,
    "A realistic total-cost model requires the inclusion of stockout cost — the financial "
    "penalty incurred when actual demand exceeds available inventory. Following Silver et al. "
    "(2017), stockout cost is modelled as:")
add_formula(doc, "SC = (target_SL − achieved_SL) · D · unit_cost · m")
add_para(doc,
    "where the achieved service level is the empirical fraction of demand that the safety "
    "stock buffer covers, target_SL is the policy's nominal service level (e.g., 0.95), D is "
    "annual demand, and m is a stockout cost multiplier capturing the effective cost per "
    "stockout unit. For the empirical (quantile-based) policy, achieved_SL ≈ target_SL by "
    "calibration. For the Gaussian policy, achieved_SL is computed by projecting the Gaussian "
    "buffer onto the empirical demand distribution observed in the quantile model — yielding a "
    "quantification of the under-allocation gap.")
add_para(doc,
    "Three stockout cost multipliers are evaluated as sensitivity scenarios: m = 0.4 (lost "
    "gross margin only, conservative), m = 1.0 (full unit cost, baseline industry assumption "
    "following Silver et al. 2017), and m = 2.0 (full cost plus reputational damage). This "
    "three-point sensitivity bound provides Chapter 5 with an honest range for the headline "
    "inventory cost reduction figure.")

# ── 3.8 DASHBOARD ────────────────────────────────────────────────────────
add_heading(doc, "3.8 Power BI Dashboard Development", 1)
add_para(doc,
    "A Power BI dashboard is developed to translate forecasting outputs and inventory "
    "parameters into an interactive, accessible decision-support tool for non-technical retail "
    "managers. The dashboard addresses the deployment gap identified in Gap 2 of the literature "
    "review and directly serves the SME accessibility objective of Chapter 1. Power BI Desktop "
    "is selected as the deployment platform because it is available under Microsoft's free "
    "tier, requires no server infrastructure, and is the most widely deployed business-"
    "intelligence tool in the SME segment.")
add_para(doc,
    "The dashboard is structured across four pages, each driven by a separate CSV feed "
    "exported from the Python modelling pipeline:")
add_bullet(doc,
    "Forecast Overview presents actual sales, ML forecasts, classical baseline forecasts, and "
    "the LightGBM quantile fan (P50/P90/P95/P99) for any selected product-store combination, "
    "with slicers for category, state, and store.")
add_bullet(doc,
    "Inventory Status presents safety stock, reorder point, and order-up-to-level values per "
    "product-store at the user-selected lead time and service level, with conditional "
    "formatting highlighting items approaching their reorder threshold.")
add_bullet(doc,
    "Cost Analysis decomposes total annual inventory cost into holding, ordering, and "
    "stockout components across lead-time scenarios and forecasting policies, enabling like-"
    "for-like comparison of Classical, ML-Gaussian, and ML-Empirical-Quantile policies.")
add_bullet(doc,
    "Sensitivity Analysis presents the curves from §3.9, showing how total cost reduction "
    "varies with MAPE level, lead time, service level, and stockout cost assumption.")
add_para(doc,
    "The dashboard's data feeds are designed for daily refresh in a production deployment, "
    "demonstrating deployment readiness consistent with the SME accessibility principle.")

# ── 3.9 SENSITIVITY ──────────────────────────────────────────────────────
add_heading(doc, "3.9 Sensitivity Analysis Design", 1)
add_para(doc,
    "A sensitivity analysis quantifies how the inventory savings from quantile-based safety "
    "stock vary across the principal assumptions of the framework. The analysis varies four "
    "parameters independently: forecast MAPE (3% to 15%, the range from state-of-the-art "
    "ensemble performance to classical baseline performance), lead time L (7, 10, 14, 21 days), "
    "service level (0.90, 0.95, 0.99), and stockout cost multiplier m (0.4, 1.0, 2.0). For each "
    "combination, the total annual cost under the Gaussian and Empirical-Quantile policies is "
    "recomputed and the percentage reduction is tabulated.")
add_para(doc,
    "This four-dimensional sweep provides decision-relevant guidance for practitioners "
    "evaluating forecasting technology investments by quantifying the marginal operational "
    "value of each percentage point of MAPE improvement, the cost of meeting higher service "
    "levels, and the robustness of the headline cost-reduction figure to the choice of "
    "stockout cost assumption. Results are reported as sensitivity curves and a summary table "
    "in Chapter 5.")

# ── 3.10 ETHICS ──────────────────────────────────────────────────────────
add_heading(doc, "3.10 Ethical Considerations and Model Transparency", 1)
add_para(doc,
    "This study uses only publicly available, aggregated, and anonymised data sourced from "
    "Kaggle. The M5 dataset contains no personally identifiable information; all records are "
    "aggregated at the product-store-day level and contain no individual customer, transaction, "
    "or demographic data. No primary data collection from human participants is involved. "
    "Accordingly, formal ethics committee approval is not required under the institution's "
    "research ethics guidelines. Data usage is governed by the Kaggle Competition Rules and the "
    "licence under which the M5 dataset is distributed for academic research.")
add_para(doc,
    "In line with the responsible AI principles identified by Ungureanu (2025), the study "
    "applies three ethical practices:")
add_para(doc,
    "Model transparency. Feature importance for the best-performing model is assessed using "
    "SHAP (SHapley Additive exPlanations) values, computed via the shap.TreeExplainer "
    "implementation on a 5,000-row stratified sample of the test set. SHAP attribution is "
    "reported in tabular form (top 15 features by mean absolute SHAP) and visually (bar and "
    "beeswarm plots), ensuring that the key drivers of model predictions can be identified and "
    "communicated to practitioners.")
add_para(doc,
    "Performance fairness. Forecast error metrics are computed disaggregated by product "
    "category (Foods, Hobbies, Household) and by state (California, Texas, Wisconsin), "
    "verifying that no category or geographic segment performs more than 2× worse than another "
    "— a heuristic threshold for systematic algorithmic bias. Variation in metrics that "
    "reflects underlying demand structure (e.g., higher predictability of Foods than Hobbies "
    "due to intermittence differences) is documented and distinguished from disparate "
    "algorithmic treatment.")
add_para(doc,
    "Data minimisation. Only features demonstrably contributing to forecasting performance are "
    "retained in the final model, with marginal-impact features identified via SHAP importance "
    "ranking dropped where their removal does not degrade evaluation metrics.")

# ── 3.11 SME TRANSFERABILITY ─────────────────────────────────────────────
add_heading(doc, "3.11 SME Transferability Specification", 1)
add_para(doc,
    "A central design principle of this study, stated in Chapter 1, is that the framework is "
    "replicable in small and medium-sized retail environments. To make this commitment concrete "
    "rather than aspirational, this section specifies the minimum requirements a typical SME "
    "would need to deploy the pipeline.")
add_para(doc,
    "Data requirements. Three data fields are required: daily unit sales by SKU-store, a "
    "calendar containing day-of-week and any locally relevant event flags (national holidays, "
    "regional promotions), and a price file containing weekly or daily selling prices by SKU-"
    "store. All three are standard outputs of any commercial point-of-sale system, including "
    "those bundled with off-the-shelf SME retail software packages (e.g., Square, Lightspeed, "
    "Vend).")
add_para(doc,
    "Tooling. The end-to-end pipeline is implemented in open-source Python using pandas, "
    "scikit-learn, lightgbm, xgboost, tensorflow, statsforecast, and shap. All libraries are "
    "free, run on standard CPU hardware (no GPU is required for the gradient boosting models "
    "that constitute the recommended deployment), and are available under permissive licences "
    "(MIT, BSD, Apache 2.0). The Power BI dashboard requires only Microsoft Power BI Desktop, "
    "which is freely distributed.")
add_para(doc,
    "Hardware. The full pipeline — preprocessing, feature engineering, model training, "
    "hyperparameter tuning, OUTL calculations, sensitivity analysis, and dashboard export — "
    "runs on a single laptop with 16 GB RAM and a modern multi-core CPU. Total runtime on the "
    "502-series stratified subsample is under three hours; on a typical SME inventory of 200–"
    "500 SKUs, runtime would be comparable or lower.")
add_para(doc,
    "Skills. A single analyst with intermediate Python competence (sufficient to install "
    "packages, execute notebooks, and modify configuration constants such as service level and "
    "lead time) can operate the pipeline. No specialist data science qualification is required "
    "for routine use. Configuration of the pipeline for a new SME would require approximately "
    "one to two working days for an experienced analyst, primarily for data ingestion "
    "adaptation and dashboard customisation.")
add_para(doc,
    "These specifications collectively demonstrate that the framework is operationally "
    "accessible to SMEs and is not gated by infrastructure, licensing, or specialist expertise "
    "barriers — addressing the SME accessibility objective in concrete terms.")

# ── 3.12 SUMMARY ─────────────────────────────────────────────────────────
add_heading(doc, "3.12 Chapter Summary", 1)
add_para(doc,
    "This chapter has presented the complete research methodology governing the empirical work "
    "of this thesis. The CRISP-DM framework structures the research across data understanding, "
    "preparation, modelling, evaluation, and deployment, each grounded in the gaps identified "
    "in Chapter 2 and designed to deliver against the objectives stated in Chapter 1.")
add_para(doc,
    "The M5-Forecasting (Walmart) dataset provides a large-scale, richly featured empirical "
    "benchmark. A stratified subsample of 502 product-store series, preserving category × state "
    "proportions, is used to enable rigorous walk-forward cross-validation and recurrent neural "
    "network training within the project's compute budget. The preprocessing pipeline addresses "
    "missing values, zero-sales records, outliers, and categorical encoding in a principled, "
    "literature-grounded manner. The feature engineering strategy constructs derived variables "
    "across lag, rolling statistic, temporal, event proximity, promotional, price, identifier, "
    "and target-encoded categories.")
add_para(doc,
    "The model comparison framework evaluates five tiers of forecasting approaches: classical "
    "statistical baselines (including Croston's method for intermittent demand), standard "
    "machine learning models with Tweedie-loss configuration, a stacked LSTM, a stacking "
    "ensemble with Ridge meta-learner, and a quantile-regression model that enables empirically "
    "calibrated safety stock. Performance is measured by MAE, RMSE, MAPE, WMAPE, and Pred(10%) "
    "— with WMAPE added explicitly to provide an honest metric for the intermittent-demand "
    "setting where per-row MAPE is unstable.")
add_para(doc,
    "The OUTL inventory framework integrates forecasting outputs into operational safety stock, "
    "reorder point, and total cost estimates across four lead-time scenarios and three service "
    "levels. The Gaussian safety stock formulation remains the primary, Chapter-1-anchored "
    "approach; an empirical quantile-based safety stock provides a robustness check that "
    "quantifies the systematic under-allocation of the Gaussian formula on intermittent demand. "
    "Stockout cost is modelled with a three-point sensitivity to provide an honest range for "
    "the headline cost-reduction figure.")
add_para(doc,
    "The Power BI dashboard delivers the framework's outputs in an accessible format for non-"
    "technical retail managers. A four-dimensional sensitivity analysis quantifies how cost "
    "savings scale with forecast accuracy, lead time, service level, and stockout cost "
    "assumptions. SHAP-based transparency, disaggregated fairness analysis, and a concrete SME "
    "transferability specification complete the responsible-AI dimension of the methodology.")
add_para(doc,
    "Chapter 4 presents the implementation and results of this framework, reporting model "
    "performance metrics and feature attribution. Chapter 5 reports the OUTL inventory "
    "parameter estimates and sensitivity analysis findings, integrating Chapter 1's "
    "quantitative targets with the empirical results across all model tiers and operating "
    "scenarios.")

# ── REFERENCES ───────────────────────────────────────────────────────────
add_heading(doc, "References (Chapter 3)", 1)
add_para(doc,
    "The following entries are additions for this chapter; the full Chapter 3 bibliography is "
    "retained as in the parent thesis document.", italic=True, font_size=10)

refs = [
    "Barghi, S. (2025). Demand forecasting and inventory improvement in supply chain "
    "management using hybrid boosting ensemble techniques. Master's thesis, École de "
    "Technologie Supérieure, Université du Québec.",
    "Bastos, A. S. T. (2023). Machine learning in digital retail demand forecasting for "
    "inventory management in a sportswear company. Master's thesis, NOVA Information "
    "Management School, Universidade Nova de Lisboa.",
    "Croston, J. D. (1972). Forecasting and stock control for intermittent demands. "
    "Operational Research Quarterly, 23(3), 289–303.",
    "Haque, M. S., Amin, M. S., & Miah, J. (2023). Retail demand forecasting: A comparative "
    "study for multivariate time series (arXiv:2308.11939). arXiv.",
    "Hyndman, R. J., & Athanasopoulos, G. (2021). Forecasting: Principles and Practice "
    "(3rd ed.). OTexts.",
    "Keith, E. (2023). Optimizing inventory management through advanced forecasting techniques "
    "in supply chains. European Journal of Supply Chain Management, 1(1), 22–30.",
    "Mitra, A., Jain, A., Kishore, A., & Kumar, P. (2022). A comparative study of demand "
    "forecasting models for a multi-channel retail company: A novel hybrid machine learning "
    "approach. Operations Research Forum, 3, 58.",
    "Nasseri, M., Falatouri, T., Brandtner, P., & Darbanian, F. (2023). Applying machine "
    "learning in retail demand prediction — A comparison of tree-based ensembles and long "
    "short-term memory-based deep learning. Applied Sciences, 13, 11112.",
    "Obi, C. I. C. (2024). Demand forecasting in retail business using the ensemble machine "
    "learning framework: A stacking approach. American Academic Scientific Research Journal "
    "for Engineering, Technology, and Sciences, 98(1), 309–329.",
    "Punia, S., Nikolopoulos, K., Prakash Singh, S., Madaan, J. K., & Litsiou, K. (2020). "
    "Deep learning with long short-term memory networks and random forests for demand "
    "forecasting in multi-channel retail. International Journal of Production Research, "
    "58(16), 4964–4979.",
    "Seyedan, M. (2023). Development of predictive analytics for demand forecasting and "
    "inventory management in supply chain using machine learning techniques (Doctoral "
    "dissertation). Concordia University, Montreal.",
    "Seyedan, M., Mafakheri, F., & Wang, C. (2022). Cluster-based demand forecasting using "
    "Bayesian model averaging: An ensemble learning approach. Decision Analytics Journal, 3, "
    "100033.",
    "Seyedan, M., Mafakheri, F., & Wang, C. (2023). Order-up-to-level inventory optimization "
    "model using time-series demand forecasting with ensemble deep learning. Supply Chain "
    "Analytics, 3, 100024.",
    "Silver, E. A., Pyke, D. F., & Thomas, D. J. (2017). Inventory and production management "
    "in supply chains (4th ed.). CRC Press.",
    "Syntetos, A. A., & Boylan, J. E. (2005). The accuracy of intermittent demand estimates. "
    "International Journal of Forecasting, 21(2), 303–314.",
    "Ungureanu, D. A. (2025). Demand forecasting and inventory optimization in mid-sized "
    "grocery retail using machine learning. Master's thesis, CCT College Dublin.",
    "Wahedi, H. J., Heltoft, M., Christophersen, G. J., Severinsen, T., Saha, S., & Nielsen, "
    "I. E. (2023). Forecasting and inventory planning: An empirical investigation of "
    "classical and machine learning approaches for Svanehøj's future software consolidation. "
    "Applied Sciences, 13, 8581.",
]
for r in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.first_line_indent = Inches(-0.4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(r)
    run.font.name = "Arial"
    run.font.size = Pt(10)

OUT = "/Users/desmond/Capstone Project/retail-demand-forecasting/Chapter_3_Methodology.docx"
doc.save(OUT)
print(f"Saved: {OUT}")
