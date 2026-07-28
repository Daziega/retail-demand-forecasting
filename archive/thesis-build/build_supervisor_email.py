"""Generate the supervisor email/report as a Word document."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def add_para(doc, text, *, bold=False, italic=False, font_size=11, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    return p


def add_table_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(11)


def add_table(doc, headers, rows, header_fill="D5E8F0"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.autofit = True

    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(10)
        set_cell_shading(hdr[i], header_fill)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for r_idx, row in enumerate(rows, start=1):
        cells = table.rows[r_idx].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ""
            p = cells[c_idx].paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = "Arial"
            run.font.size = Pt(10)
            cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return table


# ─────────────────────────────────────────────────────────────────────────────
doc = Document()

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(11)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = title.add_run("Inventory Simulation Results — Real Data")
run.bold = True
run.font.name = "Arial"
run.font.size = Pt(16)
title.paragraph_format.space_after = Pt(6)

sub = doc.add_paragraph()
run = sub.add_run("Forward (R, s, S) simulation on the M5 stratified subsample, "
                  "replacing the analytical assumed-service-level formulation")
run.italic = True
run.font.name = "Arial"
run.font.size = Pt(11)
sub.paragraph_format.space_after = Pt(18)

# Salutation / opener
add_para(doc, "Hi [Supervisor],")
add_para(doc,
    "I ran the simulation with load_real_dataset() wired to the trained-pipeline "
    "artefacts (ensemble forecasts, Croston-SBA baseline, LightGBM quantile "
    "predictions, per-series prices). The script runs on the full 502-series "
    "stratified subsample over the 56-day test window (March–April 2016). "
    "Output below.")

# Central case table
add_table_caption(doc, "Central case — L = 14 days, m = 1.0, target SL = 95%")
add_table(doc,
    headers=["Policy", "Realised SL", "SS (units)", "HC ($)", "OC ($)", "SC ($)", "Total ($)", "vs Classical"],
    rows=[
        ["Classical (Croston-SBA, Gaussian SS)", "100.0%", "4,620", "63,256", "332", "0", "63,587", "—"],
        ["ML-Gaussian", "100.0%", "4,572", "63,243", "332", "0", "63,575", "+0.0%"],
        ["ML-Empirical-Quantile", "100.0%", "6,287", "64,241", "332", "0", "64,573", "−1.5%"],
    ])

# Sensitivity to stockout multiplier
add_table_caption(doc, "Sensitivity to stockout multiplier m  (L = 14, target SL = 95%)")
add_table(doc,
    headers=["m", "Interpretation", "ML-Quantile vs Classical", "ML-Gaussian vs Classical"],
    rows=[
        ["0.4", "lost margin only (conservative)",       "−1.5%", "0.0%"],
        ["1.0", "full unit cost (baseline)",              "−1.5%", "0.0%"],
        ["2.0", "unit cost + reputation (aggressive)",    "−1.5%", "0.0%"],
    ])
add_para(doc,
    "The reduction is invariant to m because no policy stocks out — SC is zero "
    "across the board, so the multiplier never engages.",
    italic=True, font_size=10)

# Sensitivity to lead time
add_table_caption(doc, "Sensitivity to lead time L  (m = 1.0, target SL = 95%)")
add_table(doc,
    headers=["L (days)", "Classical TC ($)", "ML-Quantile TC ($)", "Reduction"],
    rows=[
        ["7",  "61,020", "61,825", "−1.3%"],
        ["10", "62,125", "63,011", "−1.4%"],
        ["14", "63,587", "64,573", "−1.5%"],
        ["21", "66,184", "67,321", "−1.7%"],
    ])

# Reading this
add_para(doc, "")
add_para(doc, "Reading this", bold=True, font_size=12, space_after=4)
add_para(doc,
    "Realised service level is 100% for all three policies — measured by counting "
    "unmet demand day-by-day, not assigned. The thesis's prior assumption of 83% "
    "(Gaussian) / 95% (Quantile) does not hold on the test window: there are "
    "essentially no stockouts.")
add_para(doc,
    "The 21% reduction collapses to ±0–2%. ML-Gaussian moves total cost by 0.0%, "
    "and ML-Empirical-Quantile actually increases cost by ~1.5% because it "
    "allocates ~36% more safety stock than the Gaussian policy without preventing "
    "any additional stockouts.")
add_para(doc,
    "Your secondary finding is confirmed: better point forecasts barely move "
    "inventory cost under the standard analytical (R, s, S) policy on SKU-daily "
    "M5. The quantile-based safety-stock reformulation, sound in principle, does "
    "not pay off when the policy is already well-served by point forecasts.")

# Next steps
add_para(doc, "")
add_para(doc, "Next steps", bold=True, font_size=12, space_after=4)
add_para(doc,
    "I will rewrite Chapter 5 around this finding as a disconfirmation result — "
    "explicitly stating that the Chapter 1 5–15% inventory cost target was not "
    "achieved, and reframing the contribution as bounding the conditions under "
    "which ML forecasting affects inventory outcomes. Chapter 4 (forecast "
    "accuracy, SHAP, fairness analysis) is unaffected. The full simulation grid "
    "(36 rows covering all L × m combinations) is saved to "
    "data/processed/simulation_results.csv for inclusion in the appendix.")
add_para(doc,
    "Let me know if you want me to proceed with the rewrite, or whether you'd "
    "like to discuss the framing first.")

add_para(doc, "")
add_para(doc, "Best regards,")
add_para(doc, "[Your name]")

OUT = "/Users/desmond/Capstone Project/retail-demand-forecasting/Supervisor_Email_Simulation_Results.docx"
doc.save(OUT)
print(f"Saved: {OUT}")
