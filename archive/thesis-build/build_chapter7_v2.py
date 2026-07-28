"""Generate Chapter 7 - Conclusion (revised to align with the supervisor's
rewritten Chapter 5 disconfirmation finding)."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_para(doc, text, *, italic=False, bold=False, indent_left=None, font_size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if indent_left:
        p.paragraph_format.left_indent = Inches(indent_left)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(font_size)
    run.italic = italic
    run.bold = bold
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    return p


# ─────────────────────────────────────────────────────────────────────────
doc = Document()
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(11)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("CHAPTER 7")
run.bold = True; run.font.name = "Arial"; run.font.size = Pt(20)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("Conclusion")
run.bold = True; run.font.name = "Arial"; run.font.size = Pt(16)
sub.paragraph_format.space_after = Pt(18)

# ── 7.1 INTRODUCTION ─────────────────────────────────────────────────────
add_heading(doc, "7.1 Introduction", 1)
add_para(doc,
    "This chapter concludes the thesis. It returns to the research objectives stated "
    "in Chapter 1, evaluates the extent to which each has been achieved through the "
    "empirical work documented in Chapters 3 to 6, restates the principal scholarly "
    "and practical contributions of the thesis, and offers closing reflections on the "
    "operational significance of the findings for the small and medium-sized retail "
    "sector. The chapter is structured as follows. Section 7.2 restates the five "
    "research objectives. Section 7.3 evaluates achievement against each objective. "
    "Section 7.4 consolidates the principal contributions of the thesis. Section 7.5 "
    "offers closing remarks and identifies the broader significance of the work.")

# ── 7.2 RESTATING OBJECTIVES ─────────────────────────────────────────────
add_heading(doc, "7.2 Restatement of the Research Objectives", 1)
add_para(doc, "Chapter 1 set out five research objectives that have guided the thesis:")
add_bullet(doc,
    "Objective 1: Conduct a comparative evaluation of forecasting approaches spanning "
    "classical statistical, standard machine learning, deep learning, ensemble, and "
    "probabilistic methods on a hierarchical retail demand benchmark.")
add_bullet(doc,
    "Objective 2: Develop a systematic feature engineering strategy capturing the "
    "principal drivers of retail demand — lags, rolling statistics, calendar and "
    "event signals, price movement, and identifier-level prior information.")
add_bullet(doc,
    "Objective 3: Integrate the best-performing forecasting approach with the Order-"
    "Up-To-Level (OUTL) inventory policy and quantify the resulting operational "
    "parameters across lead-time scenarios.")
add_bullet(doc,
    "Objective 4: Test whether the operational gains reported in closely related "
    "studies — improvements of 10–25% in forecast accuracy and inventory cost "
    "reductions of 5–15% — transfer to this setting at SKU-daily granularity.")
add_bullet(doc,
    "Objective 5: Develop a Power BI dashboard that translates forecasting outputs and "
    "inventory recommendations into an accessible, interactive decision-support tool "
    "suitable for deployment in resource-constrained retail environments, including "
    "small and medium-sized enterprises.")

# ── 7.3 ACHIEVEMENT ──────────────────────────────────────────────────────
add_heading(doc, "7.3 Achievement Against Research Objectives", 1)

add_heading(doc, "7.3.1 Objective 1 — Comparative evaluation of forecasting approaches", 2)
add_para(doc,
    "The thesis evaluated twelve forecasting models organised across five tiers on a "
    "stratified subsample of the M5-Forecasting (Walmart) dataset. Tier 1 implemented "
    "six classical baselines (Moving Average, Simple Exponential Smoothing, AutoARIMA, "
    "Croston's Classic, Croston-SBA, and Seasonal Naive); Tier 2 implemented three "
    "standard ML models (LightGBM and XGBoost with Tweedie loss, and Random Forest); "
    "Tier 3 implemented a stacked LSTM with log1p target transformation; Tier 4 "
    "implemented a stacking ensemble of the Tier 2 base learners through a non-negative "
    "Ridge meta-learner; and Tier 5 implemented a five-quantile LightGBM model "
    "supporting the inventory analysis.")
add_para(doc,
    "The comparative evaluation produced a ranked leaderboard demonstrating that the "
    "Tier 4 stacking ensemble achieves the best MAE (0.952) and WMAPE (73.58%) across "
    "all twelve models. The empirical ranking reproduces the consensus of the modern "
    "retail forecasting literature: tree-based gradient boosting and ensembles thereof "
    "at the top, classical baselines (with Croston-class methods particularly well-"
    "suited to intermittent demand) in the middle, and the LSTM at the bottom — a "
    "finding that directly replicates Nasseri et al. (2023). Objective 1 is fully "
    "achieved.")

add_heading(doc, "7.3.2 Objective 2 — Systematic feature engineering", 2)
add_para(doc,
    "The thesis constructed thirty-three engineered features organised across six "
    "categories (lag, rolling, temporal, event proximity, promotional/event, price) "
    "plus eight target-encoded mean features for high-cardinality categoricals. The "
    "feature engineering strategy is documented to the level required for independent "
    "replication in §3.5 and §3.4.6.")
add_para(doc,
    "SHAP analysis on the LightGBM model (§4.11) demonstrates that the engineered "
    "features are the principal predictive drivers: the top five SHAP-ranked features "
    "account for 78.6% of model predictions, dominated by rolling means (rolling_mean_7 "
    "at 29.9%, rolling_mean_28 at 20.4%), target-encoded item-level seasonality "
    "(item_month_mean at 15.3%, item_dow_mean at 5.3%), and demand volatility "
    "(rolling_std_7 at 7.7%). Objective 2 is fully achieved.")

add_heading(doc, "7.3.3 Objective 3 — OUTL inventory integration", 2)
add_para(doc,
    "The thesis integrated the Tier 4 stacking ensemble (for point forecasts) and the "
    "Tier 5 quantile model (for empirical safety stock) with a periodic (R, s, S) "
    "review policy. The OUTL parameters — safety stock, reorder point, order-up-to-"
    "level, order quantity, and annual cost components — are computed for each of the "
    "502 product-store series across four lead-time scenarios (L = 7, 10, 14, 21 days). "
    "Critically, the operational outcomes of each policy are measured by direct forward "
    "simulation of the (R, s, S) policy over the test window, rather than inferred "
    "analytically. This simulation-based measurement is itself a methodological choice "
    "that distinguishes the thesis from the comparable studies in the literature, "
    "which typically project an assumed service-level gap onto a parametric cost "
    "formula. Objective 3 is fully achieved.")

add_heading(doc, "7.3.4 Objective 4 — Quantification of operational impact", 2)
add_para(doc,
    "Objective 4 specified that the 10–25% MAPE improvement and 5–15% inventory cost "
    "reduction reported in the closely related studies of Barghi (2025) and Seyedan "
    "et al. (2023) be tested rather than treated as targets the thesis commits to "
    "meeting. The empirical results against each are as follows.")
add_para(doc,
    "On forecast accuracy, the stacking ensemble achieves 4.5% MAE improvement and "
    "9.5% RMSE improvement over the best classical baseline, with the MAPE improvement "
    "essentially flat (−0.3%). The MAE and MAPE results sit below the 10–25% range; "
    "the RMSE result is close to the lower bound. These figures are consistent with "
    "the M5 retrospective literature at SKU-daily granularity, which indicates that "
    "double-digit MAPE improvements at this granularity are achievable only under "
    "weekly or category-level aggregation. The reported literature figures do not, on "
    "the present evidence, transfer to SKU-daily intermittent demand.")
add_para(doc,
    "On inventory cost reduction, the simulation-based measurement reveals that all "
    "three forecasting policies — Classical, ML-Gaussian, and ML-Empirical-Quantile — "
    "achieve a realised service level of 100% on the test window and contain no "
    "stockouts. Total annual cost is dominated by holding cost; ML-Gaussian ties the "
    "Classical baseline to within 0.0% in total cost, and ML-Empirical-Quantile costs "
    "approximately 1.5% more because it allocates a larger safety stock without "
    "preventing any additional stockouts. The 5–15% inventory cost reduction reported "
    "in the literature does not therefore transfer to this dataset under direct "
    "simulation. The Chapter 1 target is not achieved.")
add_para(doc,
    "The thesis treats both shortfalls as findings rather than failures. The published "
    "benchmarks of 10–25% MAPE and 5–15% cost reduction were obtained on aggregated "
    "data where demand smoothing produces variance levels that the baseline policy "
    "cannot fully absorb; at SKU-daily granularity, with baseline (R, s, S) buffers "
    "generous enough to meet demand in full, the conditions under which those "
    "benchmarks hold do not obtain. The boundary condition identified in §5.6 specifies "
    "the regime in which ML forecasting and quantile-based safety stock would deliver "
    "operational value: namely, where the baseline policy faces genuine stockout risk. "
    "Objective 4 is therefore achieved in the form the supervisor explicitly approved — "
    "as a test of literature benchmarks rather than as a commitment to meet them — and "
    "the disconfirmation is itself the substantive finding.")

add_heading(doc, "7.3.5 Objective 5 — Power BI dashboard for SME deployment", 2)
add_para(doc,
    "The thesis implemented a Power BI dashboard with four pages (Forecast Overview, "
    "Inventory Status, Cost Analysis, Sensitivity Analysis), consuming five CSV data "
    "feeds exported from the Python pipeline and refreshed on a daily cycle. The "
    "dashboard is implemented in Microsoft Power BI Service (free tier), requires no "
    "commercial licence, and is delivered as a report that can be opened in any "
    "browser without further dependencies.")
add_para(doc,
    "Each page is designed to support a distinct operational decision: Forecast "
    "Overview surfaces actual demand, ML forecasts, classical baseline forecasts, and "
    "the LightGBM quantile fan for any selected product-store combination; Inventory "
    "Status presents safety stock, reorder point, and order-up-to-level parameters; "
    "Cost Analysis decomposes the simulated total annual cost by policy and lead time, "
    "making visible the §5.4 finding that holding cost dominates and stockout cost is "
    "zero across the three policies; Sensitivity Analysis presents the invariance "
    "result of §5.7 through coordinated line and bar charts. The dashboard supports "
    "the SME transferability commitment by translating the modelling outputs into a "
    "form accessible to non-technical managers. Objective 5 is fully achieved.")

# ── 7.4 PRINCIPAL CONTRIBUTIONS ──────────────────────────────────────────
add_heading(doc, "7.4 Principal Contributions", 1)
add_para(doc,
    "The thesis advances the demand forecasting and inventory optimisation literature "
    "through three theoretical contributions and three practical contributions, each "
    "summarised below and discussed in detail in §6.5 and §6.6.")
add_para(doc,
    "Theoretical contribution 1: First simulation-based characterisation on the M5 "
    "benchmark of the operational consequence of the Gaussian-versus-empirical safety-"
    "stock choice. The empirical quantile spread is 1.36 times the Gaussian buffer, but "
    "simulation shows that this distributional mismatch is operationally immaterial "
    "when the baseline buffer already meets demand in full. This refines the textbook "
    "treatment of safety stock (Silver et al., 2017) by supplying the empirical "
    "condition under which the Gaussian assumption fails operationally — and the "
    "condition under which it does not.")
add_para(doc,
    "Theoretical contribution 2: Demonstration that the operational value of ML "
    "forecasting in retail inventory management is conditional on the demand regime "
    "rather than automatic. The ML-Gaussian policy ties the Classical baseline at 0.0% "
    "in total cost, and the ML-Empirical-Quantile policy costs 1.5% more, on the M5 "
    "test window. The boundary condition under which the comparison would flip — "
    "namely, exposure of the baseline policy to genuine stockout risk — is itself the "
    "theoretical contribution. This refines the common framing in the forecasting "
    "literature in which improved point-forecast accuracy is presumed to translate "
    "directly into operational inventory savings.")
add_para(doc,
    "Theoretical contribution 3: Empirical robustness of the cost comparison to MAPE "
    "level and to the stockout-cost multiplier within the saturated-service regime. "
    "Across the 3–15% MAPE range, the simulated cost comparison varies by less than "
    "0.1 percentage points; across the 0.4× to 2.0× stockout-cost multiplier range, "
    "the comparison is unchanged because stockouts are zero and the multiplier never "
    "engages. This robustness profile clarifies that further investment in MAPE "
    "improvement is unlikely to yield inventory benefit in regimes of saturated "
    "service.")
add_para(doc,
    "Practical contribution 1: An end-to-end, reproducible, SME-deployable framework "
    "implemented in free and open-source tooling, runnable on a standard laptop, with "
    "a total runtime under three hours. The framework is documented to the level "
    "required for independent replication.")
add_para(doc,
    "Practical contribution 2: A forward-simulation engine usable as a pre-investment "
    "diagnostic for SME practitioners. Before committing to the more complex ML-"
    "quantile pipeline, an SME can run the engine on its own demand data to test "
    "whether its baseline policy is exposed to genuine stockout risk and therefore "
    "whether the additional methodology is likely to recover its cost. This diagnostic "
    "capability is itself a practitioner-facing deliverable of the thesis.")
add_para(doc,
    "Practical contribution 3: An interactive Power BI dashboard organised into four "
    "operational views, providing non-technical SME managers with access to forecast "
    "outputs, inventory parameters, cost decomposition, and sensitivity curves without "
    "requiring the user to understand the underlying methodology.")

# ── 7.5 CLOSING REMARKS ──────────────────────────────────────────────────
add_heading(doc, "7.5 Closing Remarks", 1)
add_para(doc,
    "The thesis began with a question that motivated both its academic positioning and "
    "its practical orientation: whether machine learning-driven demand forecasting, "
    "integrated with classical inventory policy theory, can produce operationally "
    "meaningful cost reductions in retail environments where small and medium-sized "
    "enterprises must compete with the analytics infrastructure of larger competitors. "
    "The answer that has emerged from the empirical work is nuanced and conditional: "
    "on the M5 SKU-daily test window evaluated here, when inventory outcomes are "
    "measured by direct simulation rather than inferred analytically, machine learning "
    "forecasting alone yields modest accuracy improvement and the quantile-based "
    "safety-stock reformulation does not lower inventory cost — because the baseline "
    "(R, s, S) policy already meets demand in full and the additional buffer protects "
    "against stockouts that do not occur. The contribution is therefore not a quoted "
    "cost saving but a quantified boundary: ML forecasting and distributional safety "
    "stock affect inventory outcomes when, and only when, the baseline policy is "
    "exposed to genuine stockout risk.")
add_para(doc,
    "This reframes the operational case for ML in retail inventory management. The "
    "conventional argument — that better forecasts produce better inventory decisions "
    "through improved point-forecast accuracy — is empirically conditional at the "
    "granularity at which inventory decisions are actually made. The stronger argument, "
    "supported by the data, is that the value of ML forecasting in retail inventory "
    "management is realised at the intersection of probabilistic forecasting and "
    "inventory-formulation refinement, but only in demand regimes where the baseline "
    "policy is not already adequate. The quantile model is the bridge between "
    "forecasting and inventory; the simulation engine is the diagnostic that "
    "determines whether the bridge is worth crossing in any particular setting.")
add_para(doc,
    "For the small and medium-sized retail sector specifically, the practical "
    "implication is qualified but actionable. The framework demonstrated in this "
    "thesis is implementable with free and open-source tooling on a standard laptop, "
    "requires only the data that off-the-shelf point-of-sale systems already produce, "
    "and is delivered with a pre-investment simulation diagnostic that the SME can run "
    "on its own data before committing further resource. For SMEs whose existing "
    "policy meets demand in full, the diagnostic will indicate that the more complex "
    "ML-quantile pipeline is unlikely to lower inventory cost and that the investment "
    "is better directed elsewhere. For SMEs whose existing policy faces genuine "
    "stockout risk — the regime in which the prior literature's cost-reduction "
    "benchmarks were obtained — the diagnostic will indicate that the framework is "
    "likely to pay off, and the dashboard supports the operational what-if analysis "
    "that follows.")
add_para(doc,
    "The thesis has therefore addressed both the academic question — how should ML "
    "forecasting be integrated with inventory policy theory for intermittent retail "
    "demand — and the practical question — how should SMEs decide whether such "
    "integration is worth the investment for their specific demand regime. The answer "
    "to both, with the qualifications documented in Chapters 5 and 6, is that the "
    "value of ML in retail inventory management is regime-dependent rather than "
    "automatic, and that a measurement-based diagnostic is more useful than a "
    "general-purpose cost-reduction claim. Future work should extend the framework to "
    "live deployment, validate it in actual SME contexts where the baseline policy "
    "faces stockout risk, and test whether the cost ranking of the policies flips in "
    "those regimes as the §5.6 boundary condition predicts.")
add_para(doc,
    "The wider research programme to which this thesis contributes — the operational "
    "integration of probabilistic machine learning with classical inventory theory — "
    "sits at the intersection of two mature literatures (demand forecasting and "
    "inventory optimisation) that have, for historical reasons, evolved largely in "
    "parallel. The opportunity to close the gap between them, both academically and "
    "operationally, is significant. The thesis offers one concrete step in that "
    "direction, in the specific context of small and medium-sized retail. The "
    "framework, the methodology, the simulation engine, and the empirical findings — "
    "including the disconfirmation finding — are delivered in a form intended to "
    "support both academic extension and operational deployment.")

OUT = "/Users/desmond/Capstone Project/retail-demand-forecasting/Chapter_7_Conclusion_v2.docx"
doc.save(OUT)
print(f"Saved: {OUT}")
