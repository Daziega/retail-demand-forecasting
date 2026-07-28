"""Insert the 4 Power BI dashboard screenshots into the supervisor's rewritten Chapter 5.

Strategy: open TFM_Chapter5_rewritten.docx, find the §5.8 placeholder paragraph
('Insert screenshots of each page...'), and insert four figure+caption blocks
right after the §5.8 prose, before the §5.9 Discussion heading.
"""
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy

SRC = "/Users/desmond/Capstone Project/retail-demand-forecasting/TFM_Chapter5_rewritten.docx"
OUT = "/Users/desmond/Capstone Project/retail-demand-forecasting/TFM_Chapter5_with_dashboards.docx"

FIG_DIR = "/Users/desmond/Capstone Project/retail-demand-forecasting/figures"

# Figure list: (filename, caption text)
FIGURES = [
    (
        f"{FIG_DIR}/dashboard_forecast_overview_sku.png",
        "Figure 5.1. Power BI dashboard — Forecast Overview page. KPI tiles show "
        "forecast-accuracy and realised-service-level summaries; the time series shows "
        "daily actual, ML-ensemble (Tier 4), and Croston-SBA classical forecasts over "
        "the test window for a selected SKU-store combination.",
    ),
    (
        f"{FIG_DIR}/dashboard_inventory_status.png",
        "Figure 5.2. Power BI dashboard — Inventory Status page. Summary cards report "
        "mean safety stock, reorder point, order-up-to-level, and total annual cost at "
        "the user-selected lead time, service level, and policy. The per-SKU table is "
        "sorted by total cost descending with conditional formatting flagging the "
        "highest safety-stock items. The bar chart compares total annual cost across "
        "the three policies.",
    ),
    (
        f"{FIG_DIR}/dashboard_cost_analysis.png",
        "Figure 5.3. Power BI dashboard — Cost Analysis page. KPI tiles report the "
        "central-case simulated cost decomposition (ML-Empirical-Quantile policy, "
        "L = 14, m = 1.0): $64,573 total annual cost, $0 stockout cost, holding cost "
        "dominating. The table reports realised service level by policy (all policies "
        "tie at 100%), the matrix breaks total cost down by lead time × policy, and "
        "the bar chart compares cost composition across the three policies at L = 14.",
    ),
    (
        f"{FIG_DIR}/dashboard_sensitivity_analysis.png",
        "Figure 5.4. Power BI dashboard — Sensitivity Analysis page. The line chart "
        "shows ML-Quantile cost reduction by lead time, sitting at approximately "
        "−1.5% across all four scenarios — well below the Chapter 1 5–15% target band "
        "(shown in green). The bar chart confirms the §5.7.1 invariance: cost change is "
        "identical across the three stockout-cost multipliers because realised stockouts "
        "are zero across all policies. The full sensitivity grid reports per-condition "
        "values, matching the appendix simulation results.",
    ),
]


def main():
    doc = Document(SRC)

    # Locate the §5.9 heading — figures will be inserted right before it
    # (i.e., at the end of §5.8 content)
    next_heading_idx = None
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        style = p.style.name if p.style else ""
        if text.startswith("5.9 ") and "Heading" in style:
            next_heading_idx = i
            break
    print(f"§5.9 heading found at paragraph index: {next_heading_idx}")

    # Insert the four figure+caption blocks just before §5.9.
    # python-docx adds at end; we insert by manipulating the underlying XML.
    if next_heading_idx is None:
        print("Could not locate §5.9 heading — appending to end of document.")
        next_p_element = None
    else:
        next_p_element = doc.paragraphs[next_heading_idx]._p

    for img_path, caption in FIGURES:
        # Create a centred paragraph for the image
        img_para = doc.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = img_para.add_run()
        run.add_picture(img_path, width=Inches(6.0))

        # Create a centred italic caption paragraph
        cap_para = doc.add_paragraph()
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_para.paragraph_format.space_after = Pt(18)
        cap_run = cap_para.add_run(caption)
        cap_run.italic = True
        cap_run.font.name = "Arial"
        cap_run.font.size = Pt(10)

        # Move these new paragraphs to just before §5.9
        if next_p_element is not None:
            next_p_element.addprevious(deepcopy(img_para._p))
            next_p_element.addprevious(deepcopy(cap_para._p))
            # Remove the duplicates we appended at end
            img_para._p.getparent().remove(img_para._p)
            cap_para._p.getparent().remove(cap_para._p)

    doc.save(OUT)
    print(f"\nSaved: {OUT}")
    print(f"Inserted {len(FIGURES)} figures with captions before §5.9.")


if __name__ == "__main__":
    main()
