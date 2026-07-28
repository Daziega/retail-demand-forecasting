"""Build a cleaned, deduplicated, alphabetised references list as a standalone docx.

Issues fixed (from supervisor feedback):
- Namwad et al. (2024) appeared twice (identical) — keep one
- Farooq et al. (2024) appeared twice (one abbreviated) — keep the full one
- Kalusivalingam et al. (2024) "IEEE Transactions... preprint" flagged for
  venue verification — annotated, not removed
- Croston 1972, Hyndman & Athanasopoulos 2021, Syntetos & Boylan 2005 added
  as new references from the revised Chapter 3

User instruction: open this docx, copy the references body, paste over the
References section of the assembled thesis (replacing the duplicated/messy
existing one). Confirm in-text citations still match.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = "/Users/desmond/Capstone Project/retail-demand-forecasting/TFM_References_Cleaned.docx"


# Sorted alphabetically by first author surname.
# [VERIFY] flags entries the supervisor said need checking.
REFERENCES = [
    "Aryyama, J. K. (2021). Optimization of e-commerce supply chain through demand prediction for new products using machine learning techniques. Journal of Artificial Intelligence and Machine Learning Data Science, 1(1), 565–569. https://doi.org/10.51219/JAIMLD/149",
    "Bansal, A. (2023). Role of machine learning in inventory optimization using time-series forecasting. IIT Delhi Research Reports.",
    "Barghi, S. (2025). Demand forecasting and inventory improvement in supply chain management using hybrid boosting ensemble techniques. Master's thesis, École de Technologie Supérieure, Université du Québec.",
    "Bastos, A. S. T. (2023). Machine learning in digital retail demand forecasting for inventory management in a sportswear company. Master's thesis, NOVA Information Management School, Universidade Nova de Lisboa.",
    "Bendhi, M. R. (2023). Inventory management: Machine learning predicts demand, reducing excess stock by up to 20%. International Journal of Artificial Intelligence, Data Science, and Machine Learning, 4(2), 67–83. https://doi.org/10.63282/3050-9262.IJAIDSML-V4I2P108",
    "Chetlapalli, H., & Vinayagam, S. (2019). BERT based demand forecasting for e-commerce: Enhancing inventory management and sales optimization using SSA. International Journal of Multidisciplinary and Current Research, 7, 459–466.",
    "Croston, J. D. (1972). Forecasting and stock control for intermittent demands. Operational Research Quarterly, 23(3), 289–303.",
    "Farooq, A., Abbey, A. B. N., & Onukwulu, E. C. (2024). Inventory optimization and sustainability in retail: A conceptual approach to data-driven resource management. International Journal of Multidisciplinary Research and Growth Evaluation, 5(6), 1356–1363.",
    "Harris, F. W. (1913). How many parts to make at once. The Magazine of Management, 10(2), 135–136.",
    "Haque, M. S., Amin, M. S., & Miah, J. (2023). Retail demand forecasting: A comparative study for multivariate time series (arXiv:2308.11939). arXiv. https://arxiv.org/abs/2308.11939",
    "Hughes, R. (2023). Machine learning for optimizing inventory management in retail. American Journal of Machine Learning, 4(1), 9–18.",
    "Hyndman, R. J., & Athanasopoulos, G. (2021). Forecasting: Principles and practice (3rd ed.). OTexts.",
    "Islam, M. T., Ayon, E. H., Ghosh, B. P., Chowdhury, M. S., Shahid, R., Roy Puja, A., Rahman, S., Bhuiyan, M. S., & Nguyen, T. N. (2024). Revolutionizing retail: A hybrid machine learning approach for precision demand forecasting and strategic decision-making in global commerce. Journal of Computer Science and Technology Studies.",
    "Ismail, U., Khosa, S. N., Tahir, S., Ahmad, M. A., Hussain, W., Akram, U., & Mushtaq, M. F. (2025). Hybrid machine learning models for optimizing retail market and inventory forecasting. Journal of Computing & Biomedical Informatics, 9(1). https://doi.org/10.56979/901/2025",
    "Jampani, S., Avancha, S., Mangal, A., Singh, S. P., Jain, S., & Agarwal, R. (2023). Machine learning algorithms for supply chain optimisation. International Journal of Research in Modern Engineering and Emerging Technology, 11(4), 216–224.",
    "Kalisetty, S. (2023). Harnessing big data and deep learning for real-time demand forecasting in retail: A scalable AI-driven approach. American Online Journal of Science and Engineering, 1(1), 1–15.",
    "Kalusivalingam, A. K., Sharma, A., Patel, N., & Singh, V. (2024). Optimizing inventory management with AI: Leveraging deep reinforcement learning and neural networks for enhanced demand forecasting and stock replenishment. [VERIFY VENUE — flagged by supervisor: original entry said 'IEEE Transactions on Industrial Informatics (preprint)' which is unusual for that journal. Confirm against original or remove.]",
    "Kar, A., Ganipaneni, S., Kshirsagar, R. P., Goel, O., Jain, A., & Goel, P. (2021). Demand forecasting optimization: Advanced ML models for retail and inventory planning. International Research Journal of Modernization in Engineering Technology and Science, 3(10), 1401–1410.",
    "Keith, E. (2023). Optimizing inventory management through advanced forecasting techniques in supply chains. European Journal of Supply Chain Management, 1(1), 22–30.",
    "Khan, M. A., Saqib, S., Alyas, T., Rehman, A. U., Saeed, Y., Zeb, A., Zareei, M., & Mohamed, E. M. (2020). Effective demand forecasting model using business intelligence empowered with machine learning. IEEE Access, 8, 116013–116023. https://doi.org/10.1109/ACCESS.2020.3003790",
    "Khayyat, M. M., & Gupta, S. K. (2025). Advanced predictive model for optimizing inventory management and demand forecasting in smart logistics. SGS Engineering & Sciences, 1(1).",
    "Kilimci, Z. H., Akyuz, A. O., Uysal, M., Akyokus, S., Uysal, M. O., Bulbul, B. A., & Ekmis, M. A. (2019). An improved demand forecasting model using deep learning approach and proposed decision integration strategy for supply chain. Complexity, 2019, Article 9067367. https://doi.org/10.1155/2019/9067367",
    "Kumar, V. D., Maheswari, S., Raman, Y. S., Iniyavan, S., & Hashim, A. S. (2024). Inventory optimization and demand forecasting using machine learning. International Journal of Research and Review.",
    "Loa, J., Wiratama, J., & Halim, F. A. (2024). Optimizing inventory management in retail companies through sales prediction using XGBoost. Proceedings of the International Conference on Data Science.",
    "Maharana, P., Choudhury, R., & Vimaladevi, M. (2025). Inventory supply optimization with demand forecasting. SSRN. https://ssrn.com/abstract=5832182",
    "Marripudugala, M. (2024). Optimizing inventory management with deep learning in retail supply chains. International Research Journal of Modernization in Engineering Technology and Science, 6(7), 3109–3116. https://doi.org/10.56726/IRJMETS60329",
    "Micheal, L. (2023). Machine learning for demand forecasting and inventory optimization in CPG. Industry white paper.",
    "Mitra, A., Jain, A., Kishore, A., & Kumar, P. (2022). A comparative study of demand forecasting models for a multi-channel retail company: A novel hybrid machine learning approach. Operations Research Forum, 3, 58. https://doi.org/10.1007/s43069-022-00166-4",
    "Mittal, V. K. (2025). Inventory optimization using machine learning: Advanced forecasting for multi-channel supply chains. SSRN. https://ssrn.com/abstract=5386001",
    "Mohammed, I. A., & Mandal, J. (2022). Forecasting accuracy through machine learning in supply chain management. International Journal of Supply Chain Management, 7(2), 60–77.",
    "Mweshi, G. K. (2025). Optimizing smart commerce: A machine learning approach to dynamic demand forecasting and inventory control. International Journal of Advances in Engineering and Management, 7(10), 450–460. https://doi.org/10.35629/5252-0710450460",
    "Namwad, R. S., Mishra, N. K., Ranu, & Jain, P. (2024). Optimizing inventory management with seasonal demand forecasting in a fuzzy environment. Journal of the European Systems Automation, 57(4). https://doi.org/10.18280/jesa.570416",
    "Nasseri, M., Falatouri, T., Brandtner, P., & Darbanian, F. (2023). Applying machine learning in retail demand prediction — A comparison of tree-based ensembles and long short-term memory-based deep learning. Applied Sciences, 13(19), 11112. https://doi.org/10.3390/app131911112",
    "Nasution, A. A., Matondang, N., & Ishak, A. (2022). Inventory optimization model design with machine learning approach in feed mill company. Jurnal Sistem Teknik Industri, 24(2), 254–272. https://doi.org/10.32734/jsti.v24i2.8637",
    "Obi, C. I. C. (2024). Demand forecasting in retail business using the ensemble machine learning framework: A stacking approach. American Academic Scientific Research Journal for Engineering, Technology, and Sciences, 98(1), 309–329.",
    "Omniful. (n.d.). Inventory forecasting: Challenges, benefits, and best practices. [Industry source, cited for contextual framing only.]",
    "Pasupuleti, V., Thuraka, B., Kodete, C. S., & Malisetty, S. (2024). Enhancing supply chain agility and sustainability through machine learning: Optimization techniques for logistics and inventory management. Logistics, 8(3), 73. https://doi.org/10.3390/logistics8030073",
    "Praveen, K. B., Pradyumna Kumar, Prateek, J., Pragathi, G., & Madhuri, J. (2023). Inventory management using machine learning. IEEE conference proceedings.",
    "Punia, S., Nikolopoulos, K., Prakash Singh, S., Madaan, J. K., & Litsiou, K. (2020). Deep learning with long short-term memory networks and random forests for demand forecasting in multi-channel retail. International Journal of Production Research, 58(16), 4964–4979. https://doi.org/10.1080/00207543.2020.1735666",
    "Razzak, A., Paul, R., & Rozony. (2025). Demand forecasting in retail e-commerce: A systematic review. International Journal of Scientific Interdisciplinary Research, 6(1), 1–27. https://doi.org/10.63125/mbbfw637",
    "Sekhar, C. (2022). Optimizing retail inventory management with AI: A predictive approach to demand forecasting, stock optimization, and automated reordering. European Journal of Advances in Engineering and Technology, 9(11), 89–94.",
    "Seyedan, M. (2023). Development of predictive analytics for demand forecasting and inventory management in supply chain using machine learning techniques [Doctoral dissertation]. Concordia University, Montreal.",
    "Seyedan, M., Mafakheri, F., & Wang, C. (2022). Cluster-based demand forecasting using Bayesian model averaging: An ensemble learning approach. Decision Analytics Journal, 3, 100033. https://doi.org/10.1016/j.dajour.2022.100033",
    "Seyedan, M., Mafakheri, F., & Wang, C. (2023). Order-up-to-level inventory optimization model using time-series demand forecasting with ensemble deep learning. Supply Chain Analytics, 3, 100024. https://doi.org/10.1016/j.sca.2023.100024",
    "Sihotang, J. (2023). Optimization of inventory ordering decision in retail business using exponential smoothing approach and decision support system. International Journal of Mechanical Computational and Manufacturing Research, 12(2), 45–52.",
    "Silver, E. A., Pyke, D. F., & Thomas, D. J. (2017). Inventory and production management in supply chains (4th ed.). CRC Press.",
    "Singhal, K., Singh, V., & Kaul, A. (2024). Smart retail: Utilizing machine learning for demand prediction, price strategy, and inventory management. 16th IEEE International Conference on Computational Intelligence and Communication Networks, 485–490.",
    "Sivilai, H. (2025). Application of machine learning in demand forecasting and inventory optimization. Journal of Cyber-Physical Security and Robotics, 1(1), 18–21.",
    "Stanelyte, G. (2021). Inventory optimization in retail network by creating a demand prediction model [Master's thesis]. Vilnius Gediminas Technical University.",
    "Syntetos, A. A., & Boylan, J. E. (2005). The accuracy of intermittent demand estimates. International Journal of Forecasting, 21(2), 303–314.",
    "Taparia, V., Mishra, P., Gupta, N., & Chandiramani, H. (2023). Data-driven retail excellence: Machine learning for demand forecasting and price optimization. Journal of Graphic Era University, 12(1), 37–60.",
    "Taparia, V., Mishra, P., Gupta, N., & Kumar, D. (2023). Improved demand forecasting of a retail store using a hybrid machine learning model. Journal of Graphic Era University, 12(1), 15–36. https://doi.org/10.13052/jgeu0975-1416.1212",
    "Tredence. (n.d.). Retail demand forecasting: Methods, benefits, and challenges. [Industry source, cited for contextual framing only.]",
    "Ungureanu, D. A. (2025). Demand forecasting and inventory optimization in mid-sized grocery retail using machine learning: A data-driven approach to minimizing stock-outs and waste [Master's thesis]. CCT College Dublin.",
    "Wahedi, H. J., Heltoft, M., Christophersen, G. J., Severinsen, T., Saha, S., & Nielsen, I. E. (2023). Forecasting and inventory planning: An empirical investigation of classical and machine learning approaches for Svanehøj's future software consolidation. Applied Sciences, 13(15), 8581. https://doi.org/10.3390/app13158581",
    "Yadav, P. (2025). Demand forecasting in retail using machine learning and big data. International Journal of Advanced Research in Computer and Communication Engineering, 14(2), 274–278. https://doi.org/10.17148/IJARCCE.2025.14235",
    "Zhou, Z. H., Wu, J., & Tang, W. (2002). Ensembling neural networks: Many could be better than all. Artificial Intelligence, 137(1–2), 239–263.",
    "Zohdi, M., Rafiee, M., Kayvanfar, V., & Salamiraad, A. (2022). Demand forecasting based machine learning algorithms on customer information: An applied approach. International Journal of Information Technology, 14, 1937–1947. https://doi.org/10.1007/s41870-022-00875-3",
]


def main():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # Title
    h = doc.add_heading("References", level=1)
    for run in h.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(0, 0, 0)

    # Editorial note
    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(12)
    nrun = note.add_run(
        "Editorial note (delete before submission): cleaned and deduplicated against "
        "the prior bibliography per supervisor feedback. Removed duplicate Namwad "
        "(2024) and abbreviated duplicate of Farooq et al. (2024); added Croston "
        "(1972), Hyndman & Athanasopoulos (2021) and Syntetos & Boylan (2005) "
        "introduced in the revised Chapter 3; flagged Kalusivalingam et al. (2024) "
        "for venue verification."
    )
    nrun.italic = True
    nrun.font.name = "Arial"
    nrun.font.size = Pt(9)
    nrun.font.color.rgb = RGBColor(180, 0, 0)

    # Reference list — APA hanging indent
    for ref in REFERENCES:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.first_line_indent = Inches(-0.4)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(ref)
        r.font.name = "Arial"
        r.font.size = Pt(10)

    doc.save(OUT)
    print(f"Saved: {OUT}")
    print(f"Total references: {len(REFERENCES)}")


if __name__ == "__main__":
    main()
