"""Build the Reproducibility appendix: pinned library versions, chosen
hyperparameters per model, and pointer to the .pbix file.

Versions are pinned to the project's working environment at submission time.
The user should regenerate the version table from their final environment via
`pip freeze` and replace the table values if anything has shifted.
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/Users/desmond/Capstone Project/retail-demand-forecasting/TFM_Appendix_Reproducibility.docx"


def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_para(doc, text, *, italic=False, size=11, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.italic = italic
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.name = "Arial"
        r.font.size = Pt(10)
        set_cell_shading(hdr_cells[i], "D5E8F0")
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for r_idx, row in enumerate(rows, start=1):
        cells = table.rows[r_idx].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ""
            p = cells[c_idx].paragraphs[0]
            r = p.add_run(str(val))
            r.font.name = "Arial"
            r.font.size = Pt(10)
    return table


# ─────────────────────────────────────────────────────────────────────────
doc = Document()
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(11)

# Appendix title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("APPENDIX A")
run.bold = True; run.font.name = "Arial"; run.font.size = Pt(18)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub.add_run("Reproducibility — Library Versions, Hyperparameters, and Deliverables")
sub_run.bold = True; sub_run.italic = True; sub_run.font.name = "Arial"; sub_run.font.size = Pt(14)
sub.paragraph_format.space_after = Pt(18)

# A.1 Library versions
add_heading(doc, "A.1 Library Versions", 1)
add_para(doc,
    "All experiments documented in Chapters 3 through 5 were executed on a "
    "single MacBook running macOS 14 with 16 GB RAM, no GPU, using Python 3.11 in "
    "a dedicated virtual environment. The pinned versions of all major dependencies "
    "are listed below. The full list (including transitive dependencies) is "
    "obtainable by running `pip freeze` inside the project's retail_forecasting_env "
    "environment.")

add_para(doc, "Table A.1: Pinned library versions used in the experiments",
         italic=True, size=10, space_after=4)

add_table(doc,
    headers=["Package", "Pinned version", "Role in pipeline"],
    rows=[
        ["Python",        "3.11",     "Interpreter"],
        ["pandas",        "2.2.3",    "Tabular data manipulation"],
        ["numpy",         "1.26.4",   "Numerical arrays, simulation engine"],
        ["scikit-learn",  "1.5.2",    "Random Forest, train/test split utilities, scaling"],
        ["lightgbm",      "4.5.0",    "Gradient boosting (Tier 2) and quantile regression (Tier 5)"],
        ["xgboost",       "2.1.2",    "Gradient boosting with Tweedie loss (Tier 2)"],
        ["tensorflow",    "2.16.2",   "LSTM (Tier 3) via Keras API"],
        ["statsforecast", "1.7.6",    "Tier 1 classical baselines (Croston, AutoARIMA, etc.)"],
        ["statsmodels",   "0.14.2",   "Reserved for legacy ARIMA cross-checks"],
        ["shap",          "0.46.0",   "Feature importance attribution (§3.10, §4.11)"],
        ["matplotlib",    "3.9.2",    "Plot generation for thesis figures"],
        ["seaborn",       "0.13.2",   "Plot styling"],
        ["pyarrow",       "17.0.0",   "Parquet I/O"],
        ["joblib",        "1.4.2",    "Model serialisation (.pkl)"],
        ["python-docx",   "1.1.2",    "Programmatic generation of chapter docx files"],
    ])

add_para(doc, "", space_after=12)
add_para(doc,
    "All packages are MIT, BSD, or Apache 2.0 licensed and freely available via "
    "pip. No commercial or enterprise dependency is required.", italic=True, size=10)

# A.2 Hyperparameters
add_heading(doc, "A.2 Final Chosen Hyperparameters", 1)
add_para(doc,
    "Table A.2 documents the final hyperparameter values selected for each model "
    "after the two-stage search described in §3.6.3 (coarse random search followed "
    "by grid refinement under walk-forward CV).")

add_para(doc, "Table A.2: Final hyperparameter values per model",
         italic=True, size=10, space_after=4)

add_table(doc,
    headers=["Model", "Hyperparameter", "Final value"],
    rows=[
        ["LightGBM (Tweedie)",   "n_estimators",         "2000 (early-stopped on val MAE)"],
        ["",                     "learning_rate",        "0.05"],
        ["",                     "num_leaves",           "63"],
        ["",                     "min_child_samples",    "50"],
        ["",                     "subsample",            "0.8"],
        ["",                     "colsample_bytree",     "0.8"],
        ["",                     "objective",            "tweedie"],
        ["",                     "tweedie_variance_power", "1.1"],
        ["XGBoost (Tweedie)",    "n_estimators",         "2000 (early-stopped)"],
        ["",                     "learning_rate",        "0.05"],
        ["",                     "max_depth",            "6"],
        ["",                     "subsample",            "0.8"],
        ["",                     "colsample_bytree",     "0.8"],
        ["",                     "reg_lambda",           "1.0"],
        ["",                     "objective",            "reg:tweedie"],
        ["",                     "tweedie_variance_power", "1.1"],
        ["Random Forest",        "n_estimators",         "300"],
        ["",                     "max_depth",            "20"],
        ["",                     "min_samples_leaf",     "10"],
        ["LSTM",                 "layers",               "Stacked LSTM(128) → Dropout(0.2) → LSTM(64) → Dropout(0.2) → Dense(1)"],
        ["",                     "lookback window",      "28 days"],
        ["",                     "target transform",     "log1p (inverse: expm1, clip ≥ 0)"],
        ["",                     "loss",                 "MSE in log space"],
        ["",                     "optimiser",            "Adam (lr 1e-3)"],
        ["",                     "batch size",           "256"],
        ["",                     "max epochs",           "50 (early stopping, patience 5)"],
        ["Stacking ensemble",    "base learners",        "LightGBM, XGBoost, Random Forest"],
        ["",                     "meta-learner",         "Ridge regression (sklearn)"],
        ["",                     "Ridge α",              "1.0"],
        ["",                     "Ridge fit_intercept",  "True"],
        ["",                     "Ridge positive",       "True (non-negative weights)"],
        ["LightGBM quantile",    "quantile levels α",    "0.10, 0.50, 0.90, 0.95, 0.99"],
        ["",                     "n_estimators",         "2000 per quantile (early-stopped)"],
        ["",                     "learning_rate",        "0.05"],
        ["",                     "num_leaves",           "63"],
        ["",                     "objective",            "quantile"],
        ["AutoARIMA",            "season_length",        "7 (weekly)"],
        ["",                     "selection criterion",  "AIC (statsforecast default)"],
        ["Croston-SBA",          "α (smoothing)",        "default (statsforecast auto)"],
    ])

# A.3 OUTL/economic constants
add_heading(doc, "A.3 OUTL and Economic Constants", 1)
add_para(doc,
    "The (R, s, S) inventory policy and cost calculations use the constants "
    "specified in methodology §3.7.2:")

add_table(doc,
    headers=["Constant", "Symbol", "Value", "Source"],
    rows=[
        ["Review period",          "R",          "7 days",                    "§3.7 (weekly review, SME standard)"],
        ["Carrying rate",          "h%",         "25% of unit cost / year",    "Silver, Pyke & Thomas (2017)"],
        ["Ordering cost",          "K",          "$50 per order",              "SME typical"],
        ["Gross margin",           "m",          "40% (unit_cost = 0.6 × price)", "Retail industry typical"],
        ["Target service level",   "SL",         "95%",                        "Chapter 1 primary target"],
        ["z-score @ 95% SL",       "z",          "1.6449",                     "Standard normal inverse CDF"],
        ["Days per year",          "—",          "365",                        "Annualisation"],
        ["Lead-time scenarios",    "L",          "7, 10, 14, 21 days",         "§3.7.3"],
        ["Stockout cost multipliers","m_sc",     "0.4, 1.0, 2.0",              "§3.7.5"],
    ])

# A.4 Deliverables
add_heading(doc, "A.4 Deliverables Inventory", 1)
add_para(doc,
    "The following artefacts accompany the thesis (available in the project "
    "repository):")

add_table(doc,
    headers=["Artefact", "Path", "Purpose"],
    rows=[
        ["Source code (notebooks)",        "notebooks/01_eda.ipynb … 12_dashboard_exports.ipynb",
         "End-to-end pipeline from raw M5 ingestion to Power BI exports"],
        ["Shared utilities",                "src/forecasting_utils.py",
         "Chronological split, walk-forward folds, MAE/RMSE/MAPE/WMAPE/Pred10/Pinball"],
        ["Simulation engine (supervisor's script)", "inventory_simulation.py",
         "Forward (R, s, S) simulation used for Chapter 5 measurements"],
        ["Falsification tests",             "falsification_tests.py",
         "Engine validation for §5.3"],
        ["Stratified subsample data",       "data/processed/subsample_features.parquet",
         "502 product-store series used in all experiments"],
        ["Simulation results",              "data/processed/simulation_results.csv",
         "Full grid of (L × m × policy) simulated cost outcomes"],
        ["Dashboard CSV exports",           "data/processed/powerbi/*.csv",
         "Five data feeds for Power BI: forecasts, inventory_params, cost_analysis, sensitivity, kpi_summary"],
        ["Power BI dashboard file",         "Retail_Demand_Dashboard_Report.pbix",
         "Live dashboard with the four pages described in §5.8. Submitted alongside the thesis; open in Power BI Desktop to interact with all four report pages."],
        ["Trained model artefacts",         "models/lightgbm.pkl, xgboost.pkl, random_forest.pkl, lstm_model.keras, lgbm_quantile_models.pkl",
         "Serialised models for re-prediction without retraining"],
        ["Build scripts (this thesis)",     "build_chapter3.py, …, build_chapter7.py, build_front_matter.py, build_references_clean.py, build_reproducibility_appendix.py, insert_*.py",
         "Reproducible regeneration of the docx chapters from source data"],
        ["Final thesis assembly",           "[combined docx — assembled by user from chapter files]",
         "Single-document thesis for submission"],
    ])

# A.5 How to reproduce
add_heading(doc, "A.5 How to Reproduce", 1)
add_para(doc,
    "To reproduce the experiments end-to-end:")
add_para(doc,
    "  1. Clone the repository and create a Python 3.11 virtual environment "
    "(retail_forecasting_env).")
add_para(doc,
    "  2. Install pinned dependencies via `pip install -r requirements.txt` (the "
    "requirements file in the repository root lists the same packages as Table A.1, "
    "with versions adjusted to the user's platform).")
add_para(doc,
    "  3. Download the M5-Forecasting (Walmart) dataset from Kaggle and place the "
    "three raw files (sales_train_validation.csv, calendar.csv, sell_prices.csv) "
    "into data/raw/.")
add_para(doc,
    "  4. Run the notebooks in numerical order (01_eda.ipynb → 12_dashboard_"
    "exports.ipynb). Total runtime: under three hours on the specified hardware.")
add_para(doc,
    "  5. Run inventory_simulation.py (after setting `use_real = True` via "
    "the default behaviour of the script's main()) to regenerate "
    "simulation_results.csv.")
add_para(doc,
    "  6. Run falsification_tests.py to regenerate the engine-validation table "
    "of §5.3.")
add_para(doc,
    "  7. Open Retail_Demand_Dashboard_Report in Power BI Service (or load the "
    ".pbix in Power BI Desktop), refresh the data connection to the five CSVs in "
    "data/processed/powerbi/, and export each of the four pages as PNG for the "
    "figures of §5.8.")

# A.6 Full simulation grid (referenced by Chapter 5 §5.2)
add_heading(doc, "A.6 Full Simulation Grid", 1)
add_para(doc,
    "Table A.4 reproduces the complete forward-simulation output referenced in §5.2 "
    "(simulation_results.csv): every combination of lead time (7, 10, 14, 21 days) and "
    "stockout-cost multiplier (0.4, 1.0, 2.0) for each of the three policies, at the "
    "primary 95% service-level target. All monetary values are annualised. Because no "
    "policy stocks out on the test window, the realised service level is 100% and the "
    "stockout cost (SC) is zero throughout, and the total cost (TC) is invariant to the "
    "stockout multiplier — the three multiplier rows for a given lead time and policy "
    "are therefore identical.")

import csv as _csv
_sim_rows = list(_csv.reader(open(
    "/Users/desmond/Capstone Project/retail-demand-forecasting/"
    "data/processed/simulation_results.csv")))
_hdr = _sim_rows[0]
_data = _sim_rows[1:]

# Prettify policy labels and format numbers
_policy_map = {
    "CLASSICAL": "Classical",
    "ML_GAUSSIAN": "ML-Gaussian",
    "ML_EMP_QUANTILE": "ML-Emp-Quantile",
}


def _fmt(v):
    try:
        f = float(v)
        if f == int(f):
            return f"{int(f):,}"
        return f"{f:,.2f}"
    except (ValueError, TypeError):
        return v


_table_rows = []
for r in _data:
    lead, mult, policy, sl, ss, hc, oc, sc, tc = r
    _table_rows.append([
        lead,
        mult,
        _policy_map.get(policy, policy),
        f"{float(sl):.1f}%",
        _fmt(ss),
        _fmt(hc),
        _fmt(oc),
        _fmt(sc),
        _fmt(tc),
    ])

add_para(doc, "Table A.4: Full forward-simulation grid (SL target = 95%)",
         italic=True, size=10)
add_table(doc,
    headers=["L (d)", "m", "Policy", "Realised SL", "SS (units)",
             "HC ($)", "OC ($)", "SC ($)", "TC ($)"],
    rows=_table_rows)

# A.7 Falsification test — per-policy detail
add_heading(doc, "A.7 Falsification Test — Per-Policy Detail", 1)
add_para(doc,
    "The §5.3 falsification test establishes that the simulation engine registers "
    "stockouts correctly. Table 5.1 in the main text reports the aggregate (baseline-"
    "policy) result under each stress condition. Table A.5 supplements this with the "
    "per-policy breakdown under the demand ×3 stress condition (central case L = 14, "
    "target SL = 95%, stockout multiplier m = 1.0), which is used in the boundary-"
    "condition discussion of §5.6. Under a threefold demand shock, realised service "
    "falls to approximately 94% under all three policies; the ML-Empirical-Quantile "
    "policy ties the Classical baseline (+0.0% in total cost) while the ML-Gaussian "
    "policy is marginally worse (−0.5%). Even a threefold demand shock is therefore "
    "insufficient, on this subsample, to make the empirically calibrated quantile "
    "buffer pay off operationally.")

add_para(doc, "Table A.5: Per-policy outcome under the demand ×3 stress "
         "(L = 14, SL = 95%, m = 1.0)", italic=True, size=10)
add_table(doc,
    headers=["Policy", "Realised SL", "Stockout units", "Stockout cost ($)",
             "Total cost ($)", "vs Classical"],
    rows=[
        ["Classical",        "94.1%", "6,328", "87,691", "154,905", "—"],
        ["ML-Gaussian",      "94.0%", "6,418", "88,447", "155,659", "−0.5%"],
        ["ML-Emp-Quantile",  "94.1%", "6,322", "86,629", "154,832", "+0.0%"],
    ])

add_para(doc, "")
add_para(doc,
    "The four stress conditions of Table 5.1 (§5.3) confirm that the engine measures "
    "stockouts across the full range of severity: from a 74.0% realised service level "
    "and $124,338 stockout cost when safety stock is set to zero, through 94.1% under "
    "the demand ×3 shock, to a 0.0% service level and $469,556 stockout cost when "
    "replenishment is fully suppressed. The 100% realised service level reported for "
    "the baseline configuration therefore reflects demand genuinely being met, not a "
    "stockout counter that never fires.", italic=True, size=10)

doc.save(OUT)
print(f"Saved: {OUT}")
