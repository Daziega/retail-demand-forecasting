"""Post-process the compiled thesis:
  1. Fix 2 Option A — insert hidden SEQ Figure / SEQ Table fields into every
     figure/table caption so Word's Table-of-Figures collection populates the
     List of Figures and List of Tables automatically.
  2. Rebuild the malformed TOC / List-of-Tables / List-of-Figures fields as
     well-formed Word fields (so right-click → Update Field works).
  3. Normalise stray Normal-styled section headings to Heading styles so the
     TOC is complete.
  4. Page break before each chapter / appendix / references.
  5. Justify all body text and set 1.5 line spacing.
"""
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = "/Users/desmond/Capstone Project/retail-demand-forecasting/TFM_Final.docx"
OUT = "/Users/desmond/Capstone Project/retail-demand-forecasting/TFM_Final_formatted.docx"

FIG_RE  = re.compile(r'^Figure\s+\d')
TAB_RE  = re.compile(r'^Table\s+[A-Za-z0-9.]+:')          # true caption -> colon
SEC_RE  = re.compile(r'^\d+\.\d+\s+[A-Za-z(]')            # e.g. "4.2 Title"
SUB_RE  = re.compile(r'^\d+\.\d+\.\d+\s+[A-Za-z(]')       # e.g. "4.2.1 Title"

CHAPTER_TITLES = {
    "Chapter 1: Introduction", "CHAPTER 2", "CHAPTER 3", "CHAPTER 4",
    "CHAPTER 5", "CHAPTER 6", "CHAPTER 7", "APPENDIX A", "References",
}


def _run(txt=None, field=None, hidden=False):
    """Build a <w:r>. If field in {begin,separate,end} make a fldChar run;
    if field=='instr' put instrText=txt; else a plain text run."""
    r = OxmlElement("w:r")
    if hidden:
        rpr = OxmlElement("w:rPr")
        v = OxmlElement("w:vanish")
        rpr.append(v)
        r.append(rpr)
    if field in ("begin", "separate", "end"):
        fc = OxmlElement("w:fldChar")
        fc.set(qn("w:fldCharType"), field)
        r.append(fc)
    elif field == "instr":
        it = OxmlElement("w:instrText")
        it.set(qn("xml:space"), "preserve")
        it.text = txt
        r.append(it)
    else:
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = txt or ""
        r.append(t)
    return r


def insert_seq_at_start(p, label):
    """Insert a hidden `SEQ <label> \\h` field at the very start of p. The \\h
    switch suppresses all visible output, so the caption text is untouched and
    the field exists only as a collection marker for TOC \\c."""
    pElem = p._p
    pPr = pElem.find(qn("w:pPr"))
    anchor = pPr if pPr is not None else None
    seq_runs = [
        _run(field="begin"),
        _run(txt=f" SEQ {label} \\h ", field="instr"),
        _run(field="end"),
    ]
    for i, r in enumerate(seq_runs):
        if anchor is not None:
            anchor.addnext(r)
            anchor = r
        else:
            pElem.insert(i, r)


def rebuild_field(p, instr, placeholder):
    """Clear all runs in p and rebuild a well-formed field."""
    pElem = p._p
    # remove existing runs (keep pPr)
    for r in pElem.findall(qn("w:r")):
        pElem.remove(r)
    for r in [
        _run(field="begin"),
        _run(txt=instr, field="instr"),
        _run(field="separate"),
        _run(txt=placeholder),
        _run(field="end"),
    ]:
        pElem.append(r)


def main():
    doc = Document(SRC)

    # ---- 1 + 3: captions get SEQ fields; stray section headings normalised ----
    n_fig = n_tab = n_sec = n_sub = 0
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        if FIG_RE.match(t):
            insert_seq_at_start(p, "Figure")
            n_fig += 1
        elif TAB_RE.match(t):
            insert_seq_at_start(p, "Table")
            n_tab += 1
        # heading normalisation (only promote Normal-styled numbered headings)
        if p.style.name == "Normal":
            if SUB_RE.match(t):
                p.style = doc.styles["Heading 2"]
                n_sub += 1
            elif SEC_RE.match(t):
                p.style = doc.styles["Heading 1"]
                n_sec += 1

    # ---- 2: rebuild the three malformed front-matter fields ----
    fixed = {"toc": False, "table": False, "figure": False}
    to_delete = []
    for p in doc.paragraphs:
        xml = p._p.xml
        # classify by instr content
        if "instrText" in xml and "TOC" in xml:
            if '"Table"' in xml:
                rebuild_field(p, r' TOC \h \z \c "Table" ',
                              "List of Tables — right-click and Update Field.")
                fixed["table"] = True
            elif '"Figure"' in xml:
                rebuild_field(p, r' TOC \h \z \c "Figure" ',
                              "List of Figures — right-click and Update Field.")
                fixed["figure"] = True
            else:
                rebuild_field(p, r' TOC \o "1-3" \h \z \u ',
                              "Table of Contents — right-click and Update Field.")
                fixed["toc"] = True
        # mark helper-note paragraphs for deletion
        tt = p.text.strip()
        if tt.startswith("(Place cursor") or tt.startswith("(Right-click and"):
            to_delete.append(p)

    for p in to_delete:
        p._p.getparent().remove(p._p)

    # ---- 4: page break before each chapter/appendix/references ----
    n_break = 0
    for p in doc.paragraphs:
        if p.text.strip() in CHAPTER_TITLES:
            p.paragraph_format.page_break_before = True
            n_break += 1

    # ---- 5: justify body text + 1.5 line spacing ----
    n_just = 0
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        style = p.style.name
        if style.startswith("Heading"):
            p.paragraph_format.line_spacing = 1.5
            continue
        # skip centred paragraphs (title page, captions)
        if p.alignment == WD_ALIGN_PARAGRAPH.CENTER:
            continue
        # skip the caption paragraphs even if not centred (they carry SEQ fields)
        if FIG_RE.match(t) or TAB_RE.match(t):
            continue
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.5
        n_just += 1

    doc.save(OUT)
    print(f"Saved: {OUT}")
    print(f"  Figure captions tagged (SEQ Figure): {n_fig}")
    print(f"  Table  captions tagged (SEQ Table):  {n_tab}")
    print(f"  Section headings promoted -> Heading 1: {n_sec}")
    print(f"  Subsection headings promoted -> Heading 2: {n_sub}")
    print(f"  Front-matter fields rebuilt: {fixed}")
    print(f"  Chapter page-breaks set: {n_break}")
    print(f"  Body paragraphs justified + 1.5 spacing: {n_just}")


if __name__ == "__main__":
    main()
