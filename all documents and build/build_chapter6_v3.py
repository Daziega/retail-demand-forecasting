"""Generate Chapter 6 v3 — Discussion (aligned with supervisor's revised Ch1
Gap 4 framing: 'Insufficient integration of uncertainty in inventory
calibration', NOT 'Responsible AI'). Responsible-AI material kept as a short
note at the end of §6.3 rather than as the Gap 4 contribution."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_para(doc, text, *, italic=False, bold=False, font_size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(font_size)
    run.italic = italic
    run.bold = bold
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    return p


# ─────────────────────────────────────────────────────────────────────────
doc = Document()
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(11)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("CHAPTER 6")
run.bold = True; run.font.name = "Arial"; run.font.size = Pt(20)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("Discussion")
run.bold = True; run.font.name = "Arial"; run.font.size = Pt(16)
sub.paragraph_format.space_after = Pt(18)

# 6.1 Introduction
add_heading(doc, "6.1 Introduction", 1)
add_para(doc,
    "This chapter discusses the empirical findings of the thesis in the context of the "
    "research gaps identified in Chapter 2 and the broader academic literature. The "
    "discussion proceeds in seven parts. Section 6.2 synthesises the principal results "
    "from Chapters 4 and 5. Section 6.3 revisits each of the four research gaps from "
    "the literature review and evaluates the extent to which the thesis has closed "
    "them. Section 6.4 positions the findings against the comparable studies "
    "referenced in Chapter 2. Sections 6.5 and 6.6 discuss the theoretical and "
    "practical contributions of the work. Section 6.7 addresses validity threats. "
    "Section 6.8 expands on the limitations introduced in Chapters 4 and 5. Section "
    "6.9 identifies promising directions for future research. Section 6.10 summarises "
    "the chapter.")

# 6.2 Synthesis
add_heading(doc, "6.2 Synthesis of Empirical Findings", 1)
add_para(doc, "Three findings emerge from the empirical work and frame the discussion that follows.")
add_para(doc,
    "Finding 1: Forecast accuracy improvements at SKU-daily granularity are modest. "
    "The stacking ensemble of LightGBM, XGBoost, and Random Forest base learners "
    "combined through a non-negative Ridge meta-learner achieves the best point-"
    "forecast accuracy across all twelve models evaluated (MAE 0.952, WMAPE 73.58%). "
    "Improvements over the best classical baselines are 4.5% on MAE, 9.5% on RMSE, "
    "and approximately flat on MAPE. These figures sit below the 10–25% target band "
    "specified in Chapter 1 — an outcome consistent with the granularity-dependence "
    "documented in the M5 retrospective literature. The headline accuracy figures "
    "reported by Barghi (2025) and Seyedan et al. (2023) are obtained at weekly or "
    "category-level aggregation, where averaging across observations softens the "
    "relative weight of zero-actual days that distort per-row percentage metrics.")
add_para(doc,
    "Finding 2: Under direct simulation of the (R, s, S) policy, all three forecasting "
    "policies achieve full service on the test window. When the achieved service level "
    "of each policy is measured by simulating the policy forward day-by-day over the "
    "56-day test window — rather than inferred from a Gaussian projection onto the "
    "empirical demand distribution — all three policies (Classical, ML-Gaussian, and "
    "ML-Empirical-Quantile) achieve a realised service level of 100%. No demand goes "
    "unmet under any of the three policies, and the stockout-cost term is therefore "
    "zero in all cases. The same quantity that the earlier analytical formulation read "
    "as evidence of Gaussian under-allocation — namely, the empirical quantile spread "
    "of 1.36 times the Gaussian-implied buffer — is shown by simulation to be surplus "
    "inventory in a regime where the smaller Gaussian buffer already meets demand in "
    "full. The 1.36-times ratio is unchanged; only its interpretation changes.")
add_para(doc,
    "Finding 3: ML forecasting and the quantile-based safety-stock reformulation do "
    "not lower inventory cost in this regime. When service is saturated, the larger "
    "empirically calibrated buffer protects against stockouts that do not occur, and "
    "the additional holding cost it incurs is not recovered. The ML-Gaussian policy "
    "ties the Classical baseline to within $12 in total annual cost (a 0.0% change), "
    "and the ML-Empirical-Quantile policy costs approximately 1.5% more. The 21% "
    "reduction reported in the earlier draft of Chapter 5 was an artefact of the "
    "assumed service-level gap in that draft's cost formula and does not survive "
    "measurement. The Chapter 1 5–15% inventory-cost reduction target is therefore "
    "not achieved on this dataset, a result the chapter treats as a finding rather "
    "than a failure.")

# 6.3 Research Gaps
add_heading(doc, "6.3 Revisiting the Four Research Gaps", 1)
add_para(doc,
    "The literature review in Chapter 2 identified four research gaps that motivated "
    "the design of the thesis. This section returns to each gap and evaluates the "
    "extent to which the empirical work has closed it.")

# 6.3.1 Gap 1
add_heading(doc, "6.3.1 Gap 1 — Disconnect between forecast accuracy and inventory decision-making", 2)
add_para(doc,
    "Gap 1, articulated explicitly by Seyedan et al. (2023), identifies that most "
    "comparative forecasting studies report accuracy metrics without translating "
    "improved forecasts into operational inventory outcomes. This thesis addresses "
    "Gap 1 directly by integrating the forecasting framework with the OUTL policy and "
    "measuring inventory outcomes through forward simulation rather than inferring "
    "them analytically.")
add_para(doc,
    "The finding the thesis contributes to Gap 1 is conditional rather than universal: "
    "the value of ML forecasting and a quantile-based safety-stock reformulation for "
    "inventory cost reduction depends on the demand regime, not on forecast accuracy "
    "alone. Where the baseline (R, s, S) policy is exposed to genuine stockout risk, "
    "the larger empirically calibrated buffer earns its holding cost back by averting "
    "stockouts. Where service is already saturated — as on the M5 SKU-daily test "
    "window evaluated here — it does not. The decisive variable is therefore the "
    "frequency of stockouts under the baseline policy, which is itself governed by "
    "the demand regime and the generosity of the buffers, not by forecast accuracy.")
add_para(doc,
    "This thesis therefore reframes Gap 1 from 'how does ML forecasting translate "
    "into inventory savings' to 'under what conditions does ML forecasting translate "
    "into inventory savings'. The contribution is a rigorous, simulation-based "
    "boundary condition that practitioners and researchers can use to test their own "
    "settings, rather than a single quoted percentage saving.")

# 6.3.2 Gap 2
add_heading(doc, "6.3.2 Gap 2 — Real-world deployment and operational validation", 2)
add_para(doc,
    "Gap 2 concerns the limited focus in the literature on real-world deployment, "
    "model maintenance, and the practical constraints of live operation. The "
    "literature is dominated by historical-simulation evaluations that do not address "
    "the deployment layer.")
add_para(doc,
    "The thesis addresses Gap 2 through three concrete deliverables. First, the end-"
    "to-end pipeline produces a set of CSV data feeds designed for daily refresh in a "
    "production deployment — five files exported from the Python pipeline corresponding "
    "to the four dashboard views and a KPI summary. Second, the Power BI dashboard "
    "provides an interactive, accessible interface for non-technical retail managers, "
    "organising the framework's outputs into four operational views (Forecast Overview, "
    "Inventory Status, Cost Analysis, Sensitivity Analysis) and exposing the underlying "
    "parameters as configurable filters. Third, the inventory analysis is itself a "
    "simulation-based deployment validation: the (R, s, S) policy is run forward over "
    "real test demand day-by-day, providing the kind of operational measurement that "
    "the literature's analytical projections typically substitute for.")
add_para(doc,
    "Live deployment and incremental model updating remain out of scope, but the "
    "framework is structurally deployable and the simulation engine is reusable as a "
    "pre-deployment diagnostic tool for testing whether an SME's specific demand "
    "regime falls into the stockout-risk regime where ML forecasting pays off "
    "operationally.")

# 6.3.3 Gap 3
add_heading(doc, "6.3.3 Gap 3 — Underrepresentation of SME retailers and practical constraints", 2)
add_para(doc,
    "Gap 3 concerns the dominance in the literature of frameworks designed for "
    "enterprise-scale analytics infrastructure, with limited attention to whether the "
    "same methodologies are viable for small and medium-sized retailers operating with "
    "constrained resources. Wahedi et al. (2023) and Ungureanu (2025) identify the SME "
    "gap as a substantive research need. The thesis addresses this gap through the "
    "explicit SME transferability commitment of Chapter 1 and the corresponding "
    "specification in §3.11.")
add_para(doc,
    "Three specific design choices implement this commitment. First, the entire pipeline "
    "is built on free and open-source tooling — Python with mainstream libraries "
    "(pandas, scikit-learn, lightgbm, xgboost, tensorflow, statsforecast, shap) and "
    "Microsoft Power BI Service free tier. No commercial software licence or cloud-"
    "computing subscription is required. Second, the framework runs in under three "
    "hours on a standard 16 GB-RAM laptop with no GPU. Third, the data inputs required "
    "are standard outputs of off-the-shelf SME point-of-sale systems, ensuring that no "
    "additional data infrastructure investment is needed.")
add_para(doc,
    "An important nuance emerges from the inventory analysis: SME accessibility is not "
    "only a tooling question but also a regime-fit question. The simulation engine "
    "delivered with the thesis allows an SME to test, on its own data, whether its "
    "demand regime exhibits genuine stockout risk under the baseline policy and "
    "therefore whether investment in the more complex ML-quantile pipeline is "
    "worthwhile. This regime-fit diagnostic is itself part of what the thesis "
    "contributes to Gap 3.")

# 6.3.4 Gap 4 — REVISED to match supervisor's Ch1 framing
add_heading(doc, "6.3.4 Gap 4 — Integration of uncertainty in inventory calibration", 2)
add_para(doc,
    "Gap 4, as articulated in the supervisor's revised Chapter 1, concerns the "
    "insufficient integration of forecast uncertainty into the inventory-calibration "
    "step. Classical safety-stock formulas assume point-forecast inputs and treat the "
    "mean forecast as certain, systematically underrepresenting the distributional "
    "structure of demand variability. The thesis addresses this gap in two "
    "complementary ways.")
add_para(doc,
    "First, the OUTL safety-stock calculation in §3.7.2 uses the standard deviation "
    "of forecast residuals derived from the validation set as the σ input, rather "
    "than the standard deviation of historical demand. This couples the inventory "
    "calculation directly to the forecasting model's measured prediction error, "
    "ensuring that improvements in forecast accuracy propagate into tighter safety-"
    "stock allocation through a measured, model-specific σ rather than a generic "
    "historical-demand variance.")
add_para(doc,
    "Second, the §3.7.4 robustness check replaces the parametric Gaussian formula "
    "z·σ·√(R+L) with an empirically calibrated quantile spread derived from the "
    "Tier 5 LightGBM quantile-regression model. This eliminates the assumption of "
    "normally distributed residuals — an assumption that intermittent retail demand "
    "systematically violates — and substitutes a distribution-free buffer that reads "
    "the actual shape of the empirical demand quantiles. The simulation analysis of "
    "Chapter 5 shows that this substitution does not lower inventory cost in the "
    "saturated-service regime of the M5 test window; the buffer it specifies is "
    "larger than the regime requires, and the additional holding cost is not recovered "
    "because there are no stockouts to avert. The contribution to Gap 4 is therefore "
    "not the operational saving claimed in the earlier draft but rather the boundary "
    "condition under which uncertainty integration in inventory calibration yields "
    "measurable benefit — namely, when the demand distribution's right tail is heavy "
    "enough, or the baseline buffer thin enough, that the Gaussian formula leaves "
    "demand uncovered. On the present evidence, that condition does not obtain on "
    "M5 SKU-daily data; it would be expected to obtain in regimes of high demand "
    "intermittency combined with tight baseline buffers.")

# Responsible-AI commitments — kept as a SEPARATE design note, not labelled as Gap 4
add_para(doc, "")
add_para(doc, "Responsible-AI commitments (separate to the four research gaps).", bold=True, font_size=11)
add_para(doc,
    "In addition to the four gaps the thesis is designed to address, it makes three "
    "responsible-AI commitments that are not labelled as research gaps but are "
    "discussed in §3.10 and reported in §4.11 and §4.12. Model transparency is "
    "established through SHAP analysis of the best-performing model (the top five "
    "features account for 78.6% of model behaviour). Performance fairness is "
    "established through disaggregated error metrics by product category and by state, "
    "with the observation that performance variation reflects underlying demand "
    "structure rather than systematic algorithmic bias. Data minimisation is "
    "implemented by retaining only features with non-trivial SHAP contribution. These "
    "commitments support the responsible deployment of the framework in SME "
    "environments and complement, but do not substitute for, the four gap-closures "
    "discussed above.")

# 6.4 Literature Comparison
add_heading(doc, "6.4 Comparison with the Literature", 1)
add_para(doc,
    "This section positions the thesis findings against the principal comparable studies "
    "reviewed in Chapter 2 along three axes: forecast accuracy, modelling architecture, "
    "and inventory cost impact.")
add_para(doc,
    "Forecast accuracy. Barghi (2025) reports MAPE of 6.48% (hybrid ensemble) reduced "
    "to 3.53% with fuzzy generalisation preprocessing on the Rossmann Store Sales "
    "dataset at weekly aggregation. Seyedan et al. (2023) report MAPE of 5.22% (sports) "
    "and 9.58% (electronics) with their heterogeneous MLP-LSTM-CNN stacking ensemble "
    "on weekly retail data. Nasseri et al. (2023) report MAPE in a similar range with "
    "Extra Tree Regressors on perishable retail categories. The MAPE figures in this "
    "thesis — 53–58% across the model tiers at SKU-daily granularity — are an order of "
    "magnitude higher in absolute terms but consistent with the M5 retrospective "
    "literature at the same granularity. Single-digit MAPE on M5-style data is "
    "achievable only at weekly or category-level aggregation; the SKU-daily granularity "
    "used in this thesis is methodologically more demanding and operationally more "
    "informative, because inventory decisions are made at the SKU-store-day level.")
add_para(doc,
    "Modelling architecture. The thesis adopts a stacking ensemble of LightGBM, XGBoost, "
    "and Random Forest with a non-negative Ridge meta-learner — the conventional "
    "heterogeneous ensemble architecture used by Seyedan et al. (2023) and Mitra et al. "
    "(2022). This was preferred over the sequential residual-correction design used in "
    "Barghi (2025) on the grounds that sequential chaining tends to overfit successive "
    "layers to earlier-layer noise and stacking with non-negative weights is the more "
    "standard contemporary architecture. The empirical results support this choice in "
    "the sense that the stacking ensemble achieves the best MAE and WMAPE among all "
    "twelve models evaluated.")
add_para(doc,
    "Inventory cost impact. Seyedan et al. (2023) and Barghi (2025) report inventory "
    "cost reductions in the 8–14% range against single-model baselines, on weekly "
    "aggregated retail data. The thesis does not reproduce this finding on the M5 SKU-"
    "daily window when inventory outcomes are measured by direct simulation: all "
    "policies achieve full service, no stockouts occur, and the ML-Empirical-Quantile "
    "policy in fact incurs approximately 1.5% more total annual cost than the Classical "
    "baseline because it over-provisions safety stock without preventing any additional "
    "stockouts. This is a substantive divergence from the prior literature, and it is "
    "the substantive contribution of the thesis. The divergence is most plausibly "
    "attributable to two factors. First, the prior studies report cost reductions on "
    "aggregated data where demand smoothing produces variance levels that the baseline "
    "policy cannot fully absorb; on SKU-daily intermittent demand, the baseline (R, s, "
    "S) buffers are generous enough that this no longer holds. Second, the prior "
    "studies infer rather than measure the stockout-rate gap between policies; the "
    "present analysis measures it, and the gap is zero on the test window.")

# 6.5 Theoretical
add_heading(doc, "6.5 Theoretical Contributions", 1)
add_para(doc, "The thesis makes three theoretical contributions.")
add_para(doc,
    "First, it provides the first simulation-based measurement on the M5 benchmark of "
    "the operational consequence of the Gaussian-versus-empirical safety-stock choice. "
    "The empirical quantile spread is 1.36 times the Gaussian buffer, but the simulation "
    "shows that this distributional mismatch is regime-dependent in its operational "
    "consequence: the under-allocation is real as a property of the demand distribution "
    "yet operationally immaterial when the baseline buffer already meets demand in "
    "full. This finding refines the textbook observation (Silver, Pyke & Thomas, 2017) "
    "that the Gaussian formula assumes symmetric normally distributed demand by "
    "supplying the empirical condition under which the assumption fails operationally.")
add_para(doc,
    "Second, the thesis demonstrates that the operational value of ML forecasting in "
    "retail inventory management is conditional on the demand regime rather than "
    "automatic. The ML-Gaussian policy, using the most accurate point forecasts "
    "produced in Chapter 4, ties the Classical baseline to within 0.0% in total annual "
    "cost, and the ML-Empirical-Quantile policy costs 1.5% more. This contradicts the "
    "common framing in the forecasting literature in which improved point-forecast "
    "accuracy is presumed to translate directly into operational inventory savings. "
    "The boundary condition under which the translation does and does not hold is "
    "itself the theoretical contribution.")
add_para(doc,
    "Third, the thesis shows that within the regime studied, the operational comparison "
    "is robust to MAPE level and to the stockout-cost multiplier. Across the 3–15% MAPE "
    "sensitivity range, the simulated cost comparison varies by less than 0.1 "
    "percentage points; across the 0.4× to 2.0× stockout cost multiplier range, the "
    "comparison is unchanged because stockouts are zero in all three policies and the "
    "multiplier never engages. This robustness profile clarifies that further "
    "investment in MAPE improvement is unlikely to yield inventory benefit in regimes "
    "of saturated service.")

# 6.6 Practical
add_heading(doc, "6.6 Practical Contributions", 1)
add_para(doc, "Three practical contributions complement the theoretical findings.")
add_para(doc,
    "An end-to-end, reproducible, SME-deployable framework. The thesis delivers a "
    "complete workflow from raw M5 ingestion through to the Power BI dashboard, "
    "implemented in free and open-source tooling and runnable on a standard laptop. "
    "The framework is documented to the level of detail required for independent "
    "replication, including the stratified subsample construction (§3.3.3), the "
    "chronological train/validation/test split (§3.6.2), the feature engineering "
    "specification (§3.5), the model configurations (§3.6.1), the OUTL calculations "
    "(§3.7), and the forward-simulation engine (§5.2).")
add_para(doc,
    "A simulation engine as a pre-investment diagnostic. The forward (R, s, S) "
    "simulation engine accompanying the thesis is itself a deliverable for SME "
    "practitioners. Before committing to the more complex ML-quantile pipeline, an SME "
    "can run the simulation engine on its own demand data to test whether its policy "
    "is exposed to genuine stockout risk under classical inputs. If the simulation "
    "shows that the baseline (R, s, S) policy already meets demand in full — as on the "
    "M5 test window — the additional investment in ML forecasting and quantile-based "
    "safety stock is unlikely to recover its cost; the resource is better directed "
    "elsewhere. This diagnostic capability is more operationally valuable than the "
    "headline cost-reduction figures typically reported in comparable studies.")
add_para(doc,
    "A decision-support dashboard. The Power BI dashboard translates the framework's "
    "outputs into four operational views (Forecast Overview, Inventory Status, Cost "
    "Analysis, Sensitivity Analysis) accessible to non-technical SME managers. The "
    "dashboard supports practitioner what-if analysis — changing the service-level "
    "target, the assumed lead time, or the stockout cost multiplier and observing the "
    "impact on the comparison — without requiring the user to understand the "
    "underlying methodology.")

# 6.7 Validity
add_heading(doc, "6.7 Validity Threats", 1)
add_para(doc,
    "Three classes of validity threat are considered: internal validity, external "
    "validity, and construct validity.")
add_para(doc,
    "Internal validity concerns whether the observed outcomes are attributable to the "
    "framework rather than to confounding factors in the implementation. The principal "
    "internal-validity threat is the chronological split between training, validation, "
    "and test sets. The methodology in §3.6.2 specifies strict chronological "
    "partitioning to eliminate data leakage from future to past. A secondary threat is "
    "the target encoding of high-cardinality categoricals; §3.4.6 specifies that "
    "target encodings are computed exclusively on the training window and propagated "
    "unchanged to validation and test, eliminating this leakage channel. A third "
    "threat — specific to the simulation engine — was raised by the §5.3 falsification "
    "test, which establishes that the engine can register stockouts (and the size of "
    "the stockout cost it would produce in non-saturated regimes), removing the "
    "concern that the 100% realised service level reflects a counter that fails to "
    "fire rather than genuinely-met demand.")
add_para(doc,
    "External validity concerns the generalisability of the findings to other retail "
    "contexts. The M5 dataset originates from a single large enterprise retailer "
    "(Walmart) operating primarily in three US states; the SME accessibility argument "
    "therefore rests on the transferability of the methodology, not on direct empirical "
    "validation in an SME environment. The boundary condition identified in §5.6 — "
    "that ML forecasting and quantile safety stock deliver operational value only when "
    "the baseline policy is exposed to genuine stockout risk — is likely to be the "
    "principal mediator of generalisability. Sensitivity analysis with the simulation "
    "engine (§5.7 and the §5.3 falsification results) shows that stockouts reappear "
    "under demand shocks and depleted buffers, and the cost ranking of the policies "
    "may flip in those regimes. The §3.11 SME transferability specification is "
    "therefore most plausibly satisfied for SMEs whose baseline policy is exposed to "
    "stockout risk; the regime-fit diagnostic of §6.6 should be the first step in any "
    "transfer.")
add_para(doc,
    "Construct validity concerns whether the metrics used measure what they are "
    "intended to measure. The Chapter 1 MAPE improvement target is the principal "
    "construct-validity concern at the forecasting level: at SKU-daily granularity on "
    "intermittent demand, MAPE is known to be an unstable metric, and the supplementary "
    "WMAPE used in this thesis is more robust for the data structure. The cost-"
    "reduction target is the principal construct-validity concern at the inventory "
    "level: the earlier draft's analytical projection of an under-allocation gap onto "
    "the empirical demand distribution conflated a property of the demand distribution "
    "with an operational outcome. The present analysis measures the operational outcome "
    "directly by simulation, eliminating this construct ambiguity.")

# 6.8 Extended Limitations
add_heading(doc, "6.8 Extended Limitations", 1)
add_para(doc,
    "Four limitations specific to the analysis in Chapters 4 and 5 were noted in §4.13 "
    "and §5.10. This section adds three further limitations spanning the thesis as a "
    "whole.")
add_para(doc,
    "First, the simulation analysis uses a single 56-day test window (March–April "
    "2016). A longer or differently situated window, or one containing a demand spike, "
    "could expose the policies to stockouts and could materially affect the cost "
    "ranking. The §5.3 falsification runs probe the alternative regime through stress "
    "conditions, but as illustrative checks rather than a systematic sweep across "
    "demand-intensity scenarios.")
add_para(doc,
    "Second, the LSTM model in Tier 3 underperforms the tree-based models. While this "
    "finding is consistent with the M5 retrospective literature (Nasseri et al., 2023), "
    "it leaves open the possibility that a more elaborate sequence architecture "
    "(Temporal Fusion Transformers, N-BEATS) or per-series LSTM specialisation would "
    "close the gap. The thesis evaluates only the methodology-specified LSTM architecture.")
add_para(doc,
    "Third, the stockout cost model is linear in unmet units, with the multiplier "
    "evaluated as a three-point sensitivity. In practice, stockout costs may be non-"
    "linear (for example, customer churn after repeated stockouts) and the multiplier "
    "may interact with longer-run market dynamics that a single-window simulation does "
    "not capture. The irrelevance of the multiplier in the regime studied makes the "
    "point moot here, but it would matter in a higher-stockout regime.")

# 6.9 Future Research
add_heading(doc, "6.9 Implications for Future Research", 1)
add_para(doc,
    "The findings of this thesis suggest five productive directions for future research.")
add_para(doc,
    "First, the simulation-based comparison should be repeated on retail benchmarks "
    "with more variable demand (Rossmann, Favorita, regional grocery datasets) and on "
    "alternative time windows where seasonal peaks expose the baseline policy to "
    "stockouts. The boundary condition identified in §5.6 predicts that the cost "
    "ranking of the policies will shift in those regimes; testing this prediction is "
    "the most directly actionable extension.")
add_para(doc,
    "Second, the quantile-based safety stock formulation should be evaluated against "
    "more sophisticated probabilistic forecasting methods (DeepAR, Bayesian neural "
    "networks, conformal prediction intervals) to assess whether quantile regression "
    "is the right probabilistic-forecasting model class for safety-stock calibration, "
    "or whether alternative specifications yield further operational improvement once "
    "the regime is one in which any quantile method would pay off.")
add_para(doc,
    "Third, the framework should be extended to incorporate live deployment routines: "
    "rolling retraining schedules, model drift detection, and the operational "
    "monitoring of forecast calibration over time. This extension would more fully "
    "close Gap 2 from Chapter 2.")
add_para(doc,
    "Fourth, the simulation engine should be extended to support non-linear stockout "
    "cost functions and supplier lead-time uncertainty. Both refinements would produce "
    "a more realistic operational cost model, although at the expense of analytical "
    "tractability.")
add_para(doc,
    "Fifth, the framework's transferability to SME environments should be validated "
    "through deployment in actual SME retail operations. The §3.11 transferability "
    "specification documents the requirements for such deployment; an empirical case "
    "study in an SME context where the baseline policy is exposed to stockouts would "
    "test both the SME transferability claim and the boundary condition identified in "
    "Chapter 5.")

# 6.10 Summary
add_heading(doc, "6.10 Chapter Summary", 1)
add_para(doc,
    "This chapter has discussed the empirical findings of Chapters 4 and 5 in the "
    "context of the four research gaps identified in Chapter 2 and the broader academic "
    "literature. Three principal findings were synthesised: forecast accuracy "
    "improvements at SKU-daily granularity are modest (4.5% MAE, 9.5% RMSE) and below "
    "the Chapter 1 target band; all three forecasting policies achieve full service on "
    "the test window under direct simulation; and ML forecasting and quantile-based "
    "safety-stock reformulation do not lower inventory cost in this regime, with the "
    "ML-Empirical-Quantile policy in fact costing 1.5% more than the Classical baseline.")
add_para(doc,
    "The thesis reframes Gap 1 (forecast-inventory disconnect) from a quantified cost "
    "saving to a quantified boundary condition: ML forecasting and distributional "
    "safety stock affect inventory outcomes when, and only when, the baseline policy "
    "is exposed to genuine stockout risk. It partially closes Gap 2 (deployment) "
    "through the dashboard and CSV refresh architecture, addresses Gap 3 (SME "
    "underrepresentation) through tooling and infrastructure choices and by delivering "
    "the simulation engine as a pre-investment diagnostic, and addresses Gap 4 "
    "(uncertainty integration in inventory calibration) through residual-σ coupling of "
    "forecast to inventory and through the LightGBM quantile model and empirical "
    "safety-stock buffer — with the operational consequence of the latter shown to be "
    "regime-dependent. Responsible-AI commitments (model transparency through SHAP, "
    "performance fairness, data minimisation) are honoured separately from the four "
    "gap closures. The literature comparison establishes that the headline cost-"
    "reduction figures reported by Barghi (2025) and Seyedan et al. (2023) on "
    "aggregated data do not transfer to SKU-daily simulation, and the divergence is "
    "itself the substantive contribution of the thesis.")
add_para(doc,
    "Three theoretical contributions are identified: the first simulation-based "
    "characterisation of the operational consequence of the Gaussian-versus-empirical "
    "safety-stock choice; the demonstration that the operational value of ML "
    "forecasting in retail inventory management is conditional on the demand regime "
    "rather than automatic; and the empirical robustness of the comparison to MAPE "
    "level and stockout-cost multiplier within the saturated-service regime. Three "
    "practical contributions complement these: an end-to-end deployable framework; a "
    "simulation engine as a pre-investment diagnostic for SME practitioners; and a "
    "decision-support dashboard. Validity threats and limitations are acknowledged, "
    "and five directions for future research are identified.")
add_para(doc,
    "Chapter 7 concludes the thesis. It returns to the research objectives stated in "
    "Chapter 1, summarises the extent to which each has been achieved, and consolidates "
    "the contribution to the demand forecasting and inventory optimisation literature.")

OUT = "/Users/desmond/Capstone Project/retail-demand-forecasting/Chapter_6_Discussion_v3.docx"
doc.save(OUT)
print(f"Saved: {OUT}")
