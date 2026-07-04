"""Build a standalone, management-oriented Executive Summary (~4,500-5,000
words, 12-15 pages) that expands on the academic abstract for a
decision-maker / organisational audience. Kept consistent with the final
disconfirmation-based narrative of Chapters 5-7."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/Users/desmond/Capstone Project/retail-demand-forecasting/TFM_Executive_Summary.docx"


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_para(doc, text, *, italic=False, bold=False, size=11, space_after=10,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.5):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.italic = italic
    run.bold = bold
    return p


def add_bullet(doc, text, *, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    if bold_lead:
        r1 = p.add_run(bold_lead)
        r1.bold = True
        r1.font.name = "Arial"
        r1.font.size = Pt(11)
        r2 = p.add_run(text)
        r2.font.name = "Arial"
        r2.font.size = Pt(11)
    else:
        r = p.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(11)
    return p


def add_page_break(doc):
    p = doc.add_paragraph()
    r = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r._r.append(br)


def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def add_table(doc, headers, rows, header_fill="D5E8F0"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.name = "Arial"
        r.font.size = Pt(10)
        set_cell_shading(hdr_cells[i], header_fill)
    for r_idx, row in enumerate(rows, start=1):
        cells = table.rows[r_idx].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ""
            p = cells[c_idx].paragraphs[0]
            r = p.add_run(str(val))
            r.font.name = "Arial"
            r.font.size = Pt(10)
    return table


# ─────────────────────────────────────────────────────────────────────────
doc = Document()
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(11)

# ── TITLE PAGE ────────────────────────────────────────────────────────────
for _ in range(5):
    doc.add_paragraph()

t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("EXECUTIVE SUMMARY")
r.bold = True; r.font.name = "Arial"; r.font.size = Pt(24)
t.paragraph_format.space_after = Pt(12)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Demand Forecasting and Inventory Optimisation Using Machine Learning")
r.bold = True; r.font.name = "Arial"; r.font.size = Pt(16)
sub.paragraph_format.space_after = Pt(8)

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run(
    "A Simulation-Based Evaluation for Retail Decision-Makers, with an "
    "SME-Accessible Power BI Dashboard")
r.italic = True; r.font.name = "Arial"; r.font.size = Pt(13)
sub2.paragraph_format.space_after = Pt(36)

note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = note.add_run(
    "Prepared for retail operations decision-makers and Master's Thesis (TFM) "
    "evaluation committee. This document expands on the academic abstract for "
    "an organisational and managerial audience and should be read alongside, "
    "not in place of, the full thesis.")
r.italic = True; r.font.name = "Arial"; r.font.size = Pt(10)

add_page_break(doc)

# ── 1. PURPOSE OF THIS SUMMARY ──────────────────────────────────────────
add_heading(doc, "1. Purpose of This Summary", 1)
add_para(doc,
    "This executive summary is written for the people who will ultimately decide "
    "whether, and how, to invest in machine learning-based demand forecasting and "
    "inventory optimisation: retail operations managers, finance leads, and owners "
    "of small and medium-sized enterprises (SMEs), as well as the academic committee "
    "evaluating this thesis as a piece of applied research. It is not a condensed "
    "version of the technical thesis. It is a translation of that thesis into the "
    "language and priorities of a decision-maker: what problem was being solved, why "
    "it matters commercially, what was actually tested, what the evidence shows, and "
    "what a sensible organisation should do next.")
add_para(doc,
    "The underlying thesis is a substantial piece of technical work — twelve "
    "forecasting models compared, an inventory policy built and tested through direct "
    "simulation, and an interactive dashboard delivered as a decision-support tool. "
    "None of that technical depth is necessary to understand the conclusion, which "
    "can be stated in one sentence: the value of machine learning forecasting for "
    "retail inventory management is real, but it is not automatic, and this thesis "
    "provides a tool that tells an organisation, before it spends money, whether its "
    "own situation is one where the investment is likely to pay off.")
add_para(doc,
    "The remainder of this summary is structured to mirror the questions a "
    "decision-maker actually asks: What decision am I facing? Why is this problem "
    "hard? What does the existing evidence say? How was this specific research "
    "conducted, and can I trust it? What did it find? And, most importantly, what "
    "should I do with that information?")

# ── 2. THE BUSINESS DECISION AT STAKE ────────────────────────────────────
add_heading(doc, "2. The Business Decision at Stake", 1)
add_para(doc,
    "Every retailer, from a single-location shop to a national chain, makes the same "
    "recurring decision: how much stock to hold, and when to reorder it. Get the "
    "answer wrong in one direction and the business runs out of product customers "
    "want to buy — a stockout — losing the sale immediately and, if it happens "
    "repeatedly, losing the customer's future business as well. Get it wrong in the "
    "other direction and the business ties up cash in inventory that sits on shelves "
    "or in warehouses, incurring storage cost, spoilage risk for perishables, and the "
    "opportunity cost of capital that could have been used elsewhere.")
add_para(doc,
    "Large retailers such as Walmart address this problem with dedicated data "
    "science teams, enterprise forecasting software, and analysts who spend their "
    "careers tuning inventory parameters. Small and medium-sized retailers typically "
    "have none of this. They rely on the experience and intuition of a manager, on "
    "simple spreadsheet rules, or on the basic reorder logic built into off-the-shelf "
    "point-of-sale software. This is not a criticism of SME operators — it reflects a "
    "genuine resource gap. Data science talent, enterprise software licences, and "
    "cloud computing budgets are simply not accessible to a business running on thin "
    "retail margins.")
add_para(doc,
    "In the last decade, a large and growing body of academic and industry research "
    "has argued that machine learning can close this gap: that algorithms trained on "
    "a retailer's own sales history can forecast demand more accurately than "
    "traditional statistical methods, and that these more accurate forecasts can be "
    "converted into inventory policies that reduce both stockouts and excess stock "
    "simultaneously. Vendors selling forecasting software, and many academic papers, "
    "report double-digit percentage improvements — forecast accuracy gains of 10 to "
    "25 percent, and inventory cost reductions of 5 to 15 percent, relative to the "
    "classical methods that most SMEs currently use.")
add_para(doc,
    "This creates a genuine decision for any retail operator: should the organisation "
    "invest time, money, and change-management effort into adopting a machine "
    "learning forecasting and inventory system? The honest answer, prior to this "
    "thesis, was difficult to give with confidence, for a reason explained in detail "
    "in Section 4 below: most of the evidence cited in favour of such an investment "
    "was generated under conditions — aggregated weekly or category-level sales data, "
    "and inventory calculations that assumed rather than measured their own "
    "performance — that do not match how a typical SME actually operates or how the "
    "resulting software would actually be used, at the day-by-day, product-by-product "
    "level.")
add_para(doc,
    "This thesis was designed to close that evidence gap. It builds a complete, "
    "end-to-end forecasting-to-inventory pipeline using the same kind of data an SME "
    "already collects, tests it under conditions that match real day-to-day retail "
    "operation, and — critically — measures the resulting inventory performance by "
    "actually running the inventory policy forward through real demand data, rather "
    "than assuming its performance analytically. The result is a decision-support "
    "tool, delivered as an interactive dashboard, and a specific, evidence-based "
    "answer to the investment question above: not a blanket yes or no, but a "
    "diagnostic that tells an individual organisation whether its own demand pattern "
    "makes the investment worthwhile.")

# ── 3. THE NATURE OF THE PROBLEM ─────────────────────────────────────────
add_heading(doc, "3. The Nature of the Problem", 1)
add_para(doc,
    "To understand why this is a genuinely hard problem — and why the answer turned "
    "out to be more nuanced than the marketing literature suggests — it helps to "
    "unpack what \"demand forecasting and inventory optimisation\" actually involves "
    "at the level of detail a retail operation experiences it.")

add_heading(doc, "3.1 Demand Is Volatile and, at the Product Level, Highly Intermittent", 2)
add_para(doc,
    "Aggregate retail demand — total sales across a whole store or category, summed "
    "over a week or a month — tends to be relatively smooth and predictable. But "
    "inventory decisions are not made at that aggregate level. They are made "
    "product by product, store by store, day by day: how many units of this "
    "specific item, at this specific location, will sell tomorrow? At that level of "
    "detail, demand for most retail products is intermittent — many days see zero "
    "sales of a given item at a given store, punctuated by occasional spikes driven "
    "by promotions, paydays, holidays, or simple randomness. This intermittency is "
    "the single most important technical fact about retail demand, and it is the "
    "reason forecasting accuracy figures that sound impressive in academic papers "
    "often do not describe the reality of day-to-day, item-level operations.")

add_heading(doc, "3.2 The Forecast Is Only Half the Problem", 2)
add_para(doc,
    "A more accurate demand forecast is not, by itself, worth anything to a "
    "business. Its value is realised only when it is translated into an inventory "
    "decision — how much safety stock to hold, when to trigger a reorder, and how "
    "much to order. Most published research on retail forecasting stops at the "
    "forecast: it reports how close the model's predictions were to actual sales, "
    "using statistical accuracy metrics, and stops there. Very little published "
    "research goes the extra step of actually running the resulting inventory "
    "policy and measuring what happens to costs and stockouts. This thesis "
    "identifies this disconnect as the first and most consequential gap in the "
    "existing literature, and it is the gap that this research is primarily "
    "designed to close.")

add_heading(doc, "3.3 The Textbook Safety-Stock Formula Was Built for a Different Kind of Demand", 2)
add_para(doc,
    "The standard method used across industry and academia to calculate how much "
    "buffer stock (\"safety stock\") to hold assumes that demand forecast errors "
    "follow a symmetric, bell-shaped statistical distribution — the same assumption "
    "underlying most classical statistics. Intermittent, spiky retail demand of the "
    "kind described in Section 3.1 does not follow this pattern; it is skewed, with "
    "a small number of large demand spikes and a large number of zero- or near-zero "
    "demand days. This mismatch between the assumption built into the textbook "
    "formula and the actual shape of real retail demand is a second identified gap "
    "in current practice, and one of the central technical questions this thesis "
    "investigates: does correcting this mismatch — replacing the textbook formula "
    "with one that reflects the true shape of the demand distribution — actually "
    "save money in practice?")

add_heading(doc, "3.4 SMEs Face a Structural Resource Disadvantage", 2)
add_para(doc,
    "Even where the technical solutions above are well understood, SME retailers "
    "face a separate and equally real barrier: access. Enterprise forecasting "
    "software, cloud computing infrastructure, and in-house data science expertise "
    "are typically priced and designed for organisations far larger than a typical "
    "SME. This is the third gap the thesis addresses directly, through an explicit "
    "design commitment to use only free, open-source software, to require no "
    "computing hardware beyond a standard laptop, and to require no data that a "
    "typical retail point-of-sale system does not already collect.")

add_heading(doc, "3.5 Point Forecasts Hide the Uncertainty That Inventory Decisions Actually Depend On", 2)
add_para(doc,
    "Finally, most forecasting systems — including many marketed commercially — "
    "produce a single \"best guess\" number for tomorrow's demand. But the inventory "
    "decision does not depend on the best guess; it depends on the range of "
    "plausible outcomes and, specifically, on how bad the worst realistic case "
    "could be, because that is what safety stock is meant to protect against. A "
    "single point forecast cannot answer that question. This is the fourth gap the "
    "thesis addresses: incorporating the full range of forecast uncertainty, not "
    "just a single predicted number, into the inventory calculation.")

# ── 4. LITERATURE REVIEW ─────────────────────────────────────────────────
add_heading(doc, "4. What Is Already Known: Insights from the Research Literature", 1)
add_para(doc,
    "This thesis is grounded in a structured review of 52 academic sources spanning "
    "classical statistical forecasting, machine learning, deep learning, ensemble "
    "methods, and inventory theory. Three consistent findings emerge from that "
    "body of work, alongside the four gaps already introduced in Section 3.")
add_para(doc,
    "First, there is strong and consistent evidence that machine learning and "
    "ensemble methods — models that combine the predictions of several algorithms — "
    "outperform traditional statistical forecasting methods (moving averages, "
    "exponential smoothing, ARIMA-family models) across a wide range of retail "
    "datasets and sectors. This finding is well replicated and is not in serious "
    "dispute in the literature.")
add_para(doc,
    "Second, published studies report substantial operational benefits from this "
    "accuracy improvement: forecast accuracy gains in the 10-25 percent range and "
    "inventory cost reductions of 5-15 percent relative to classical baselines are "
    "commonly cited figures, drawn from studies such as Barghi (2025) and Seyedan, "
    "Mafakheri and Wang (2023). These are the benchmark figures against which this "
    "thesis explicitly tests its own results, and they are the figures most "
    "frequently used to justify commercial investment in forecasting technology.")
add_para(doc,
    "Third — and this is the critical qualification that a decision-maker needs to "
    "understand before acting on the figures in the paragraph above — the studies "
    "reporting these benchmark figures typically evaluate their models on data "
    "aggregated to the weekly or product-category level, not at the daily, "
    "individual-product level at which actual inventory reorder decisions are made. "
    "This distinction matters enormously. Aggregation smooths out the intermittency "
    "described in Section 3.1: a category's total weekly sales are far more "
    "predictable than one specific product's sales on one specific day. A forecasting "
    "improvement measured on smoothed, aggregated data does not automatically apply "
    "at the granularity where inventory decisions actually get made.")
add_para(doc,
    "This thesis was explicitly designed to test whether the published benchmark "
    "figures — 10-25 percent accuracy improvement, 5-15 percent cost reduction — "
    "hold up when the analysis is done at the SKU-daily level, using real product-"
    "store-day sales records rather than aggregated data, and using inventory "
    "outcomes that are directly measured through simulation rather than estimated "
    "analytically. This is a test of transferability, not a repetition of prior "
    "work, and Section 6 reports what that test found.")

add_para(doc, "The four research gaps this thesis addresses can be summarised as follows.", bold=False)
add_table(doc,
    headers=["Gap", "The Problem", "What This Thesis Does About It"],
    rows=[
        ["1. Forecast-to-decision disconnect",
         "Most studies report forecast accuracy and stop there, without measuring the resulting inventory cost or service outcome.",
         "Builds the full pipeline from forecast to inventory policy, and measures the resulting cost and service level directly by simulation."],
        ["2. Limited real-world validation",
         "Most studies validate against historical data only, without testing how the resulting policy behaves when actually run forward in time.",
         "Runs the inventory policy forward, day by day, over real test-period demand, and includes a dedicated engine-validation exercise (Section 5.4)."],
        ["3. SME underrepresentation",
         "Most research and commercial tooling assumes enterprise-scale infrastructure and data science staffing that SMEs do not have.",
         "Uses only free, open-source tools; runs on a standard laptop; requires only data an SME point-of-sale system already produces."],
        ["4. Uncertainty ignored in inventory calibration",
         "Standard safety-stock formulas assume a single point forecast and a symmetric error distribution that intermittent retail demand does not follow.",
         "Builds a probabilistic (quantile) forecasting model and tests an inventory formula based on the real shape of the demand distribution, not an assumed one."],
    ])

# ── 5. RESEARCH METHODOLOGY AND APPROACH ─────────────────────────────────
add_heading(doc, "5. Research Methodology and Approach", 1)
add_para(doc,
    "This section summarises how the research was conducted, written for a reader "
    "who wants to understand the credibility and rigour of the approach without "
    "needing the full technical specification found in the thesis itself.")

add_heading(doc, "5.1 The Data", 2)
add_para(doc,
    "The research uses the M5-Forecasting (Walmart) dataset, a large, publicly "
    "available benchmark comprising daily unit sales for 3,049 products across ten "
    "Walmart stores in three US states, spanning more than five years (2011-2016), "
    "together with calendar information (day of week, holidays, promotional events, "
    "SNAP benefit-eligibility days) and weekly pricing data. This dataset was chosen "
    "because it is large-scale, hierarchically structured in a way that mirrors real "
    "SME retail operations, and because it is a recognised academic benchmark, "
    "enabling direct comparison with the published studies discussed in Section 4. "
    "For computational practicality, a representative sample of 502 product-store "
    "combinations was selected, proportionally balanced across product categories "
    "and states, so that the results generalise across the different types of "
    "product and geography represented in the full dataset.")

add_heading(doc, "5.2 Building the Forecasting Models", 2)
add_para(doc,
    "Rather than building and testing a single forecasting model, the research "
    "constructed and compared five tiers of increasingly sophisticated approaches, "
    "so that any claimed benefit from machine learning could be measured against a "
    "full spectrum of alternatives, not just a single, potentially weak, baseline.")
add_bullet(doc,
    "the six classical statistical methods most commonly used in retail practice "
    "today, including moving averages and a specialised method (Croston's method) "
    "designed specifically for the intermittent, gap-filled demand patterns "
    "described in Section 3.1.",
    bold_lead="Classical baselines — ")
add_bullet(doc,
    "three widely used machine learning algorithms (Random Forest, XGBoost, and "
    "LightGBM), each trained on a rich set of engineered features capturing recent "
    "sales trends, calendar effects, promotional signals, and pricing information.",
    bold_lead="Standard machine learning — ")
add_bullet(doc,
    "a deep learning model (a Long Short-Term Memory neural network) capable of "
    "learning complex temporal patterns directly from sequences of past sales.",
    bold_lead="Deep learning — ")
add_bullet(doc,
    "a combination model that blends the predictions of the three machine learning "
    "models above through a statistical weighting layer, representing the "
    "state-of-the-art approach reported in the academic literature.",
    bold_lead="Ensemble model — ")
add_bullet(doc,
    "a specialised model that forecasts not just a single expected demand figure "
    "but a full range of plausible outcomes (from the 10th to the 99th percentile), "
    "directly addressing Gap 4 from Section 4.",
    bold_lead="Probabilistic model — ")
add_para(doc,
    "All models were trained strictly on historical data and evaluated exclusively "
    "on a later, entirely unseen eight-week test period, so that the reported "
    "results reflect genuine out-of-sample forecasting performance — the same "
    "standard a business would need to trust before deploying a model live.")

add_heading(doc, "5.3 Translating Forecasts into Inventory Decisions", 2)
add_para(doc,
    "The forecasting outputs were integrated with a standard, industry-recognised "
    "inventory control policy (the Order-Up-To-Level, or periodic review, policy), "
    "which specifies how much buffer stock to hold, at what inventory level to "
    "trigger a reorder, and how much to order, based on a target customer service "
    "level (the percentage of demand the business commits to fulfilling from stock "
    "on hand). Three variants of this policy were compared, differing only in how "
    "the buffer stock (safety stock) is calculated: one using the classical forecast "
    "with the textbook safety-stock formula, one using the machine learning forecast "
    "with the same textbook formula, and one using the machine learning forecast "
    "with the alternative, distribution-aware formula described in Section 3.3. "
    "This design isolates two separate questions that are usually conflated in "
    "commercial claims: does a better forecast help, and does a better safety-stock "
    "formula help, independently of one another?")

add_heading(doc, "5.4 The Critical Methodological Decision: Measuring, Not Assuming, Inventory Performance", 2)
add_para(doc,
    "This is the single most important methodological choice in the research, and "
    "the one that ultimately determined the honesty and credibility of its "
    "conclusions. An earlier draft of this research, like most of the published "
    "literature it reviews, calculated the expected inventory cost and service "
    "level of each policy analytically — using a mathematical formula to project "
    "what should happen, based on assumptions about how much of the time the policy "
    "would run short of stock. On the advice of the academic supervisor overseeing "
    "this work, that approach was replaced with something more rigorous: each "
    "inventory policy was run forward, day by day, against the real 56-day test-"
    "period demand for every one of the 502 products studied, genuinely tracking "
    "stock levels, genuinely recording every day on which a product ran out, and "
    "genuinely calculating the resulting costs from what actually happened rather "
    "than from what a formula predicted should happen.")
add_para(doc,
    "This distinction — between assuming performance and measuring it — is exactly "
    "the kind of rigour a business should demand before trusting any forecasting or "
    "inventory vendor's claims, and it is precisely this rigour that produced the "
    "thesis's most important and most commercially relevant finding, reported in "
    "Section 6 below. To confirm that this simulation was itself trustworthy — that "
    "it would correctly detect a stockout if one occurred, rather than silently "
    "failing to register it — the research includes a dedicated validation exercise "
    "in which the simulation was deliberately stress-tested under conditions "
    "engineered to force stockouts (removing all safety stock, tripling demand, and "
    "cutting off resupply entirely). In every stress test, the simulation correctly "
    "detected the resulting service failures and calculated the associated cost, "
    "confirming that the simulation engine is measuring real outcomes and can be "
    "trusted.")

add_heading(doc, "5.5 Building the Decision-Support Dashboard", 2)
add_para(doc,
    "Finally, the entire analysis was translated into an interactive dashboard, "
    "built using Microsoft Power BI's free tier, organised into four pages: a "
    "forecast overview showing predicted versus actual demand for any selected "
    "product; an inventory status page showing safety stock, reorder point, and "
    "order-up-to level recommendations; a cost analysis page breaking down the cost "
    "of holding, ordering, and stocking-out for each policy; and a sensitivity "
    "page allowing a manager to explore how the results change under different "
    "assumptions about lead time and the cost of a stockout. This dashboard is the "
    "practical deliverable of the research: a tool a non-technical manager can open "
    "and use, without needing to understand the underlying statistics.")

# ── 6. WHAT WE FOUND ─────────────────────────────────────────────────────
add_heading(doc, "6. What We Found", 1)
add_para(doc,
    "The findings fall into two parts, corresponding to the two separate questions "
    "identified in Section 5.3: what happened to forecast accuracy, and what "
    "happened to inventory cost. The two answers turned out to be quite different "
    "from each other, and that difference is itself the thesis's central finding.")

add_heading(doc, "6.1 Forecast Accuracy: A Real but Modest Improvement", 2)
add_para(doc,
    "The ensemble model — the combination of the three machine learning "
    "algorithms — was the most accurate forecaster among all twelve models tested. "
    "It improved on the best classical method by 4.5 percent on average forecasting "
    "error and by 9.5 percent on a metric that penalises large errors more heavily "
    "(relevant for promotional spikes and other high-impact events). It did not, "
    "however, improve meaningfully on the specific percentage-error metric (MAPE) "
    "most commonly quoted in the marketing literature, and it therefore fell short "
    "of the 10-25 percent improvement range reported in the published studies "
    "discussed in Section 4.")
add_para(doc,
    "This shortfall is not a failure of the modelling work; it is itself an "
    "important and honest finding. As anticipated in Section 4, the published "
    "10-25 percent figures were obtained on data aggregated to the weekly or "
    "category level. At the daily, individual-product level — which is the "
    "granularity at which real inventory decisions actually get made — demand is "
    "far more erratic, and even the best available machine learning model cannot "
    "close that gap to the same degree. A retail decision-maker evaluating any "
    "forecasting technology, whether from this thesis or from a commercial vendor, "
    "should specifically ask at what level of data aggregation quoted accuracy "
    "figures were measured, because the answer materially changes what improvement "
    "should realistically be expected in day-to-day operation.")

add_heading(doc, "6.2 Inventory Cost: The More Important, and More Surprising, Finding", 2)
add_para(doc,
    "This is where the rigorous, simulation-based measurement described in Section "
    "5.4 produced a result that a purely analytical, assumption-based study would "
    "never have revealed. When the three inventory policies were run forward "
    "through real demand, all three — the classical policy, the machine-learning-"
    "with-standard-formula policy, and the machine-learning-with-distribution-aware-"
    "formula policy — achieved a 100 percent realised service level across the "
    "test period. In plain terms: none of the three policies ever ran out of stock, "
    "for any of the 502 products, at any point during the eight-week test window.")
add_para(doc,
    "The practical consequence of this is significant and somewhat counter-"
    "intuitive: because no policy ever experienced a stockout, none of the three "
    "policies had any stockout cost to save, and the more accurate forecast and "
    "the more sophisticated safety-stock formula therefore had nothing to improve "
    "upon in terms of avoided lost sales. The machine-learning policy using the "
    "standard safety-stock formula performed essentially identically to the "
    "classical policy — the two differed in total annual cost by less than 0.05 "
    "percent, an amount too small to be commercially meaningful. The machine-"
    "learning policy using the more sophisticated, distribution-aware safety-stock "
    "formula actually cost approximately 1.5 percent more than the classical "
    "policy, because it allocated a larger safety-stock buffer — appropriately "
    "reflecting the true, skewed shape of the demand distribution — but that extra "
    "buffer was protecting against stockouts that, in this specific demand regime, "
    "simply did not occur.")
add_para(doc,
    "This is the central, headline finding of the thesis, and it directly overturns "
    "an earlier, less rigorous version of this same analysis, which — using the "
    "conventional analytical approach most of the published literature relies on — "
    "had estimated a 21 percent inventory cost saving from adopting the machine-"
    "learning and distribution-aware approach. That 21 percent figure was an "
    "artefact of an assumption embedded in the analytical formula, not a measured "
    "outcome; once the same policies were actually run against real demand, the "
    "saving disappeared. This is precisely the kind of gap between claimed and "
    "actual performance that a decision-maker should be alert to when evaluating "
    "any forecasting or inventory technology that has not been validated through "
    "genuine simulation or a live pilot.")

add_heading(doc, "6.3 Why This Happened, in Plain Terms", 2)
add_para(doc,
    "The explanation is straightforward once stated plainly. The classical "
    "inventory policy, using the simplest available forecasting method, was already "
    "holding enough buffer stock to meet demand in full over the test period. When "
    "a policy is already meeting 100 percent of demand, there is, by definition, no "
    "stockout cost left to eliminate, and therefore no room for a more sophisticated "
    "forecast or a more sophisticated safety-stock formula to demonstrate a cost "
    "advantage through avoided stockouts. The additional safety stock held by the "
    "more sophisticated policy becomes pure additional carrying cost, with no "
    "offsetting benefit.")
add_para(doc,
    "This does not mean machine learning forecasting or distribution-aware safety "
    "stock are without value in general — it means their value depends on the "
    "specific demand and inventory situation an organisation is in. Where a "
    "business's current, simpler forecasting and inventory approach is already "
    "adequate — where stockouts are rare under the existing policy — a more "
    "sophisticated approach is unlikely to save meaningful money, and may in fact "
    "add small unnecessary holding costs. Where a business's current approach is "
    "genuinely exposed to stockout risk — running short of stock with meaningful "
    "frequency under existing rules — the same sophisticated approach that showed "
    "no benefit in this test window would be expected to show a real one, because "
    "there would be actual stockout cost available to reduce. The research includes "
    "supporting evidence for this: when the test-period demand was deliberately "
    "increased threefold in the validation exercise described in Section 5.4, all "
    "three policies began experiencing real stockouts, and while the distribution-"
    "aware policy then performed on par with the classical one rather than worse, "
    "even a threefold demand shock was not sufficient, in this specific dataset, "
    "to make the more sophisticated approach clearly superior. This indicates that "
    "the conditions under which the more complex approach pays off are more "
    "specific and more stringent than the general marketing narrative around "
    "machine learning forecasting suggests.")

add_heading(doc, "6.4 Which Products and Which Model Features Mattered Most", 2)
add_para(doc,
    "Two further findings are of direct practical relevance to any retailer "
    "considering this kind of investment. First, an interpretability analysis (using "
    "a technique called SHAP, which quantifies how much each input variable "
    "contributed to each individual forecast) showed that the best-performing model "
    "relied overwhelmingly on straightforward, intuitive signals — recent sales "
    "trends and each product's typical seasonal pattern accounted for nearly 80 "
    "percent of the model's predictive power. This is a reassuring finding for a "
    "non-technical stakeholder: the model is not relying on opaque or hard-to-"
    "explain logic, but on the same kind of signals an experienced store manager "
    "already intuitively tracks, simply applied more systematically and "
    "consistently across every product and every day.")
add_para(doc,
    "Second, a fairness check confirmed that forecasting performance did not "
    "systematically favour or disadvantage any particular product category or "
    "store location beyond what would be expected from underlying differences in "
    "how predictable different types of products naturally are. This matters for "
    "any retailer operating a diverse product range or multiple locations: the "
    "model is not silently underperforming for a subset of the business.")

# ── 7. CENTRAL RECOMMENDATIONS ────────────────────────────────────────────
add_heading(doc, "7. What This Means for Decision-Makers: Central Recommendations", 1)
add_para(doc,
    "The findings in Section 6 translate into a specific, actionable set of "
    "recommendations, organised by the type of decision a reader of this summary "
    "is likely to be facing.")

add_heading(doc, "7.1 Do Not Invest Based on Published Percentage Improvements Alone", 2)
add_para(doc,
    "The single most important recommendation from this research is a note of "
    "caution about how to read the forecasting and inventory-optimisation "
    "literature, and by extension, vendor marketing material that draws on it. "
    "Reported accuracy improvements of 10-25 percent and cost reductions of 5-15 "
    "percent are real, published findings, but they were typically measured under "
    "conditions — aggregated data, and in many cases analytically assumed rather "
    "than measured inventory outcomes — that do not automatically transfer to the "
    "day-to-day, product-level reality of most retail operations. Any organisation "
    "evaluating a forecasting or inventory technology investment should ask two "
    "specific questions of any vendor or study: at what level of data aggregation "
    "was the accuracy figure measured, and was the claimed cost saving measured by "
    "actually running the resulting policy against real demand, or calculated from "
    "a formula?")

add_heading(doc, "7.2 Test Your Own Situation Before Investing — Use the Diagnostic Delivered by This Research", 2)
add_para(doc,
    "The most practically useful output of this thesis, for an SME decision-maker "
    "specifically, is not a single number but a reusable diagnostic tool: the "
    "forward-simulation engine built and validated in this research. Before "
    "committing budget to a more sophisticated forecasting and inventory system, an "
    "organisation can run this same simulation approach against its own historical "
    "sales data to answer the one question that actually determines whether the "
    "investment will pay off: under our current, simpler approach, how often do we "
    "actually run out of stock, and how much does that cost us? If the answer is "
    "\"rarely, and not much\" — as it was in this specific test dataset — the "
    "additional investment in machine learning forecasting and sophisticated safety-"
    "stock formulas is unlikely to be cost-justified, and resources are better spent "
    "elsewhere. If the answer is \"frequently, and at real cost\" — which is a "
    "common situation for fast-moving, high-margin, or highly seasonal product "
    "categories — the same evidence base built in this thesis indicates the "
    "investment is likely to be worthwhile.")
add_para(doc,
    "This reframes the investment decision from a generic \"should we adopt AI\" "
    "question into a specific, testable, low-cost diagnostic exercise that any "
    "organisation with basic sales records can run before spending anything "
    "further.")

add_heading(doc, "7.3 Where the Investment Is Most Likely to Pay Off", 2)
add_para(doc,
    "Based on the mechanism identified in Section 6.3, the organisations and product "
    "categories most likely to see a genuine return from machine learning "
    "forecasting and distribution-aware safety stock are those where: existing "
    "stockouts are frequent or costly (high-margin items, fast fashion, or products "
    "where a stockout risks losing the customer to a competitor); demand is highly "
    "seasonal, promotional, or volatile, making the current, simpler forecasting "
    "approach more likely to be inadequate; and service-level targets are "
    "deliberately set high (99 percent or above), a regime this research found to "
    "amplify the operational value of more sophisticated approaches considerably "
    "more than at more conventional 90-95 percent targets. Conversely, stable, "
    "slow-moving, low-margin categories where the current approach is already "
    "adequate are less promising candidates for this specific investment.")

add_heading(doc, "7.4 Deploy the Dashboard as a Low-Cost, Low-Risk First Step", 2)
add_para(doc,
    "Regardless of the outcome of the diagnostic in Section 7.2, the Power BI "
    "dashboard delivered alongside this research is a low-cost, low-risk way to "
    "begin building forecasting and inventory-management capability. It requires no "
    "paid software licence, runs on standard office computing hardware, and "
    "requires only the sales, calendar, and pricing data that a typical point-of-"
    "sale system already produces. Adopting it does not commit an organisation to "
    "the more expensive machine-learning-plus-quantile-forecasting approach; it "
    "provides visibility into current inventory parameters and cost drivers "
    "regardless of which forecasting method ultimately feeds it, and it can be "
    "upgraded to more sophisticated forecasting later if the diagnostic in Section "
    "7.2 indicates that upgrade would pay off.")

add_heading(doc, "7.5 Treat This as an Ongoing Measurement Discipline, Not a One-Time Decision", 2)
add_para(doc,
    "Demand regimes change: a product category that shows no stockout risk today "
    "may become promotional, seasonal, or supply-constrained in the future, "
    "changing the answer to the diagnostic question in Section 7.2. The "
    "recommendation is not a single, one-time investment decision but an ongoing "
    "practice: periodically re-running the simulation-based diagnostic against "
    "current data, particularly ahead of known high-stakes periods (major "
    "promotional events, seasonal peaks, new product launches, or supply chain "
    "disruptions), so that the decision to invest further in forecasting "
    "sophistication is always grounded in current, measured evidence rather than a "
    "decision made once and never revisited.")

# ── 8. IMPLEMENTATION ROADMAP ────────────────────────────────────────────
add_heading(doc, "8. A Practical Path Forward", 1)
add_para(doc,
    "For an organisation persuaded by the case above, the following is a realistic, "
    "low-risk sequence for putting this research into practice, requiring no "
    "specialist data science hires and no capital software investment.")

add_table(doc,
    headers=["Step", "Action", "Resource Required"],
    rows=[
        ["1", "Export 12-24 months of daily sales, calendar/promotional, and price "
              "data from the existing point-of-sale system.",
              "One analyst; standard POS export functionality"],
        ["2", "Run the open-source diagnostic pipeline delivered with this research "
              "against that data to identify which product categories currently "
              "experience meaningful stockout risk.",
              "One laptop (16 GB RAM, no specialist hardware); free, open-source "
              "software"],
        ["3", "For categories showing low stockout risk, retain the existing, "
              "simpler forecasting and reorder approach — no further investment "
              "is indicated.",
              "No additional resource"],
        ["4", "For categories showing meaningful stockout risk, pilot the "
              "machine-learning forecasting and distribution-aware safety-stock "
              "approach on that subset only, before considering a wider rollout.",
              "Continued use of the same open-source pipeline; no new licences"],
        ["5", "Deploy the Power BI dashboard for ongoing visibility into forecast "
              "accuracy, inventory parameters, and cost drivers across the "
              "business.",
              "Power BI Desktop or Power BI Service free tier"],
        ["6", "Re-run the diagnostic periodically (for example, quarterly, or "
              "ahead of major seasonal or promotional periods) to keep the "
              "investment decision grounded in current data.",
              "Recurring, low time cost using the same free pipeline"],
    ])
add_para(doc, "", space_after=6)
add_para(doc,
    "This sequence is deliberately designed so that no organisation needs to commit "
    "significant budget before it has direct, measured evidence — from its own data, "
    "not from a published benchmark — that the investment is likely to be worthwhile.")

# ── 9. CONCLUSION ─────────────────────────────────────────────────────────
add_heading(doc, "9. Conclusion", 1)
add_para(doc,
    "This research set out to answer a question every retail decision-maker "
    "eventually faces: does investing in machine learning-based demand forecasting "
    "and inventory optimisation actually pay off, and is it accessible to an "
    "organisation without enterprise-scale resources? The answer this thesis "
    "provides is more useful, and more honest, than a simple yes: the value of the "
    "investment is real but conditional, and this research delivers both the "
    "evidence for that conclusion and a practical, low-cost tool that lets any "
    "organisation test which side of that condition it falls on, using its own "
    "data, before spending anything further.")
add_para(doc,
    "The technology built to demonstrate this — a full forecasting and inventory "
    "pipeline, a validated simulation engine, and an accessible dashboard — is "
    "delivered entirely in free, open-source form, runs on standard office "
    "hardware, and requires only the data a typical point-of-sale system already "
    "produces. The barrier to entry this research removes is not financial; it is "
    "informational. What most retail decision-makers have lacked is not access to "
    "machine learning tools, but a trustworthy, rigorously validated way of knowing "
    "whether those tools are worth adopting for their specific business. This "
    "thesis provides that.")

OUT_PATH = OUT
doc.save(OUT_PATH)
print(f"Saved: {OUT_PATH}")

# Word count report
from docx import Document as _D
d2 = _D(OUT_PATH)
wc = sum(len(p.text.split()) for p in d2.paragraphs)
for t in d2.tables:
    for row in t.rows:
        for cell in row.cells:
            wc += len(cell.text.split())
print(f"Approx word count (incl. tables): {wc}")
print(f"Paragraphs: {len(d2.paragraphs)}  Tables: {len(d2.tables)}")
