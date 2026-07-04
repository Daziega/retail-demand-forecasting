"""Generate Chapter 7 - Conclusion."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_para(doc, text, *, italic=False, bold=False, indent_left=None, font_size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if indent_left:
        p.paragraph_format.left_indent = Inches(indent_left)
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(font_size)
    run.italic = italic
    run.bold = bold
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)
    return p


def add_table_caption(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(10)


def add_simple_table(doc, headers, rows, header_fill="D5E8F0"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.autofit = True

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(10)
        set_cell_shading(hdr_cells[i], header_fill)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for r_idx, row in enumerate(rows, start=1):
        cells = table.rows[r_idx].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ""
            p = cells[c_idx].paragraphs[0]
            run = p.add_run(str(val))
            run.font.name = "Arial"
            run.font.size = Pt(10)
            cells[c_idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return table


# ─────────────────────────────────────────────────────────────────────────────
doc = Document()

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

style = doc.styles["Normal"]
style.font.name = "Arial"
style.font.size = Pt(11)

# ── TITLE ─────────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("CHAPTER 7")
run.bold = True
run.font.name = "Arial"
run.font.size = Pt(20)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Conclusion")
run.bold = True
run.font.name = "Arial"
run.font.size = Pt(16)
subtitle.paragraph_format.space_after = Pt(18)

# ── 7.1 INTRODUCTION ─────────────────────────────────────────────────────
add_heading(doc, "7.1 Introduction", 1)
add_para(doc,
    "This chapter concludes the thesis. It returns to the research objectives stated in "
    "Chapter 1, evaluates the extent to which each has been achieved through the empirical "
    "work documented in Chapters 3 to 6, restates the principal scholarly and practical "
    "contributions of the thesis, and offers closing reflections on the operational "
    "significance of the findings for the small and medium-sized retail sector. The chapter is "
    "structured as follows. Section 7.2 restates the five research objectives. Section 7.3 "
    "evaluates achievement against each objective. Section 7.4 consolidates the principal "
    "contributions of the thesis. Section 7.5 offers closing remarks and identifies the broader "
    "significance of the work.")

# ── 7.2 RESTATING THE OBJECTIVES ─────────────────────────────────────────
add_heading(doc, "7.2 Restatement of the Research Objectives", 1)
add_para(doc,
    "Chapter 1 set out five research objectives that have guided the thesis:")
add_bullet(doc,
    "Objective 1: Conduct a comparative evaluation of forecasting approaches spanning classical "
    "statistical, standard machine learning, deep learning, ensemble, and probabilistic "
    "methods on a hierarchical retail demand benchmark.")
add_bullet(doc,
    "Objective 2: Develop a systematic feature engineering strategy capturing the principal "
    "drivers of retail demand — lags, rolling statistics, calendar and event signals, price "
    "movement, and identifier-level prior information.")
add_bullet(doc,
    "Objective 3: Integrate the best-performing forecasting approach with the Order-Up-To-"
    "Level (OUTL) inventory policy and quantify the resulting operational parameters across "
    "lead-time scenarios.")
add_bullet(doc,
    "Objective 4: Quantify the operational impact of ML-driven forecasting relative to "
    "classical baselines, targeting improvements of 10–25% in forecast accuracy (measured by "
    "MAPE, MAE, and RMSE) and inventory cost reductions of 5–15% through more precise demand "
    "predictions.")
add_bullet(doc,
    "Objective 5: Develop a Power BI dashboard that translates forecasting outputs and "
    "inventory recommendations into an accessible, interactive decision-support tool suitable "
    "for deployment in resource-constrained retail environments, including small and medium-"
    "sized enterprises.")

# ── 7.3 ACHIEVEMENT ──────────────────────────────────────────────────────
add_heading(doc, "7.3 Achievement Against Research Objectives", 1)

add_heading(doc, "7.3.1 Objective 1 — Comparative evaluation of forecasting approaches", 2)
add_para(doc,
    "The thesis evaluated twelve forecasting models organised across five tiers on a stratified "
    "subsample of the M5-Forecasting (Walmart) dataset. Tier 1 implemented six classical "
    "baselines (Moving Average, Simple Exponential Smoothing, AutoARIMA, Croston's Classic, "
    "Croston-SBA, and Seasonal Naive); Tier 2 implemented three standard ML models (LightGBM "
    "and XGBoost with Tweedie loss, and Random Forest); Tier 3 implemented a stacked LSTM with "
    "log1p target transformation; Tier 4 implemented a stacking ensemble of the Tier 2 base "
    "learners through a non-negative Ridge meta-learner; and Tier 5 implemented a five-quantile "
    "LightGBM model supporting the inventory analysis.")
add_para(doc,
    "The comparative evaluation produced a ranked leaderboard (Table 4.7) demonstrating that "
    "the Tier 4 stacking ensemble achieves the best MAE (0.952) and WMAPE (73.58%) across all "
    "models. The empirical ranking reproduces the consensus of the modern retail forecasting "
    "literature: tree-based gradient boosting and ensembles thereof at the top, classical "
    "baselines (with Croston-class methods particularly well-suited to intermittent demand) in "
    "the middle, and the LSTM at the bottom — a finding that directly replicates Nasseri et "
    "al. (2023). Objective 1 is fully achieved.")

add_heading(doc, "7.3.2 Objective 2 — Systematic feature engineering", 2)
add_para(doc,
    "The thesis constructed thirty-three engineered features organised across six categories "
    "(lag, rolling, temporal, event proximity, promotional/event, price) plus eight target-"
    "encoded mean features for high-cardinality categoricals. The feature engineering strategy "
    "is documented to the level required for independent replication in §3.5 and §3.4.6.")
add_para(doc,
    "SHAP analysis on the LightGBM model (§4.11) demonstrates that the engineered features are "
    "the principal predictive drivers: the top five SHAP-ranked features account for 78.6% of "
    "model predictions, dominated by rolling means (rolling_mean_7 at 29.9%, rolling_mean_28 at "
    "20.4%), target-encoded item-level seasonality (item_month_mean at 15.3%, item_dow_mean at "
    "5.3%), and demand volatility (rolling_std_7 at 7.7%). Identifier features contribute "
    "marginally on their own (each <0.5%), confirming that the target-encoded mean features "
    "absorb the cross-sectional signal more efficiently than raw identifiers and validating the "
    "methodological choice to add target encoding to the feature set. Objective 2 is fully "
    "achieved.")

add_heading(doc, "7.3.3 Objective 3 — OUTL inventory integration", 2)
add_para(doc,
    "The thesis integrated the Tier 4 stacking ensemble (for point forecasts) and the Tier 5 "
    "quantile model (for empirical safety stock) with a periodic (R, s, S) review policy. The "
    "OUTL parameters — safety stock, reorder point, order-up-to level, order quantity, and "
    "annual cost components — are computed analytically for each of the 502 product-store "
    "series across four lead-time scenarios (L = 7, 10, 14, 21 days) and three service levels "
    "(90%, 95%, 99%), yielding 6,024 distinct OUTL configurations per forecasting policy.")
add_para(doc,
    "The methodology is implemented through the equations of Silver, Pyke and Thomas (2017), "
    "with the order quantity computed via the Economic Order Quantity formula to decouple "
    "ordering frequency from forecast level. Three forecasting policies are compared on a like-"
    "for-like basis: Classical (Croston-SBA with Gaussian SS), ML-Gaussian (ensemble with "
    "Gaussian SS), and ML-Empirical-Quantile (ensemble with quantile-derived SS). Objective 3 "
    "is fully achieved.")

add_heading(doc, "7.3.4 Objective 4 — Quantification of operational impact", 2)
add_para(doc,
    "Objective 4 specified two quantitative performance targets: 10–25% improvement in MAPE "
    "(and MAE/RMSE) and 5–15% reduction in inventory cost. The empirical results against each "
    "target are nuanced and warrant detailed restatement.")
add_para(doc,
    "On forecast accuracy, the stacking ensemble achieves 4.5% MAE improvement and 9.5% RMSE "
    "improvement over the best classical baseline, with the MAPE improvement essentially flat "
    "(−0.3%). The MAE and MAPE results sit below the 10–25% target band; the RMSE result is "
    "close to the lower bound. The reasons for these shortfalls — granularity-dependence of "
    "forecast accuracy, the inherent intermittence of SKU-daily retail demand, and the upper-"
    "bound nature of the literature-derived target figures — are documented in §4.10 and §6.2. "
    "The accuracy target is therefore partially achieved, with explicit acknowledgement that "
    "double-digit MAPE improvements at SKU-daily granularity are at the upper end of what is "
    "achievable on intermittent demand benchmarks.")
add_para(doc,
    "On inventory cost reduction, the ML-Empirical-Quantile policy achieves a 21% reduction "
    "in total annual cost relative to the Classical baseline at the central scenario (L = 14, "
    "SL = 95%, baseline stockout cost multiplier m = 1.0), with the reduction robust across "
    "all four lead-time scenarios (20.8–21.2%) and rising to 27% at the sensitivity-analysis "
    "central case (MAPE = 10%). The cost target is exceeded by a substantial margin and is "
    "robust even under the most conservative stockout cost assumption (15% reduction at m = "
    "0.4, at the upper end of the target band). The inventory cost component of Objective 4 "
    "is fully achieved and exceeded.")
add_para(doc,
    "On balance, Objective 4 is achieved on the operational dimension (inventory cost) and "
    "partially achieved on the accuracy dimension (forecast metrics). The discussion in §6.5 "
    "argues that this asymmetric outcome is itself a contribution to the literature: it "
    "demonstrates that the operational value of ML forecasting in retail inventory management "
    "is mediated by the inventory formulation rather than by the forecasting model alone, and "
    "that pursuing further MAPE improvement at the SKU-daily level has limited operational "
    "return compared to investment in the inventory-decision layer.")

add_heading(doc, "7.3.5 Objective 5 — Power BI dashboard for SME deployment", 2)
add_para(doc,
    "The thesis implemented a Power BI dashboard with four pages (Forecast Overview, Inventory "
    "Status, Cost Analysis, Sensitivity Analysis), consuming five CSV data feeds exported from "
    "the Python pipeline and refreshed on a daily cycle. The dashboard is implemented in "
    "Microsoft Power BI Desktop (free tier), requires no commercial licence, and is delivered "
    "as a single .pbix file that can be opened in any installation of Power BI Desktop without "
    "further dependencies.")
add_para(doc,
    "Each page is designed to support a distinct operational decision: Forecast Overview "
    "surfaces actual demand, ML forecasts, classical baseline forecasts, and the LightGBM "
    "quantile fan for any selected product-store combination; Inventory Status presents safety "
    "stock, reorder point, and order-up-to-level parameters with conditional formatting for "
    "items approaching reorder; Cost Analysis decomposes total annual inventory cost by policy "
    "and lead time; Sensitivity Analysis exposes the four-dimensional sensitivity curves from "
    "§5.7 to enable practitioner what-if analysis. The dashboard supports the SME "
    "transferability commitment of the thesis by translating the modelling outputs into a form "
    "accessible to non-technical managers. Objective 5 is fully achieved.")

# ── 7.4 PRINCIPAL CONTRIBUTIONS ──────────────────────────────────────────
add_heading(doc, "7.4 Principal Contributions", 1)
add_para(doc,
    "The thesis advances the demand forecasting and inventory optimisation literature through "
    "three theoretical contributions and three practical contributions, each summarised below "
    "and discussed in detail in §6.5 and §6.6.")
add_para(doc,
    "Theoretical contribution 1: First quantitative measurement on the M5 benchmark of the "
    "Gaussian safety-stock under-allocation factor for intermittent retail demand. The "
    "empirical 95th-percentile demand spread is 1.36 times the Gaussian-implied buffer, "
    "corresponding to an effective service level of approximately 83% when 95% is nominally "
    "targeted. This finding extends the textbook treatment of safety stock (Silver et al., "
    "2017), which acknowledges the Gaussian assumption's validity conditions but has not "
    "previously been quantified on a standard retail benchmark to the author's knowledge.")
add_para(doc,
    "Theoretical contribution 2: Demonstration that the operational value of ML forecasting in "
    "retail inventory management is mediated by the inventory formulation, not by the "
    "forecasting model alone. The ML-Gaussian policy improves over the Classical-Gaussian "
    "policy by only 0.02% of annual cost — a negligible benefit — whereas the ML-Empirical-"
    "Quantile policy improves by 21%. This finding refines the literature's framing of "
    "forecasting-inventory integration (Seyedan et al., 2023) and identifies the inventory-"
    "formulation choice as the dominant operational lever.")
add_para(doc,
    "Theoretical contribution 3: Empirical robustness of the inventory cost reduction to "
    "forecast accuracy level. Across the 3–15% MAPE sensitivity range, the cost reduction "
    "varies by less than 0.1 percentage points. This contribution has implications for the "
    "ordering of research priorities in retail demand forecasting: the next decimal point of "
    "MAPE is less operationally valuable than improvements in the inventory-decision layer.")
add_para(doc,
    "Practical contribution 1: An end-to-end, reproducible, SME-deployable framework "
    "implemented in free and open-source tooling, runnable on a standard laptop, with a total "
    "runtime under three hours. The framework is documented to the level required for "
    "independent replication.")
add_para(doc,
    "Practical contribution 2: An interactive Power BI decision-support dashboard organised "
    "into four operational views, providing non-technical SME managers with access to forecast "
    "outputs, inventory parameters, cost decomposition, and sensitivity analysis curves "
    "without requiring the user to understand the underlying methodology.")
add_para(doc,
    "Practical contribution 3: A defensible quantitative argument for SME forecasting "
    "investment. The headline 21% cost reduction is robust to MAPE level and achieves the "
    "Chapter 1 target band (5–15%) even under the most conservative stockout cost assumption. "
    "The sensitivity analysis curves (§5.7) provide the basis for adapting the headline figure "
    "to the specific operational parameters of any individual SME.")

# ── 7.5 CLOSING REMARKS ──────────────────────────────────────────────────
add_heading(doc, "7.5 Closing Remarks", 1)
add_para(doc,
    "The thesis began with a question that motivated both its academic positioning and its "
    "practical orientation: whether machine learning-driven demand forecasting, integrated "
    "with classical inventory policy theory, can produce operationally meaningful cost "
    "reductions in retail environments where small and medium-sized enterprises must compete "
    "with the analytics infrastructure of larger competitors. The answer that has emerged from "
    "the empirical work is qualified but affirmative: machine learning forecasting alone, "
    "applied at SKU-daily granularity to intermittent retail demand, yields modest accuracy "
    "improvement and consequently modest savings under the textbook Gaussian inventory "
    "formulation; but the same machine learning toolkit, used to produce quantile forecasts "
    "feeding an empirically calibrated safety stock, produces operational cost reductions of "
    "20% or more — a result robust across lead times, service levels, and reasonable variation "
    "in stockout cost assumptions.")
add_para(doc,
    "This finding reframes the operational case for ML in retail inventory management. The "
    "conventional argument — that better forecasts produce better inventory decisions through "
    "improved point-forecast accuracy — is empirically weak at the granularity at which "
    "inventory decisions are actually made. The stronger argument — supported by the data — is "
    "that the value of ML forecasting in retail inventory management is realised at the "
    "intersection of probabilistic forecasting and inventory-formulation refinement. The "
    "quantile model is the bridge between the two: it allows the textbook (R, s, S) policy to "
    "be applied with empirically calibrated rather than parametric safety stock, closing a gap "
    "in the inventory theory literature that has been acknowledged but not previously "
    "quantified on a standard benchmark.")
add_para(doc,
    "For the small and medium-sized retail sector specifically, the practical implication is "
    "encouraging. The framework demonstrated in this thesis is implementable with free and "
    "open-source tooling on a standard laptop, requires only the data that off-the-shelf "
    "point-of-sale systems already produce, and achieves the Chapter 1 cost reduction target "
    "even under the most conservative stockout cost assumption. The barriers to adoption are "
    "therefore not infrastructural, computational, or financial; they are organisational and "
    "skills-related. The Power BI dashboard is designed specifically to lower the latter "
    "barrier by translating the framework's outputs into a form accessible to non-technical "
    "managers.")
add_para(doc,
    "The thesis has therefore addressed both the academic question — how should ML forecasting "
    "be integrated with inventory policy theory for intermittent retail demand — and the "
    "practical question — can such integration be made operationally accessible to SMEs. The "
    "answer to both, with the qualifications documented in Chapter 6 and the limitations "
    "documented in §5.10 and §6.8, is affirmative. Future work should extend the framework to "
    "live deployment, validate it in actual SME contexts, and explore whether the Gaussian "
    "under-allocation factor of 1.36 measured on M5 generalises across other retail datasets "
    "and other categories of intermittent demand.")
add_para(doc,
    "The wider research programme to which this thesis contributes — the operational "
    "integration of probabilistic machine learning with classical inventory theory — sits at "
    "the intersection of two mature literatures (demand forecasting and inventory optimisation) "
    "that have, for historical reasons, evolved largely in parallel. The opportunity to close "
    "the gap between them, both academically and operationally, is significant. The thesis "
    "offers one concrete step in that direction, in the specific context of small and medium-"
    "sized retail. The framework, the methodology, and the empirical findings are delivered in "
    "a form intended to support both academic extension and operational deployment.")

OUT = "/Users/desmond/Capstone Project/retail-demand-forecasting/Chapter_7_Conclusion.docx"
doc.save(OUT)
print(f"Saved: {OUT}")
